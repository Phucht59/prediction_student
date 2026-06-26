from __future__ import annotations

from pathlib import Path
import csv
import json

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports" / "final" / "KLTN_BAO_CAO_THEO_MAU_2026.docx"
FIG_DIR = ROOT / "reports" / "final" / "figures"


TITLE = "XÂY DỰNG MÔ HÌNH HỌC KẾT HỢP ĐỂ DỰ ĐOÁN THÀNH TÍCH HỌC TẬP SINH VIÊN"
SCHOOL = "TRƯỜNG ĐẠI HỌC NGOẠI NGỮ - TIN HỌC TP. HỒ CHÍ MINH"
FACULTY = "KHOA CÔNG NGHỆ THÔNG TIN"


def set_run_font(run, size: float | None = None, bold: bool | None = None, italic: bool | None = None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def chart_font(size: int, bold: bool = False):
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\timesbd.ttf" if bold else r"C:\Windows\Fonts\times.ttf",
        r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def text_size(draw: ImageDraw.ImageDraw, text: str, font) -> tuple[int, int]:
    box = draw.textbbox((0, 0), text, font=font)
    return box[2] - box[0], box[3] - box[1]


def draw_centered(draw, box, text, font, fill=(20, 20, 20), linespacing=8):
    x1, y1, x2, y2 = box
    max_width = x2 - x1 - 28
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        trial = (current + " " + word).strip()
        if text_size(draw, trial, font)[0] <= max_width or not current:
            current = trial
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    heights = [text_size(draw, line, font)[1] for line in lines]
    total_h = sum(heights) + linespacing * (len(lines) - 1)
    y = y1 + (y2 - y1 - total_h) / 2
    for line, h in zip(lines, heights):
        w = text_size(draw, line, font)[0]
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=font, fill=fill)
        y += h + linespacing


def load_prediction_metrics() -> list[dict[str, str]]:
    text = (ROOT / "reports" / "final" / "FINAL_PROJECT_STATUS.md").read_text(encoding="utf-8")
    rows: list[dict[str, str]] = []
    for line in text.splitlines():
        if not line.startswith("|"):
            continue
        cols = [c.strip() for c in line.strip("|").split("|")]
        if not cols or cols[0] in {"Dataset", "---"} or set(cols[0]) == {"-"}:
            continue
        if len(cols) == 7:
            rows.append(
                {
                    "dataset": cols[0],
                    "scenario": cols[1],
                    "model": cols[2],
                    "prediction_mode": cols[3],
                    "macro_f1": cols[4],
                    "recall_low": cols[5],
                    "f1_low": cols[6],
                }
            )
    return rows


def load_baseline_rows() -> list[dict[str, str]]:
    path = ROOT / "reports" / "final" / "final_baseline_comparison.csv"
    with path.open(newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def load_recommender_metrics() -> list[dict[str, float | str]]:
    rows = []
    for dataset in ["xapi", "student-por"]:
        path = ROOT / "outputs" / "recommender" / dataset / "recommender_metrics.json"
        data = json.loads(path.read_text(encoding="utf-8"))
        rows.append(
            {
                "dataset": dataset,
                "risk_macro_f1": data["risk_diagnosis"]["f1_macro"],
                "risk_micro_f1": data["risk_diagnosis"]["f1_micro"],
                "precision_at_3": data["ranking"]["precision_at_3"],
                "recall_at_3": data["ranking"]["recall_at_3"],
                "ndcg_at_3": data["ranking"]["ndcg_at_3"],
                "coverage_at_3": data["ranking"]["coverage_at_3"],
                "risk_coverage": data["path_quality"]["risk_coverage_rate"],
                "difficulty_progression": data["path_quality"]["difficulty_progression_rate"],
                "prereq_violation": data["path_quality"]["prerequisite_violation_rate"],
            }
        )
    return rows


def save_pipeline_figure(path: Path):
    w, h = 1700, 760
    img = Image.new("RGB", (w, h), "#FFFFFF")
    d = ImageDraw.Draw(img)
    title_f = chart_font(42, True)
    label_f = chart_font(25, True)
    small_f = chart_font(20)
    d.text((70, 45), "Quy trình dữ liệu - xử lý - mô hình - khuyến nghị", font=title_f, fill="#111827")
    d.text((72, 103), "Hình được sinh bằng code từ thiết kế pipeline final trong repo, không phải ảnh minh họa bên ngoài.", font=small_f, fill="#4B5563")
    boxes = [
        ("Dữ liệu", "student-mat\nstudent-por\nxAPI", "#DBEAFE"),
        ("Tiền xử lý", "feature setting\nencoding\ntrain pool", "#E0F2FE"),
        ("Đánh giá", "CV/OOF\nlocked test\nno leakage", "#DCFCE7"),
        ("CNN-BiLSTM", "sequence branch\nattention/pooling\ngated fusion xAPI", "#FEF3C7"),
        ("Kết quả", "Macro F1\nRecall Low\nF1 Low", "#FCE7F3"),
        ("RA-HLPR", "risk diagnosis\nintervention ranking\n4-week path", "#EDE9FE"),
    ]
    x0, y0, bw, bh, gap = 65, 245, 238, 250, 35
    for i, (head, body, color) in enumerate(boxes):
        x = x0 + i * (bw + gap)
        d.rounded_rectangle((x, y0, x + bw, y0 + bh), radius=24, fill=color, outline="#334155", width=3)
        d.text((x + 22, y0 + 28), head, font=label_f, fill="#0F172A")
        yy = y0 + 82
        for line in body.split("\n"):
            d.text((x + 22, yy), line, font=small_f, fill="#1F2937")
            yy += 42
        if i < len(boxes) - 1:
            ax = x + bw + 8
            ay = y0 + bh // 2
            d.line((ax, ay, ax + gap - 16, ay), fill="#334155", width=4)
            d.polygon([(ax + gap - 16, ay - 10), (ax + gap - 16, ay + 10), (ax + gap - 1, ay)], fill="#334155")
    d.rounded_rectangle((70, 565, 1630, 680), radius=20, fill="#F8FAFC", outline="#CBD5E1", width=2)
    note = "Guardrails: không dùng student-combine, không chọn model bằng locked test, threshold tuning bằng CV/OOF, baseline chỉ dùng đối chứng."
    draw_centered(d, (90, 585, 1610, 660), note, small_f, fill="#0F172A")
    img.save(path)


def save_model_architecture_figure(path: Path):
    w, h = 1600, 860
    img = Image.new("RGB", (w, h), "#FFFFFF")
    d = ImageDraw.Draw(img)
    title_f = chart_font(42, True)
    label_f = chart_font(24, True)
    small_f = chart_font(20)
    d.text((70, 45), "Kiến trúc mô hình CNN-BiLSTM final", font=title_f, fill="#111827")
    d.text((72, 104), "Student dùng sequence-only; xAPI dùng gated fusion khi có thêm context/categorical features.", font=small_f, fill="#4B5563")
    def box(x, y, ww, hh, title, body, fill):
        d.rounded_rectangle((x, y, x + ww, y + hh), radius=22, fill=fill, outline="#334155", width=3)
        d.text((x + 20, y + 20), title, font=label_f, fill="#0F172A")
        yy = y + 70
        for line in body.split("\n"):
            d.text((x + 20, yy), line, font=small_f, fill="#1F2937")
            yy += 36
    def arrow(x1, y1, x2, y2):
        d.line((x1, y1, x2, y2), fill="#334155", width=4)
        if abs(x2 - x1) > abs(y2 - y1):
            d.polygon([(x2, y2), (x2 - 16, y2 - 9), (x2 - 16, y2 + 9)], fill="#334155")
        else:
            d.polygon([(x2, y2), (x2 - 9, y2 - 16), (x2 + 9, y2 - 16)], fill="#334155")
    box(80, 190, 290, 160, "Sequence input", "Student: G1/G2\nxAPI: behavior signals", "#DBEAFE")
    box(450, 190, 250, 160, "Conv1D", "local pattern\nshort sequence fit", "#E0F2FE")
    box(780, 190, 260, 160, "BiLSTM", "two-direction\nsequence encoding", "#DCFCE7")
    box(1120, 190, 280, 160, "Pooling", "attention/last hidden\nsequence vector", "#FEF3C7")
    arrow(370, 270, 450, 270)
    arrow(700, 270, 780, 270)
    arrow(1040, 270, 1120, 270)
    box(460, 510, 320, 150, "Context features", "categorical + profile\nused mainly for xAPI", "#FCE7F3")
    box(880, 480, 300, 190, "Gated fusion", "gate = sigmoid(W[seq, ctx])\nfinal xAPI: gated_fusion_v28", "#EDE9FE")
    arrow(1260, 350, 1030, 480)
    arrow(780, 585, 880, 585)
    box(1270, 500, 230, 150, "Classifier", "p_low\np_medium\np_high", "#F8FAFC")
    arrow(1180, 585, 1270, 585)
    d.rounded_rectangle((80, 720, 1500, 795), radius=18, fill="#F8FAFC", outline="#CBD5E1", width=2)
    draw_centered(d, (95, 728, 1485, 785), "Final champion: student-mat/student-por sequence_cnn_bilstm_only; xAPI gated_fusion_v28 + low_f1_tuned.", small_f)
    img.save(path)


def save_recommender_flow_figure(path: Path):
    w, h = 1600, 760
    img = Image.new("RGB", (w, h), "#FFFFFF")
    d = ImageDraw.Draw(img)
    title_f = chart_font(42, True)
    label_f = chart_font(24, True)
    small_f = chart_font(19)
    d.text((70, 45), "RA-HLPR: từ xác suất dự đoán đến lộ trình 4 tuần", font=title_f, fill="#111827")
    d.text((72, 104), "Module downstream, không collaborative filtering và không dùng true G3/Class để sinh khuyến nghị vận hành.", font=small_f, fill="#4B5563")
    top = [
        ("CNN-BiLSTM\nprobabilities", "p_low\np_medium\np_high", "#DBEAFE"),
        ("RiskDiagnosisHead", "observable signals\nrisk probabilities", "#DCFCE7"),
        ("CandidateGenerator", "dataset-aware\nprediction-aware", "#FEF3C7"),
        ("HybridScorer", "risk_match\nperformance_need\nexpected_effect", "#FCE7F3"),
        ("PathPlanner", "Stabilize\nPractice\nReinforce\nEvaluate", "#EDE9FE"),
    ]
    x0, y0, bw, bh, gap = 80, 210, 260, 210, 55
    for i, (head, body, fill) in enumerate(top):
        x = x0 + i * (bw + gap)
        d.rounded_rectangle((x, y0, x + bw, y0 + bh), radius=24, fill=fill, outline="#334155", width=3)
        draw_centered(d, (x + 15, y0 + 20, x + bw - 15, y0 + 92), head, label_f, fill="#0F172A")
        draw_centered(d, (x + 20, y0 + 105, x + bw - 20, y0 + bh - 20), body, small_f, fill="#1F2937")
        if i < len(top) - 1:
            ax = x + bw + 8
            ay = y0 + bh // 2
            d.line((ax, ay, ax + gap - 18, ay), fill="#334155", width=4)
            d.polygon([(ax + gap - 18, ay - 10), (ax + gap - 18, ay + 10), (ax + gap - 2, ay)], fill="#334155")
    d.rounded_rectangle((110, 515, 1490, 675), radius=22, fill="#F8FAFC", outline="#CBD5E1", width=2)
    draw_centered(
        d,
        (135, 535, 1465, 650),
        "Output: mỗi sinh viên nhận risk_band, plan_intensity, top risks, top interventions và lộ trình 4 tuần có objective, actions, expected outcome.",
        small_f,
        fill="#0F172A",
    )
    img.save(path)


def save_grouped_bar(path: Path, title: str, subtitle: str, categories: list[str], series: dict[str, list[float]], colors: dict[str, str], ymax: float = 1.0):
    w, h = 1700, 950
    img = Image.new("RGB", (w, h), "#FFFFFF")
    d = ImageDraw.Draw(img)
    title_f = chart_font(42, True)
    sub_f = chart_font(21)
    axis_f = chart_font(19)
    small_f = chart_font(18)
    d.text((70, 45), title, font=title_f, fill="#111827")
    d.text((72, 104), subtitle, font=sub_f, fill="#4B5563")
    left, top, right, bottom = 140, 185, 1610, 790
    d.line((left, top, left, bottom), fill="#334155", width=3)
    d.line((left, bottom, right, bottom), fill="#334155", width=3)
    for tick in [0, 0.25, 0.5, 0.75, 1.0]:
        y = bottom - (bottom - top) * tick / ymax
        d.line((left - 8, y, right, y), fill="#E5E7EB", width=1)
        d.text((55, y - 12), f"{tick:.2f}", font=axis_f, fill="#374151")
    n_cat = len(categories)
    n_series = len(series)
    group_w = (right - left) / n_cat
    bar_w = min(70, group_w / (n_series + 1.4))
    for ci, cat in enumerate(categories):
        gx = left + ci * group_w + group_w * 0.12
        for si, (name, values) in enumerate(series.items()):
            val = values[ci]
            x1 = gx + si * (bar_w + 10)
            x2 = x1 + bar_w
            y1 = bottom - (bottom - top) * val / ymax
            d.rounded_rectangle((x1, y1, x2, bottom), radius=8, fill=colors[name], outline=None)
            d.text((x1 - 2, y1 - 26), f"{val:.3f}", font=small_f, fill="#111827")
        label = cat.replace(" ", "\n")
        tw, th = text_size(d, label, axis_f)
        d.multiline_text((left + ci * group_w + group_w / 2 - tw / 2, bottom + 18), label, font=axis_f, fill="#111827", align="center", spacing=4)
    lx, ly = left, 835
    for name in series:
        if lx > right - 210:
            lx = left
            ly += 45
        d.rounded_rectangle((lx, ly, lx + 30, ly + 20), radius=4, fill=colors[name])
        d.text((lx + 42, ly - 3), name, font=axis_f, fill="#111827")
        lx += max(210, text_size(d, name, axis_f)[0] + 90)
    img.save(path)


def save_metric_figures():
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    save_pipeline_figure(FIG_DIR / "fig_3_1_pipeline_overview.png")
    save_model_architecture_figure(FIG_DIR / "fig_3_2_cnn_bilstm_architecture.png")
    save_recommender_flow_figure(FIG_DIR / "fig_5_2_ra_hlpr_flow.png")

    pred = load_prediction_metrics()
    categories = [f"{r['dataset']} {r['scenario']}" for r in pred]
    save_grouped_bar(
        FIG_DIR / "fig_4_1_prediction_metrics.png",
        "Kết quả locked test của mô hình deep final",
        "Nguồn: reports/final/FINAL_PROJECT_STATUS.md. Các chỉ số dùng đúng giá trị final champion đã chốt.",
        categories,
        {
            "Macro F1": [float(r["macro_f1"]) for r in pred],
            "Recall Low": [float(r["recall_low"]) for r in pred],
            "F1 Low": [float(r["f1_low"]) for r in pred],
        },
        {"Macro F1": "#2563EB", "Recall Low": "#16A34A", "F1 Low": "#F97316"},
    )

    baseline = load_baseline_rows()
    xapi_rows = [r for r in baseline if r["dataset"].lower() == "xapi"]
    categories = [r["model_type"].capitalize() for r in xapi_rows]
    save_grouped_bar(
        FIG_DIR / "fig_4_2_xapi_baseline_comparison.png",
        "So sánh xAPI deep final với baseline ML",
        "Nguồn: reports/final/final_baseline_comparison.csv. Baseline chỉ dùng đối chứng, không dùng làm teacher.",
        categories,
        {"Macro F1": [float(r["macro_f1"]) for r in xapi_rows]},
        {"Macro F1": "#7C3AED"},
    )

    categories = [f"{r['dataset']} {r['scenario']}" for r in pred]
    save_grouped_bar(
        FIG_DIR / "fig_4_3_low_class_focus.png",
        "Phân tích riêng lớp Low của mô hình deep final",
        "Nguồn: reports/final/FINAL_PROJECT_STATUS.md. Recall Low và F1 Low đánh giá khả năng phát hiện nhóm nguy cơ.",
        categories,
        {
            "Recall Low": [float(r["recall_low"]) for r in pred],
            "F1 Low": [float(r["f1_low"]) for r in pred],
        },
        {"Recall Low": "#16A34A", "F1 Low": "#F97316"},
    )

    macro_sorted = sorted(pred, key=lambda r: float(r["macro_f1"]), reverse=True)
    save_grouped_bar(
        FIG_DIR / "fig_4_4_macro_f1_ranking.png",
        "Xếp hạng Macro F1 của các final champion",
        "Nguồn: reports/final/FINAL_PROJECT_STATUS.md. Student-Mat late là kết quả deep mạnh nhất.",
        [f"{r['dataset']} {r['scenario']}" for r in macro_sorted],
        {"Macro F1": [float(r["macro_f1"]) for r in macro_sorted]},
        {"Macro F1": "#2563EB"},
    )

    rec = load_recommender_metrics()
    categories = [str(r["dataset"]) for r in rec]
    save_grouped_bar(
        FIG_DIR / "fig_5_1_recommender_offline_metrics.png",
        "Đánh giá offline module khuyến nghị RA-HLPR",
        "Nguồn: outputs/recommender/*/recommender_metrics.json. Đây là weak-supervision/rule-based reference.",
        categories,
        {
            "Risk Macro F1": [float(r["risk_macro_f1"]) for r in rec],
            "Risk Micro F1": [float(r["risk_micro_f1"]) for r in rec],
            "Precision@3": [float(r["precision_at_3"]) for r in rec],
            "Recall@3": [float(r["recall_at_3"]) for r in rec],
            "NDCG@3": [float(r["ndcg_at_3"]) for r in rec],
            "Coverage@3": [float(r["coverage_at_3"]) for r in rec],
            "Risk Coverage": [float(r["risk_coverage"]) for r in rec],
        },
        {
            "Risk Macro F1": "#2563EB",
            "Risk Micro F1": "#0EA5E9",
            "Precision@3": "#16A34A",
            "Recall@3": "#84CC16",
            "NDCG@3": "#F97316",
            "Coverage@3": "#E11D48",
            "Risk Coverage": "#7C3AED",
        },
    )

    save_grouped_bar(
        FIG_DIR / "fig_5_3_risk_diagnosis_metrics.png",
        "Metric chẩn đoán rủi ro của RiskDiagnosisHead",
        "Nguồn: outputs/recommender/*/recommender_metrics.json. Giá trị được lấy trực tiếp từ JSON output.",
        categories,
        {
            "Risk Macro F1": [float(r["risk_macro_f1"]) for r in rec],
            "Risk Micro F1": [float(r["risk_micro_f1"]) for r in rec],
        },
        {"Risk Macro F1": "#2563EB", "Risk Micro F1": "#0EA5E9"},
    )

    save_grouped_bar(
        FIG_DIR / "fig_5_4_ranking_metrics.png",
        "Metric xếp hạng Top-3 can thiệp",
        "Nguồn: outputs/recommender/*/recommender_metrics.json. Precision@3, Recall@3, NDCG@3 và Coverage@3 đánh giá ranking.",
        categories,
        {
            "Precision@3": [float(r["precision_at_3"]) for r in rec],
            "Recall@3": [float(r["recall_at_3"]) for r in rec],
            "NDCG@3": [float(r["ndcg_at_3"]) for r in rec],
            "Coverage@3": [float(r["coverage_at_3"]) for r in rec],
        },
        {"Precision@3": "#16A34A", "Recall@3": "#84CC16", "NDCG@3": "#F97316", "Coverage@3": "#E11D48"},
    )

    save_grouped_bar(
        FIG_DIR / "fig_5_5_path_quality_metrics.png",
        "Metric chất lượng lộ trình học tập",
        "Nguồn: outputs/recommender/*/recommender_metrics.json. Prereq Violation càng thấp càng tốt.",
        categories,
        {
            "Risk Coverage": [float(r["risk_coverage"]) for r in rec],
            "Difficulty Progression": [float(r["difficulty_progression"]) for r in rec],
            "Prereq Violation": [float(r["prereq_violation"]) for r in rec],
        },
        {"Risk Coverage": "#7C3AED", "Difficulty Progression": "#2563EB", "Prereq Violation": "#DC2626"},
    )


def configure_section(section):
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(3.0)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(2.25)
    section.footer_distance = Cm(1.5)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)
    set_run_font(run, 11)


def set_page_numbering(section, start: int | None = None, fmt: str | None = None):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    if start is not None:
        pg_num.set(qn("w:start"), str(start))
    if fmt is not None:
        pg_num.set(qn("w:fmt"), fmt)


def clear_page_numbering(section):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is not None:
        sect_pr.remove(pg_num)


def set_header_footer(section, header_text: str | None, start: int | None = None, fmt: str | None = None):
    configure_section(section)
    if start is not None or fmt is not None:
        set_page_numbering(section, start=start, fmt=fmt)
    else:
        clear_page_numbering(section)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header_p = section.header.paragraphs[0]
    header_p.text = ""
    if header_text:
        header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = header_p.add_run(header_text)
        set_run_font(run, 11, italic=True)
    footer_p = section.footer.paragraphs[0]
    footer_p.text = ""
    add_page_number(footer_p)


def add_paragraph(doc, text: str = "", style: str | None = None, align=None, bold=False, italic=False):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, bold=bold, italic=italic)
    return p


def add_heading(doc, text: str, level: int = 1):
    p = doc.add_heading(level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, 14 if level == 1 else 13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_bullet(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, 13)
    return p


def add_numbered(doc, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, 13)
    return p


def add_caption(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, 12, bold=True, italic=True)


def add_figure(doc, image_name: str, caption: str, width_cm: float = 14.5):
    path = FIG_DIR / image_name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)
    src = doc.add_paragraph()
    src.alignment = WD_ALIGN_PARAGRAPH.CENTER
    src.paragraph_format.space_after = Pt(8)
    note = src.add_run(f"Nguồn hình: sinh bằng Python từ artifact thật trong repo ({path.as_posix()}).")
    set_run_font(note, 10, italic=True)


def set_cell_text(cell, text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, 11, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc, headers: list[str], rows: list[list[str]], widths_cm: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr_cells[i], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "D9EAF7")
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if i > 0 and len(value) < 24 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], value, align=align)
    if widths_cm:
        for row in table.rows:
            for i, width in enumerate(widths_cm):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return table


def add_cover(doc, sub: bool = False):
    configure_section(doc.sections[-1])

    def blank(lines: int = 1):
        for _ in range(lines):
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(0)
    for text in [SCHOOL, FACULTY]:
        run = p.add_run(text + "\n")
        set_run_font(run, 14, bold=True)

    blank(5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("KHÓA LUẬN TỐT NGHIỆP")
    set_run_font(run, 18, bold=True)

    blank(4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(TITLE)
    set_run_font(run, 24, bold=True)

    blank(5)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(6.0)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("GIẢNG VIÊN HƯỚNG DẪN: [BỔ SUNG]\n")
    set_run_font(run, 14, bold=True)
    run = p.add_run("SINH VIÊN THỰC HIỆN: [BỔ SUNG HỌ TÊN - MSSV]")
    set_run_font(run, 14, bold=True)

    blank(8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("TP. HỒ CHÍ MINH - THÁNG 06 NĂM 2026")
    set_run_font(run, 14, bold=True)
    if sub:
        p.add_run("")
    doc.add_page_break()


def add_static_toc(doc):
    add_heading(doc, "MỤC LỤC", 1)
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-3" \h \z \u')
    run_el = OxmlElement("w:r")
    text_el = OxmlElement("w:t")
    text_el.text = "Mục lục tự động. Khi mở bằng Microsoft Word, chọn Update Table để cập nhật số trang."
    run_el.append(text_el)
    fld.append(run_el)
    p._p.append(fld)


def front_matter(doc):
    set_header_footer(doc.add_section(WD_SECTION.NEW_PAGE), None, start=1, fmt="lowerRoman")
    add_heading(doc, "LỜI CẢM ƠN", 1)
    add_paragraph(
        doc,
        "Em xin gửi lời cảm ơn đến Quý Thầy Cô Khoa Công nghệ Thông tin, Trường Đại học Ngoại ngữ - Tin học TP. Hồ Chí Minh đã tạo điều kiện học tập và định hướng chuyên môn trong quá trình thực hiện khóa luận. Em cũng xin cảm ơn giảng viên hướng dẫn đã góp ý về phạm vi đề tài, cách tổ chức thực nghiệm và yêu cầu trình bày kết quả một cách trung thực.",
    )
    add_paragraph(
        doc,
        "Khóa luận này được hoàn thiện dựa trên quá trình xây dựng, kiểm thử và tổng hợp kết quả của hệ thống dự đoán thành tích học tập sinh viên bằng mô hình CNN-BiLSTM, kết hợp module khuyến nghị lộ trình học tập RA-HLPR. Những thiếu sót còn lại trong báo cáo là trách nhiệm của sinh viên thực hiện.",
    )
    doc.add_page_break()

    add_heading(doc, "LỜI CAM ĐOAN", 1)
    add_paragraph(
        doc,
        "Em cam đoan khóa luận này là sản phẩm nghiên cứu và triển khai của bản thân dưới sự hướng dẫn của giảng viên hướng dẫn. Các kết quả thực nghiệm được ghi nhận từ pipeline trong đồ án, không sử dụng kết quả optimistic/paper-like, không dùng locked test để chọn mô hình và không làm đẹp số bằng leakage.",
    )
    add_paragraph(
        doc,
        "Các tài liệu, công trình và bộ dữ liệu tham khảo được trích dẫn trong phần tài liệu tham khảo. Mô hình dự đoán cuối cùng, module khuyến nghị và các giới hạn của đề tài được trình bày đúng với trạng thái kỹ thuật hiện tại của dự án.",
    )
    doc.add_paragraph("\n\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Sinh viên thực hiện\n\n[BỔ SUNG HỌ TÊN]")
    set_run_font(run, 13)
    doc.add_page_break()

    add_static_toc(doc)

    add_heading(doc, "DANH MỤC CÁC KÝ HIỆU, CHỮ VIẾT TẮT VÀ THUẬT NGỮ", 1)
    add_table(
        doc,
        ["Ký hiệu/Thuật ngữ", "Diễn giải"],
        [
            ["CNN", "Convolutional Neural Network - mạng nơ-ron tích chập."],
            ["BiLSTM", "Bidirectional Long Short-Term Memory - mạng LSTM hai chiều."],
            ["RA-HLPR", "Risk-Aware Hybrid Learning Path Recommender - module khuyến nghị lộ trình học tập dựa trên rủi ro."],
            ["OOF", "Out-of-fold probabilities - xác suất dự đoán ngoài fold dùng để tinh chỉnh threshold."],
            ["Locked test", "Tập kiểm thử khóa, chỉ dùng cho đánh giá cuối cùng sau khi đã chọn mô hình bằng CV/OOF."],
            ["Macro F1", "Trung bình F1 của các lớp, không trọng số theo số mẫu từng lớp."],
            ["Recall Low", "Recall của lớp Low, phản ánh khả năng phát hiện sinh viên có nguy cơ thấp kết quả học tập."],
        ],
        [4.0, 11.0],
    )
    doc.add_page_break()

    add_heading(doc, "DANH MỤC CÁC BẢNG", 1)
    for item in [
        "Bảng 1.1. Phạm vi bộ dữ liệu và kịch bản sử dụng",
        "Bảng 3.1. Quy tắc xử lý dữ liệu và chống leakage",
        "Bảng 3.2. Các thành phần chính của kiến trúc CNN-BiLSTM",
        "Bảng 4.1. Kết quả mô hình dự đoán cuối cùng",
        "Bảng 4.2. So sánh mô hình deep với baseline có sẵn",
        "Bảng 4.3. Nguồn artifact dùng để sinh hình metric",
        "Bảng 4.4. Diễn giải vai trò của các metric chính",
        "Bảng 5.1. Định nghĩa rủi ro vận hành",
        "Bảng 5.2. Kết quả đánh giá offline module khuyến nghị",
        "Bảng 5.3. Nhóm can thiệp và phạm vi áp dụng",
        "Bảng 5.4. Case study xAPI",
        "Bảng 5.5. Case study student-por",
    ]:
        add_paragraph(doc, item)
    doc.add_page_break()

    add_heading(doc, "DANH MỤC CÁC HÌNH VẼ, SƠ ĐỒ", 1)
    for item in [
        "Hình 3.1. Quy trình dữ liệu - xử lý - mô hình - khuyến nghị",
        "Hình 3.2. Kiến trúc mô hình CNN-BiLSTM final",
        "Hình 4.1. Kết quả locked test của mô hình deep final",
        "Hình 4.2. So sánh xAPI deep final với baseline ML",
        "Hình 4.3. Phân tích riêng lớp Low của mô hình deep final",
        "Hình 4.4. Xếp hạng Macro F1 của các final champion",
        "Hình 5.1. Pipeline RA-HLPR tạo lộ trình học tập 4 tuần",
        "Hình 5.2. Đánh giá offline module khuyến nghị RA-HLPR",
        "Hình 5.3. Metric chẩn đoán rủi ro của RiskDiagnosisHead",
        "Hình 5.4. Metric xếp hạng Top-3 can thiệp",
        "Hình 5.5. Metric chất lượng lộ trình học tập",
    ]:
        add_paragraph(doc, item)
    doc.add_page_break()


def new_chapter(doc, title: str):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    if title.startswith("CHƯƠNG 1."):
        set_header_footer(section, title, start=1, fmt="decimal")
    else:
        set_header_footer(section, title)
    add_heading(doc, title, 1)


def add_pipeline_figure(doc, caption: str, lines: list[str]):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in lines:
        run = p.add_run(line + "\n")
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
        run.font.size = Pt(10)
    add_caption(doc, caption)


def chapter_1(doc):
    new_chapter(doc, "CHƯƠNG 1. MỞ ĐẦU")
    add_heading(doc, "1.1. Lý do chọn đề tài", 2)
    add_paragraph(
        doc,
        "Dự đoán thành tích học tập là một hướng ứng dụng quan trọng của khai phá dữ liệu giáo dục. Trong môi trường học tập hiện đại, dữ liệu điểm số giữa kỳ, dữ liệu hành vi học trực tuyến, mức độ tham gia và một số yếu tố hoàn cảnh có thể hỗ trợ nhà trường phát hiện sớm sinh viên có nguy cơ đạt kết quả thấp. Nếu chỉ dừng ở dự đoán nhãn Low, Medium hoặc High, hệ thống chưa tạo ra được hỗ trợ vận hành cụ thể cho người học. Vì vậy, đồ án kết hợp mô hình dự đoán với một module khuyến nghị lộ trình học tập theo rủi ro.",
    )
    add_paragraph(
        doc,
        "Đề tài lựa chọn mô hình học sâu CNN-BiLSTM làm mô hình chính vì kiến trúc này có khả năng khai thác tín hiệu tuần tự ngắn từ điểm G1/G2 của Student datasets hoặc các chỉ số hành vi học tập của xAPI, đồng thời vẫn có thể kết hợp với đặc trưng ngữ cảnh thông qua nhánh context/gated fusion. Baseline machine learning chỉ được sử dụng để đối chứng kết quả, không được dùng làm teacher, distillation source hoặc nguồn pseudo-label cho mô hình deep.",
    )
    add_heading(doc, "1.2. Mục tiêu nghiên cứu", 2)
    for text in [
        "Xây dựng pipeline dự đoán thành tích học tập sinh viên theo ba mức Low, Medium và High.",
        "Thiết kế mô hình CNN-BiLSTM và các biến thể phù hợp với đặc điểm dữ liệu Student và xAPI.",
        "Đảm bảo quy trình đánh giá công bằng: locked test chỉ dùng cho đánh giá cuối cùng, threshold được tinh chỉnh bằng CV/OOF.",
        "Xây dựng module RA-HLPR để chuyển xác suất dự đoán thành khuyến nghị can thiệp và lộ trình học tập 4 tuần.",
        "Trình bày kết quả trung thực, không sử dụng student-combine, không dùng ADASYN trực tiếp trên dữ liệu phân loại label encoding và không claim regression head khi RMSE chưa đạt yêu cầu.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "1.3. Đối tượng và phạm vi nghiên cứu", 2)
    add_paragraph(
        doc,
        "Đối tượng nghiên cứu là bài toán dự đoán thành tích học tập và tạo khuyến nghị học tập dựa trên dữ liệu giáo dục dạng bảng. Phạm vi triển khai gồm ba bộ dữ liệu: student-mat, student-por và xAPI. Đồ án không sử dụng student-combine để tránh trộn dữ liệu theo cách có thể làm sai lệch giao thức đánh giá.",
    )
    add_caption(doc, "Bảng 1.1. Phạm vi bộ dữ liệu và kịch bản sử dụng")
    add_table(
        doc,
        ["Bộ dữ liệu", "Kịch bản", "Vai trò trong đồ án"],
        [
            ["student-mat", "late", "Dự đoán dựa trên G1 và G2; mô hình deep chính là sequence_cnn_bilstm_only."],
            ["student-por", "late, midterm", "Đánh giá mô hình trên môn Portuguese; late dùng G1/G2, midterm dùng G1."],
            ["xAPI", "default", "Dữ liệu hành vi học tập trực tuyến; dùng gated_fusion_v28 do có cả tín hiệu sequence và context."],
        ],
        [3.0, 3.0, 9.0],
    )

    add_heading(doc, "1.4. Đóng góp của đề tài", 2)
    for text in [
        "Hoàn thiện pipeline dự đoán có kiểm soát rò rỉ dữ liệu, thống nhất train pool, locked test, seed, preprocessing và metrics.",
        "Chốt mô hình deep final theo kết quả thực nghiệm: sequence_cnn_bilstm_only cho Student và gated_fusion_v28 cho xAPI.",
        "Bổ sung module khuyến nghị RA-HLPR có khả năng giải thích: risk diagnosis, candidate filtering, hybrid scoring và path planning.",
        "Xây dựng báo cáo kỹ thuật và artifact final phục vụ khóa luận.",
    ]:
        add_bullet(doc, text)


def chapter_2(doc):
    new_chapter(doc, "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT VÀ TỔNG QUAN")
    add_heading(doc, "2.1. Khai phá dữ liệu giáo dục", 2)
    add_paragraph(
        doc,
        "Khai phá dữ liệu giáo dục tập trung phân tích dữ liệu học tập nhằm hỗ trợ dự đoán, chẩn đoán, cá nhân hóa và ra quyết định. Trong bài toán của đồ án, đầu ra chính là mức thành tích học tập Low, Medium hoặc High. Nhãn này giúp nhà trường hoặc người học nhận biết mức độ rủi ro, nhưng để có giá trị thực tiễn hơn cần liên kết với các can thiệp học tập cụ thể.",
    )
    add_heading(doc, "2.2. Mô hình CNN và BiLSTM", 2)
    add_paragraph(
        doc,
        "CNN được dùng để trích xuất mẫu cục bộ trên chuỗi đặc trưng. Với dữ liệu Student, chuỗi có thể rất ngắn, ví dụ gồm G1 và G2 trong kịch bản late. Với xAPI, chuỗi hành vi có thể gồm các tín hiệu như raisedhands, VisITedResources, AnnouncementsView và Discussion. BiLSTM xử lý chuỗi theo hai chiều, giúp mô hình tổng hợp thông tin trước và sau trong chuỗi trước khi đưa vào lớp phân loại.",
    )
    add_heading(doc, "2.3. Mất cân bằng lớp và threshold tuning", 2)
    add_paragraph(
        doc,
        "Bài toán dự đoán kết quả học tập thường có mất cân bằng lớp, đặc biệt với lớp Low. Đề tài ưu tiên Macro F1, Recall Low và F1 Low thay vì chỉ tối ưu Accuracy. Các chiến lược được xem xét gồm class weight, SMOTENC, random oversampling, focal loss và class-balanced focal loss. ADASYN trực tiếp trên dữ liệu Student có biến phân loại đã label encoding không được dùng vì có thể nội suy ra giá trị phân loại không hợp lệ.",
    )
    add_paragraph(
        doc,
        "Threshold tuning được thực hiện bằng xác suất CV/OOF trên train pool. Locked test chỉ được sử dụng một lần cho đánh giá cuối cùng, không dùng để chọn mô hình hoặc tinh chỉnh threshold.",
    )
    add_heading(doc, "2.4. Khuyến nghị lộ trình học tập", 2)
    add_paragraph(
        doc,
        "Do các bộ dữ liệu không có lịch sử tương tác user-item hoặc phản hồi sau khuyến nghị, đồ án không xây dựng collaborative filtering. Thay vào đó, module RA-HLPR được thiết kế như một hệ thống downstream: xác suất dự đoán từ CNN-BiLSTM được dùng để chẩn đoán rủi ro, lọc can thiệp và lập lộ trình học tập 4 tuần. Cách tiếp cận này phù hợp với dữ liệu hiện có và giữ được tính giải thích.",
    )


def chapter_3(doc):
    new_chapter(doc, "CHƯƠNG 3. PHƯƠNG PHÁP NGHIÊN CỨU VÀ MÔ HÌNH ĐỀ XUẤT")
    add_heading(doc, "3.1. Quy trình tổng thể", 2)
    add_figure(
        doc,
        "fig_3_1_pipeline_overview.png",
        "Hình 3.1. Quy trình dữ liệu - xử lý - mô hình - khuyến nghị",
        15.0,
    )
    add_heading(doc, "3.2. Tiền xử lý dữ liệu", 2)
    add_paragraph(
        doc,
        "Pipeline tiền xử lý giữ nguyên nguyên tắc cùng train pool, cùng locked test, cùng feature setting, seed, folds và metrics giữa baseline và mô hình deep. Các biến phân loại được xử lý trong pipeline thay vì áp dụng oversampling tùy tiện. Với Student datasets, các kịch bản được phân tách theo thông tin sẵn có: early không dùng G1/G2, midterm dùng G1 và late dùng cả G1/G2. Mô hình chính cuối cùng sử dụng các kịch bản đã được xác nhận qua thực nghiệm: student-mat late, student-por late và student-por midterm.",
    )
    add_caption(doc, "Bảng 3.1. Quy tắc xử lý dữ liệu và chống leakage")
    add_table(
        doc,
        ["Thành phần", "Cách xử lý trong đồ án", "Lý do"],
        [
            ["Train pool / locked test", "Tách trước khi chọn mô hình; locked test chỉ dùng final evaluation.", "Tránh chọn model bằng test."],
            ["Threshold tuning", "Dùng CV/OOF probabilities trên train pool.", "Không tune bằng locked test."],
            ["Student categorical features", "Không dùng ADASYN trực tiếp trên label encoding.", "Tránh nội suy giá trị phân loại không hợp lệ."],
            ["Baseline ML", "Chỉ dùng đối chứng cuối cùng.", "Không dùng teacher, distillation, pseudo-label hoặc feature importance."],
            ["student-combine", "Không sử dụng.", "Giữ dataset/scenario rõ ràng."],
        ],
        [3.2, 6.0, 5.8],
    )
    add_heading(doc, "3.3. Kiến trúc CNN-BiLSTM", 2)
    add_figure(
        doc,
        "fig_3_2_cnn_bilstm_architecture.png",
        "Hình 3.2. Kiến trúc mô hình CNN-BiLSTM final",
        15.0,
    )
    add_caption(doc, "Bảng 3.2. Các thành phần chính của kiến trúc CNN-BiLSTM")
    add_table(
        doc,
        ["Thành phần", "Vai trò"],
        [
            ["Conv1D", "Trích xuất mẫu cục bộ trên chuỗi điểm hoặc hành vi học tập."],
            ["BiLSTM", "Tổng hợp quan hệ tuần tự theo hai chiều."],
            ["Attention/Pooling", "Tạo vector đại diện chuỗi trước khi phân loại."],
            ["Context/Gated fusion", "Kết hợp đặc trưng chuỗi với đặc trưng ngữ cảnh, dùng cho xAPI."],
            ["Threshold tuning", "Điều chỉnh quyết định lớp Low bằng CV/OOF probabilities."],
        ],
        [4.0, 11.0],
    )
    add_heading(doc, "3.4. Lựa chọn mô hình cuối cùng", 2)
    add_paragraph(
        doc,
        "Sau các vòng thực nghiệm, đồ án không tiếp tục mở rộng kiến trúc mà chốt mô hình theo kết quả tốt nhất đã được xác nhận. Với Student datasets, sequence_cnn_bilstm_only được chọn vì nhánh chuỗi G1/G2 đã đủ mạnh và fusion không thắng theo CV. Với xAPI, gated_fusion_v28 được giữ làm final champion vì xAPI có cả hành vi học tập dạng sequence và đặc trưng ngữ cảnh/categorical.",
    )


def chapter_4(doc):
    new_chapter(doc, "CHƯƠNG 4. THỰC NGHIỆM VÀ ĐÁNH GIÁ MÔ HÌNH DỰ ĐOÁN")
    add_heading(doc, "4.1. Thiết lập thực nghiệm", 2)
    add_paragraph(
        doc,
        "Các mô hình được đánh giá bằng Accuracy, Macro Precision, Macro Recall, Macro F1, Recall Low và F1 Low. Trong báo cáo final, các bảng chính tập trung vào Macro F1, Recall Low và F1 Low vì đây là các chỉ số quyết định việc phát hiện nhóm sinh viên có nguy cơ thấp kết quả học tập. Locked test không dùng để chọn mô hình, chỉ dùng cho final evaluation sau khi mô hình và threshold đã được quyết định bằng train pool/CV/OOF.",
    )
    add_heading(doc, "4.2. Kết quả mô hình dự đoán cuối cùng", 2)
    add_caption(doc, "Bảng 4.1. Kết quả mô hình dự đoán cuối cùng")
    add_table(
        doc,
        ["Dataset", "Scenario", "Mô hình", "Prediction mode", "Macro F1", "Recall Low", "F1 Low"],
        [
            ["student-mat", "late", "sequence_cnn_bilstm_only", "low_f1_tuned", "0.9365", "0.9615", "0.8929"],
            ["student-por", "late", "sequence_cnn_bilstm_only", "low_f1_tuned", "0.8783", "0.9000", "0.8182"],
            ["student-por", "midterm", "sequence_cnn_bilstm_only", "argmax", "0.8228", "0.6500", "0.7429"],
            ["xAPI", "default", "gated_fusion_v28", "low_f1_tuned", "0.7541", "0.8846", "0.8214"],
        ],
        [2.0, 2.0, 4.0, 2.8, 1.4, 1.4, 1.4],
    )
    add_paragraph(
        doc,
        "Kết quả cho thấy mô hình deep đạt hiệu quả tốt trên Student datasets, đặc biệt ở student-mat late với Macro F1 0.9365 và Recall Low 0.9615. Trên xAPI, mô hình deep giữ Recall Low cao 0.8846 nhưng Macro F1 vẫn thấp hơn baseline Random Forest đã có trong kết quả đối chứng.",
    )
    add_figure(
        doc,
        "fig_4_1_prediction_metrics.png",
        "Hình 4.1. Kết quả locked test của mô hình deep final",
        15.0,
    )
    add_heading(doc, "4.3. So sánh với baseline", 2)
    add_caption(doc, "Bảng 4.2. So sánh mô hình deep với baseline có sẵn")
    add_table(
        doc,
        ["Dataset", "Loại mô hình", "Mô hình", "Macro F1", "Recall Low", "F1 Low", "Ghi chú"],
        [
            ["xAPI", "Deep", "gated_fusion_v28", "0.7541", "0.8846", "0.8214", "Mô hình final cho xAPI."],
            ["xAPI", "Baseline", "RandomForestClassifier", "0.8465", "not_available", "not_available", "Chỉ dùng đối chứng, không dùng huấn luyện deep."],
        ],
        [1.7, 2.0, 3.4, 1.4, 1.4, 1.4, 3.7],
    )
    add_paragraph(
        doc,
        "Baseline machine learning không được dùng làm teacher, distillation source, pseudo-label source, baseline probability source hoặc feature-importance source. Việc so sánh chỉ nhằm đặt kết quả deep learning vào bối cảnh đối chứng công bằng.",
    )
    add_figure(
        doc,
        "fig_4_2_xapi_baseline_comparison.png",
        "Hình 4.2. So sánh xAPI deep final với baseline ML",
        14.0,
    )
    add_caption(doc, "Bảng 4.3. Nguồn artifact dùng để sinh hình metric")
    add_table(
        doc,
        ["Hình", "Nguồn dữ liệu", "Cách sinh"],
        [
            ["Hình 4.1", "reports/final/FINAL_PROJECT_STATUS.md", "Script Python parse bảng final champion và vẽ bar chart bằng PIL."],
            ["Hình 4.2", "reports/final/final_baseline_comparison.csv", "Script Python đọc CSV baseline comparison và vẽ Macro F1 deep vs baseline."],
            ["Hình 5.1", "outputs/recommender/*/recommender_metrics.json", "Script Python đọc JSON metric thực tế của xAPI và student-por."],
        ],
        [2.5, 6.0, 6.5],
    )
    add_heading(doc, "4.4. Nhận xét kỹ thuật", 2)
    for text in [
        "Student-Mat late và Student-Por late sử dụng mô hình sequence-only vì kết quả cũ tốt hơn các biến thể fusion sau đó.",
        "Student-Por midterm được giữ với argmax vì đây là cấu hình deep tốt nhất đã chốt cho kịch bản này.",
        "xAPI giữ gated_fusion_v28 do các vòng V31/V32 không vượt được champion hiện tại trên locked test.",
        "Regression head không được claim là kết quả chính vì RMSE vẫn chưa đạt mức đủ tin cậy.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "4.5. Phân tích hình metric trọng tâm", 2)
    add_paragraph(
        doc,
        "Hình 4.1 trình bày đồng thời ba chỉ số được dùng để chốt mô hình deep final: Macro F1, Recall Low và F1 Low. Macro F1 phản ánh chất lượng phân loại cân bằng giữa các lớp, trong khi Recall Low và F1 Low tập trung vào nhóm sinh viên có nguy cơ đạt kết quả thấp. Việc trình bày ba chỉ số cùng lúc giúp tránh kết luận dựa trên Accuracy hoặc một metric đơn lẻ.",
    )
    add_paragraph(
        doc,
        "Student-Mat late đạt Macro F1 0.9365, Recall Low 0.9615 và F1 Low 0.8929. Đây là cấu hình mạnh nhất trong các final champion vì vừa giữ được Macro F1 cao vừa phát hiện tốt lớp Low. Student-Por late đạt Macro F1 0.8783 và Recall Low 0.9000, cho thấy mô hình sequence-only vẫn giữ khả năng phát hiện nhóm nguy cơ trên một môn học khác. Student-Por midterm có Macro F1 0.8228 nhưng Recall Low 0.6500, phản ánh việc chỉ dùng G1 làm tín hiệu giữa kỳ khiến dự đoán lớp Low khó hơn late scenario.",
    )
    add_figure(
        doc,
        "fig_4_3_low_class_focus.png",
        "Hình 4.3. Phân tích riêng lớp Low của mô hình deep final",
        15.0,
    )
    add_paragraph(
        doc,
        "Hình 4.3 tách riêng Recall Low và F1 Low để làm rõ mục tiêu phát hiện sinh viên có nguy cơ. Với student-mat late, Recall Low cao hơn F1 Low, nghĩa là mô hình ưu tiên phát hiện đúng phần lớn sinh viên Low, nhưng vẫn còn đánh đổi nhất định ở precision của lớp này. Với xAPI, Recall Low 0.8846 và F1 Low 0.8214 cho thấy mô hình deep tuy chưa thắng baseline về Macro F1 nhưng vẫn có ý nghĩa khi bài toán ưu tiên cảnh báo nhóm rủi ro.",
    )
    add_paragraph(
        doc,
        "Trong báo cáo, Recall Low không được dùng một mình để claim mô hình tốt nhất. Nếu chỉ tối ưu Recall Low, mô hình có thể dự đoán quá nhiều mẫu thành Low và làm giảm precision. Vì vậy F1 Low được đặt cạnh Recall Low để kiểm tra cân bằng giữa phát hiện và độ chính xác của cảnh báo. Đây cũng là lý do cấu hình threshold được tinh chỉnh bằng CV/OOF thay vì locked test.",
    )
    add_figure(
        doc,
        "fig_4_4_macro_f1_ranking.png",
        "Hình 4.4. Xếp hạng Macro F1 của các final champion",
        15.0,
    )
    add_paragraph(
        doc,
        "Hình 4.4 cho thấy thứ tự Macro F1 của các mô hình final: student-mat late cao nhất, kế đến là student-por late, student-por midterm và xAPI. Thứ tự này phù hợp với mức độ mạnh của tín hiệu đầu vào: late scenario có G1/G2 nên giàu thông tin hơn midterm, còn xAPI có đặc trưng hành vi và ngữ cảnh phức tạp hơn nên mô hình deep khó vượt baseline ML.",
    )
    add_caption(doc, "Bảng 4.4. Diễn giải vai trò của các metric chính")
    add_table(
        doc,
        ["Metric", "Ý nghĩa", "Cách dùng trong báo cáo"],
        [
            ["Macro F1", "Trung bình F1 của các lớp, giảm thiên lệch do mất cân bằng lớp.", "Metric ưu tiên khi chọn mô hình bằng CV/OOF."],
            ["Recall Low", "Tỷ lệ sinh viên Low được phát hiện đúng.", "Dùng để đánh giá khả năng cảnh báo sớm nhóm nguy cơ."],
            ["F1 Low", "Cân bằng precision và recall của lớp Low.", "Dùng cùng Recall Low để tránh cảnh báo quá rộng."],
            ["Baseline Macro F1", "Điểm đối chứng của mô hình ML cùng scenario.", "Chỉ dùng so sánh cuối cùng, không dùng huấn luyện deep."],
        ],
        [3.0, 6.0, 6.0],
    )
    add_heading(doc, "4.6. Diễn giải kết quả theo từng bộ dữ liệu", 2)
    add_paragraph(
        doc,
        "Đối với student-mat late, mô hình sequence_cnn_bilstm_only + low_f1_tuned được chọn vì đây là cấu hình giữ được cả Macro F1 và Recall Low cao. Late scenario sử dụng cả G1 và G2, nên chuỗi điểm số ngắn nhưng có ý nghĩa rõ ràng về tiến trình học tập. Kết quả này cũng cho thấy không cần ép fusion nếu nhánh sequence-only đã thắng trong kiểm chứng.",
    )
    add_paragraph(
        doc,
        "Đối với student-por late, cùng kiến trúc sequence-only vẫn đạt Macro F1 0.8783 và Recall Low 0.9000. Điều này cho thấy mô hình không chỉ phù hợp một môn học duy nhất, dù mức F1 Low 0.8182 thấp hơn student-mat. Sự khác biệt này có thể đến từ phân phối điểm, số lượng mẫu và mối quan hệ giữa G1/G2 với nhãn cuối.",
    )
    add_paragraph(
        doc,
        "Đối với student-por midterm, mô hình dùng argmax thay vì threshold low_f1_tuned. Kịch bản midterm chỉ có G1, do đó tín hiệu tuần tự yếu hơn late. Macro F1 0.8228 vẫn cho thấy mô hình có khả năng dự đoán ở giai đoạn giữa kỳ, nhưng Recall Low 0.6500 cho thấy cảnh báo sớm nhóm Low còn khó và không nên được trình bày như kết quả mạnh nhất.",
    )
    add_paragraph(
        doc,
        "Đối với xAPI, gated_fusion_v28 + low_f1_tuned được giữ làm final vì đây là champion deep đã chốt. Mô hình sử dụng gated fusion để kết hợp hành vi học tập và đặc trưng ngữ cảnh. Tuy nhiên, baseline RandomForestClassifier có Macro F1 0.8465, cao hơn deep 0.7541. Báo cáo vì vậy trình bày xAPI một cách thận trọng: deep model giữ Recall Low cao, nhưng chưa claim vượt baseline.",
    )


def chapter_5(doc):
    new_chapter(doc, "CHƯƠNG 5. MODULE KHUYẾN NGHỊ LỘ TRÌNH HỌC TẬP")
    add_heading(doc, "5.1. Mục tiêu module khuyến nghị", 2)
    add_paragraph(
        doc,
        "RA-HLPR là module downstream của mô hình CNN-BiLSTM. Module này nhận xác suất Low, Medium và High, sau đó kết hợp với các tín hiệu quan sát được để tạo chẩn đoán rủi ro, xếp hạng can thiệp và lập lộ trình học tập 4 tuần. Hệ thống không phải collaborative filtering vì không có dữ liệu user-item interaction hoặc feedback thực tế sau khuyến nghị.",
    )
    add_figure(
        doc,
        "fig_5_2_ra_hlpr_flow.png",
        "Hình 5.1. Pipeline RA-HLPR tạo lộ trình học tập 4 tuần",
        15.0,
    )
    add_heading(doc, "5.2. Định nghĩa rủi ro", 2)
    add_caption(doc, "Bảng 5.1. Định nghĩa rủi ro vận hành")
    add_table(
        doc,
        ["Dataset", "Risk", "Tín hiệu vận hành"],
        [
            ["Student", "R1_LOW_PRIOR_PERFORMANCE", "failures, G1"],
            ["Student", "R2_DECLINING_TREND", "G2 thấp hơn G1"],
            ["Student", "R3_ATTENDANCE_RISK", "absences"],
            ["Student", "R4_LOW_ENGAGEMENT", "goout, freetime, activities"],
            ["Student", "R5_INSUFFICIENT_STUDY_TIME", "studytime"],
            ["Student", "R6_HIGH_FAILURE_PROBABILITY", "failures, G1/G2 và trend; không dùng G3"],
            ["xAPI", "R3_ATTENDANCE_RISK", "StudentAbsenceDays"],
            ["xAPI", "R4_LOW_ENGAGEMENT", "VisITedResources, raisedhands, Discussion, AnnouncementsView"],
            ["xAPI", "R6_HIGH_FAILURE_PROBABILITY", "attendance, engagement, parent/school support; không dùng true Class"],
        ],
        [2.2, 4.5, 8.3],
    )
    add_heading(doc, "5.3. Công thức chấm điểm", 2)
    add_paragraph(
        doc,
        "Với mỗi can thiệp ứng viên, HybridScorer tính điểm theo công thức:",
    )
    add_paragraph(
        doc,
        "score = w1 * risk_match + w2 * performance_need + w3 * difficulty_fit + w4 * time_fit + w5 * prerequisite_fit + w6 * expected_effect + rule_adjustment",
        italic=True,
    )
    add_paragraph(
        doc,
        "Các trọng số thay đổi theo predicted class, p_low, p_high và max risk. Với sinh viên Low/high-risk, hệ thống ưu tiên risk_match và performance_need. Với Medium, hệ thống cân bằng giữa hỗ trợ và duy trì tiến độ. Với High/stable, hệ thống ưu tiên enrichment, difficulty fit và prerequisite fit.",
    )
    add_heading(doc, "5.4. Kết quả đánh giá offline", 2)
    add_caption(doc, "Bảng 5.2. Kết quả đánh giá offline module khuyến nghị")
    add_table(
        doc,
        ["Dataset", "Risk Macro F1", "Risk Micro F1", "Precision@3", "Recall@3", "NDCG@3", "Coverage@3", "Risk Coverage"],
        [
            ["xAPI", "0.9831", "0.9813", "0.6840", "0.4720", "0.8229", "0.6500", "0.8958"],
            ["student-por", "0.9359", "0.9094", "0.6641", "0.3185", "0.7455", "0.5500", "0.9508"],
        ],
        [1.9, 1.7, 1.7, 1.5, 1.4, 1.4, 1.4, 1.5],
    )
    add_paragraph(
        doc,
        "Các chỉ số khuyến nghị là đánh giá offline theo weak-supervision/rule-based reference. Đồ án không claim causal improvement vì chưa có dữ liệu phản hồi thực tế của sinh viên sau khi nhận khuyến nghị.",
    )
    add_figure(
        doc,
        "fig_5_1_recommender_offline_metrics.png",
        "Hình 5.2. Đánh giá offline module khuyến nghị RA-HLPR",
        15.0,
    )
    add_paragraph(
        doc,
        "Hình 5.2 gom các nhóm metric chính của module khuyến nghị. Tuy nhiên, để tránh làm rối kết luận, báo cáo tách tiếp ba nhóm metric: chẩn đoán rủi ro, xếp hạng can thiệp và chất lượng lộ trình. Việc tách nhóm giúp người đọc thấy RA-HLPR không chỉ dự đoán risk label mà còn cần xếp hạng can thiệp và lập kế hoạch theo tuần.",
    )
    add_figure(
        doc,
        "fig_5_3_risk_diagnosis_metrics.png",
        "Hình 5.3. Metric chẩn đoán rủi ro của RiskDiagnosisHead",
        15.0,
    )
    add_paragraph(
        doc,
        "Hình 5.3 cho thấy RiskDiagnosisHead đạt Risk Macro F1 0.9831 trên xAPI và 0.9359 trên student-por. Đây là kết quả offline dựa trên weak-supervision/rule-based reference, không phải ground truth phản hồi can thiệp thực tế. Vì vậy, báo cáo chỉ claim module có khả năng tái hiện logic risk rule và hỗ trợ giải thích, không claim cải thiện nhân quả kết quả học tập.",
    )
    add_figure(
        doc,
        "fig_5_4_ranking_metrics.png",
        "Hình 5.4. Metric xếp hạng Top-3 can thiệp",
        15.0,
    )
    add_paragraph(
        doc,
        "Hình 5.4 tập trung vào chất lượng ranking. Precision@3 của xAPI là 0.6840 và student-por là 0.6641, cho thấy phần lớn top-3 intervention khớp với reference rule. Recall@3 thấp hơn, đặc biệt ở student-por, vì một sinh viên có thể có nhiều can thiệp phù hợp hơn ba mục đầu. NDCG@3 cao hơn precision/recall vì thứ tự top recommendation thường hợp lý ngay cả khi chưa phủ hết mọi can thiệp tham chiếu.",
    )
    add_figure(
        doc,
        "fig_5_5_path_quality_metrics.png",
        "Hình 5.5. Metric chất lượng lộ trình học tập",
        15.0,
    )
    add_paragraph(
        doc,
        "Hình 5.5 thể hiện chất lượng lộ trình. Risk Coverage của student-por đạt 0.9508 và xAPI đạt 0.8958, nghĩa là lộ trình thường bao phủ được các rủi ro chính đã chẩn đoán. Prereq Violation của xAPI bằng 0.0000 và student-por bằng 0.0449, cho thấy thứ tự can thiệp nhìn chung không phá vỡ ràng buộc tiên quyết. Difficulty Progression chưa đạt tuyệt đối, do module ưu tiên cân bằng giữa khối lượng, rủi ro và tính khả thi của kế hoạch 4 tuần.",
    )
    add_caption(doc, "Bảng 5.3. Nhóm can thiệp và phạm vi áp dụng")
    add_table(
        doc,
        ["Nhóm can thiệp", "Ví dụ", "Phạm vi/điều kiện ưu tiên"],
        [
            ["Academic remediation", "Targeted Practice Exercises, Peer-Led Study Tutoring, Academic Coaching", "Ưu tiên Student khi R1/R2 hoặc predicted Low."],
            ["LMS engagement", "Daily LMS Resource Checklist, Guided Discussion Prompts, Interactive Quizzing", "Ưu tiên xAPI khi R4 low engagement active."],
            ["Attendance recovery", "Daily Attendance Monitoring, Absence Recovery Pack", "Chỉ tăng điểm khi R3 attendance risk active."],
            ["General maintenance", "Weekly Progress Review, Standard Practice Plan, Maintain LMS Engagement", "Ưu tiên Medium/High khi risk gần 0."],
            ["Parent/school support", "Parent-School Engagement Sync, Family Progress Contract", "Chỉ score cao khi R6/support risk active."],
        ],
        [3.4, 6.1, 5.5],
    )
    add_heading(doc, "5.5. Kiểm tra logic sau khi sửa", 2)
    for text in [
        "Student-Por high-risk với R1/R2 hiện ưu tiên academic remediation như Peer-Led Study Tutoring, Targeted Practice Exercises và Biweekly Academic Coaching.",
        "xAPI no-risk Medium ưu tiên Standard Practice Plan, Weekly Progress Review và Maintain LMS Engagement thay vì attendance/counselor khi không có rủi ro tương ứng.",
        "xAPI high-risk low engagement ưu tiên LMS/resource/discussion interventions đúng với rủi ro R4.",
        "Student-Mat recommender đang pending do thiếu metadata checkpoint: models/saved/final/student-mat_3class_ensemble_features.json.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "5.6. Case study minh họa khuyến nghị", 2)
    add_paragraph(
        doc,
        "Các case study dưới đây được lấy từ báo cáo recommender final. Mục tiêu của phần này là minh họa cách pipeline chuyển xác suất dự đoán thành khuyến nghị cụ thể, không phải tạo thêm metric mới. Những ví dụ này cũng cho thấy candidate filtering đã được chỉnh theo dataset_kind, tránh để intervention của xAPI lấn top recommendation trong trường hợp Student có rủi ro học lực R1/R2.",
    )
    add_caption(doc, "Bảng 5.4. Case study xAPI")
    add_table(
        doc,
        ["Trường hợp", "Prediction", "Rủi ro chính", "Top 3 can thiệp", "Diễn giải"],
        [
            ["High risk", "Low; p_low=0.59", "R4_LOW_ENGAGEMENT=1.00; R6=1.00", "Daily LMS Resource Checklist; Guided Discussion Prompts; LMS Interactive Quizzing", "Can thiệp tập trung LMS/resource/discussion đúng với low engagement."],
            ["Moderate", "Medium; p_medium=0.69", "Không có risk active", "Standard Practice Plan; Weekly Progress Review; Maintain LMS Engagement", "Không đẩy attendance/counselor khi không có attendance/support risk."],
            ["Stable", "High; p_high=0.90", "Không có risk active", "Advanced Subject Seminar; Maintain LMS Engagement; Standard Practice Plan", "Ưu tiên enrichment và duy trì engagement thay vì remedial nặng."],
        ],
        [2.1, 2.5, 3.2, 4.8, 4.4],
    )
    add_paragraph(
        doc,
        "Case xAPI high-risk cho thấy mô hình khuyến nghị không chỉ nhìn vào predicted Low mà còn xét risk score. Khi R4 low engagement active, các intervention liên quan LMS và discussion được đưa lên đầu. Ngược lại, case Medium không có risk active được chuyển sang nhóm general/light reinforcement, phù hợp yêu cầu không recommend attendance/counselor nếu không có risk tương ứng.",
    )
    add_caption(doc, "Bảng 5.5. Case study student-por")
    add_table(
        doc,
        ["Trường hợp", "Prediction", "Rủi ro chính", "Top 3 can thiệp", "Diễn giải"],
        [
            ["High risk", "Low; p_low=0.73", "R1=1.00; R2=1.00; R6=1.00", "Peer-Led Study Tutoring; Targeted Practice Exercises; Biweekly Academic Coaching", "R1/R2 được ưu tiên academic remediation, không để LMS/support lấn top."],
            ["Moderate", "Medium; p_medium=0.81", "R4_LOW_ENGAGEMENT=0.65", "Facilitated Study Group; Daily Attendance Monitoring; Academic Counselor Consultation", "Can thiệp ở mức guided support do risk vừa phải."],
            ["Stable", "High; p_high=0.73", "R4=0.49; các risk học lực thấp", "Advanced Subject Seminar; Standard Practice Plan; Weekly Progress Review", "Ưu tiên enrichment và duy trì kế hoạch học tập."],
        ],
        [2.1, 2.5, 3.2, 4.8, 4.4],
    )
    add_paragraph(
        doc,
        "Case student-por high-risk là ví dụ quan trọng sau khi sửa dataset-aware filtering. Với rủi ro chính R1_LOW_PRIOR_PERFORMANCE và R2_DECLINING_TREND, top recommendation phải là academic remediation như tutoring, practice exercises hoặc coaching. Điều này giúp khuyến nghị bám vào nguyên nhân học lực thay vì đưa các intervention thiên về LMS hoặc parent/school support lên quá cao khi không phù hợp.",
    )


def chapter_6(doc):
    new_chapter(doc, "CHƯƠNG 6. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN")
    add_heading(doc, "6.1. Kết luận", 2)
    add_paragraph(
        doc,
        "Khóa luận đã xây dựng và chốt pipeline dự đoán thành tích học tập sinh viên bằng mô hình CNN-BiLSTM, đồng thời phát triển module RA-HLPR để tạo lộ trình học tập 4 tuần dựa trên xác suất dự đoán và rủi ro vận hành. Mô hình final cho Student datasets là sequence_cnn_bilstm_only trong các kịch bản late/midterm đã chốt; mô hình final cho xAPI là gated_fusion_v28 với low_f1_tuned.",
    )
    add_paragraph(
        doc,
        "Các kết quả được trình bày theo nguyên tắc locked test chỉ dùng cho final evaluation, threshold tuning dùng CV/OOF, baseline chỉ dùng đối chứng, không sử dụng student-combine, không dùng ADASYN trực tiếp với categorical label encoding và không claim regression head.",
    )
    add_heading(doc, "6.2. Hạn chế", 2)
    for text in [
        "xAPI deep model chưa vượt Random Forest baseline về Macro F1, dù Recall Low vẫn cao.",
        "Đánh giá recommender là offline theo weak-supervision/rule-based reference.",
        "Chưa có dữ liệu phản hồi thực tế của sinh viên sau khi nhận khuyến nghị, nên không claim causal improvement.",
        "Student-Mat recommender chưa refresh full run do thiếu metadata checkpoint final.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "6.3. Hướng phát triển", 2)
    for text in [
        "Bổ sung dữ liệu phản hồi sau khuyến nghị để đánh giá hiệu quả can thiệp thực tế.",
        "Hoàn thiện metadata checkpoint cho Student-Mat recommender để chạy đồng nhất với xAPI và student-por.",
        "Thử nghiệm thêm các chiến lược calibration và threshold ổn định hơn cho xAPI nhưng vẫn giữ guardrail không chọn mô hình bằng locked test.",
        "Triển khai giao diện hỗ trợ cố vấn học tập xem xác suất dự đoán, rủi ro chính và lộ trình đề xuất.",
    ]:
        add_bullet(doc, text)


def references_and_appendix(doc):
    new_chapter(doc, "TÀI LIỆU THAM KHẢO")
    refs = [
        "[1] P. Cortez and A. Silva, \"Using data mining to predict secondary school student performance,\" in Proceedings of the 5th Future Business Technology Conference, 2008.",
        "[2] E. A. Amrieh, T. Hamtini, and I. Aljarah, \"Mining educational data to predict student's academic performance using ensemble methods,\" International Journal of Database Theory and Application, vol. 9, no. 8, pp. 119-136, 2016.",
        "[3] Y. LeCun, L. Bottou, Y. Bengio, and P. Haffner, \"Gradient-based learning applied to document recognition,\" Proceedings of the IEEE, vol. 86, no. 11, pp. 2278-2324, 1998.",
        "[4] S. Hochreiter and J. Schmidhuber, \"Long short-term memory,\" Neural Computation, vol. 9, no. 8, pp. 1735-1780, 1997.",
        "[5] T.-Y. Lin, P. Goyal, R. Girshick, K. He, and P. Dollar, \"Focal loss for dense object detection,\" in IEEE International Conference on Computer Vision, 2017.",
        "[6] N. V. Chawla, K. W. Bowyer, L. O. Hall, and W. P. Kegelmeyer, \"SMOTE: Synthetic minority over-sampling technique,\" Journal of Artificial Intelligence Research, vol. 16, pp. 321-357, 2002.",
        "[7] F. Pedregosa et al., \"Scikit-learn: Machine learning in Python,\" Journal of Machine Learning Research, vol. 12, pp. 2825-2830, 2011.",
        "[8] A. Paszke et al., \"PyTorch: An imperative style, high-performance deep learning library,\" in Advances in Neural Information Processing Systems, 2019.",
    ]
    for ref in refs:
        add_paragraph(doc, ref)

    new_chapter(doc, "PHỤ LỤC")
    add_heading(doc, "Phụ lục A. Các artifact final trong dự án", 2)
    for text in [
        "reports/final/final_model_manifest.json",
        "reports/final/final_deep_results_table.csv",
        "reports/final/final_baseline_comparison.csv",
        "reports/final/final_prediction_model_report.md",
        "reports/final/final_recommender_report.md",
        "reports/final/final_recommender_thesis_summary_vi.md",
        "reports/final/FINAL_PROJECT_STATUS.md",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "Phụ lục B. Lệnh kiểm thử", 2)
    add_paragraph(doc, "Lệnh kiểm thử cơ bản của dự án:")
    add_paragraph(doc, "py -3.10 -m pytest -q", italic=True)
    add_paragraph(doc, "Kết quả kiểm thử gần nhất trong quá trình chuẩn bị báo cáo: 31 passed.")

    doc.add_page_break()
    add_heading(doc, "Phụ lục C. Bảng nguồn metric dự đoán", 2)
    add_paragraph(
        doc,
        "Bảng phụ lục này lặp lại metric final theo đúng nguồn reports/final/FINAL_PROJECT_STATUS.md để người đọc có thể đối chiếu trực tiếp với các hình ở Chương 4. Các số liệu không được tính lại trong báo cáo Word mà được đọc từ artifact final đã chốt.",
    )
    add_table(
        doc,
        ["Dataset", "Scenario", "Model", "Prediction mode", "Macro F1", "Recall Low", "F1 Low"],
        [
            ["student-mat", "late", "sequence_cnn_bilstm_only", "low_f1_tuned", "0.9365", "0.9615", "0.8929"],
            ["student-por", "late", "sequence_cnn_bilstm_only", "low_f1_tuned", "0.8783", "0.9000", "0.8182"],
            ["student-por", "midterm", "sequence_cnn_bilstm_only", "argmax", "0.8228", "0.6500", "0.7429"],
            ["xAPI", "default", "gated_fusion_v28", "low_f1_tuned", "0.7541", "0.8846", "0.8214"],
        ],
        [2.0, 2.0, 4.0, 2.8, 1.4, 1.4, 1.4],
    )
    add_paragraph(
        doc,
        "Cách diễn giải: student-mat late là cấu hình deep mạnh nhất theo Macro F1. student-por late có Recall Low cao nên phù hợp khi nhấn mạnh phát hiện nhóm nguy cơ. student-por midterm có ý nghĩa ở bối cảnh cảnh báo giữa kỳ nhưng không nên dùng để claim kết quả mạnh nhất. xAPI giữ Recall Low cao nhưng Macro F1 thấp hơn baseline Random Forest.",
    )

    doc.add_page_break()
    add_heading(doc, "Phụ lục D. Bảng nguồn metric recommender", 2)
    add_paragraph(
        doc,
        "Các metric dưới đây được lấy từ outputs/recommender/xapi/recommender_metrics.json và outputs/recommender/student-por/recommender_metrics.json. Đây là đánh giá offline theo weak-supervision/rule-based reference, không phải dữ liệu phản hồi thực tế sau khuyến nghị.",
    )
    add_table(
        doc,
        ["Dataset", "Risk Macro F1", "Risk Micro F1", "Precision@3", "Recall@3", "NDCG@3", "Coverage@3"],
        [
            ["xAPI", "0.9831", "0.9813", "0.6840", "0.4720", "0.8229", "0.6500"],
            ["student-por", "0.9359", "0.9094", "0.6641", "0.3185", "0.7455", "0.5500"],
        ],
        [2.2, 2.1, 2.1, 1.8, 1.8, 1.8, 1.8],
    )
    add_table(
        doc,
        ["Dataset", "Risk Coverage", "Workload Std", "Difficulty Progression", "Prereq Violation"],
        [
            ["xAPI", "0.8958", "1.1210", "0.7153", "0.0000"],
            ["student-por", "0.9508", "1.3137", "0.6000", "0.0449"],
        ],
        [2.6, 2.8, 2.8, 3.6, 3.2],
    )
    add_paragraph(
        doc,
        "Risk Coverage đo mức độ lộ trình bao phủ các rủi ro đã chẩn đoán. Workload Std phản ánh độ phân tán khối lượng giữa các tuần. Difficulty Progression đo tính tăng dần độ khó. Prereq Violation càng thấp càng tốt vì nó cho thấy path planner không xếp can thiệp vượt điều kiện tiên quyết.",
    )

    doc.add_page_break()
    add_heading(doc, "Phụ lục E. Guardrail kiểm soát kết quả", 2)
    add_table(
        doc,
        ["Guardrail", "Trạng thái trong báo cáo", "Ý nghĩa"],
        [
            ["Locked test", "Chỉ dùng final evaluation.", "Không chọn model hoặc threshold bằng test."],
            ["Threshold tuning", "Dùng CV/OOF probabilities.", "Giảm nguy cơ overfit locked test."],
            ["Baseline ML", "Chỉ dùng đối chứng.", "Không dùng teacher, distillation, pseudo-label hoặc feature importance."],
            ["student-combine", "Không sử dụng.", "Không trộn dataset làm sai scenario."],
            ["ADASYN", "Không dùng trực tiếp trên Student label encoding.", "Tránh nội suy giá trị categorical không hợp lệ."],
            ["Regression head", "Không claim.", "RMSE chưa đủ tốt để làm kết quả chính."],
            ["Recommender", "Không dùng true G3/Class để sinh operational recommendation.", "Giữ tính triển khai thực tế."],
        ],
        [3.2, 5.0, 6.8],
    )
    add_paragraph(
        doc,
        "Các guardrail này là phần quan trọng của báo cáo vì chúng chứng minh kết quả không được làm đẹp bằng leakage. Đặc biệt, locked test chỉ xuất hiện ở giai đoạn đánh giá cuối cùng; mọi lựa chọn mô hình và threshold phải được quyết định trước đó bằng train pool hoặc CV/OOF.",
    )

    doc.add_page_break()
    add_heading(doc, "Phụ lục F. Danh sách hình metric sinh bằng code", 2)
    add_table(
        doc,
        ["File hình", "Nguồn dữ liệu", "Nội dung"],
        [
            ["fig_4_1_prediction_metrics.png", "FINAL_PROJECT_STATUS.md", "Macro F1, Recall Low, F1 Low của final champion."],
            ["fig_4_2_xapi_baseline_comparison.png", "final_baseline_comparison.csv", "Macro F1 xAPI deep vs baseline Random Forest."],
            ["fig_4_3_low_class_focus.png", "FINAL_PROJECT_STATUS.md", "Recall Low và F1 Low theo dataset/scenario."],
            ["fig_4_4_macro_f1_ranking.png", "FINAL_PROJECT_STATUS.md", "Xếp hạng Macro F1 của final champion."],
            ["fig_5_1_recommender_offline_metrics.png", "recommender_metrics.json", "Tổng hợp metric offline RA-HLPR."],
            ["fig_5_3_risk_diagnosis_metrics.png", "recommender_metrics.json", "Risk Macro F1 và Risk Micro F1."],
            ["fig_5_4_ranking_metrics.png", "recommender_metrics.json", "Precision@3, Recall@3, NDCG@3, Coverage@3."],
            ["fig_5_5_path_quality_metrics.png", "recommender_metrics.json", "Risk Coverage, Difficulty Progression, Prereq Violation."],
        ],
        [4.6, 4.2, 6.2],
    )
    add_paragraph(
        doc,
        "Các hình trên được tạo bởi scripts/build_klt_report_docx.py bằng PIL. Báo cáo không dùng ảnh metric vẽ tay, không lấy ảnh từ nguồn ngoài và không thêm số không có trong artifact final.",
    )
    doc.add_page_break()
    add_heading(doc, "Phụ lục G. Tái tạo hình metric bằng code", 2)
    add_paragraph(
        doc,
        "Phụ lục này ghi lại quy trình tái tạo hình metric để đảm bảo các biểu đồ trong báo cáo có thể kiểm chứng. Toàn bộ hình metric được sinh từ dữ liệu đã chốt trong repo, không chỉnh tay số liệu trên ảnh.",
    )
    add_table(
        doc,
        ["Bước", "Nguồn đọc", "Output sinh ra"],
        [
            ["1", "reports/final/FINAL_PROJECT_STATUS.md", "Hình prediction metrics, Low-class focus và Macro F1 ranking."],
            ["2", "reports/final/final_baseline_comparison.csv", "Hình xAPI deep final so với Random Forest baseline."],
            ["3", "outputs/recommender/*/recommender_metrics.json", "Hình recommender offline, risk diagnosis, ranking@3 và path quality."],
            ["4", "scripts/build_klt_report_docx.py", "DOCX final có bảng, hình, caption và nguồn hình."],
        ],
        [1.4, 6.8, 6.8],
    )
    add_paragraph(
        doc,
        "Khi cần cập nhật số liệu, cần cập nhật artifact nguồn rồi chạy lại script. Không nên sửa trực tiếp ảnh PNG vì sẽ làm mất truy vết giữa báo cáo và output thực nghiệm.",
    )



def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.bold = True
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 3"].font.size = Pt(13)


def main():
    save_metric_figures()
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])

    add_cover(doc)
    add_cover(doc, sub=True)
    front_matter(doc)
    chapter_1(doc)
    chapter_2(doc)
    chapter_3(doc)
    chapter_4(doc)
    chapter_5(doc)
    chapter_6(doc)
    references_and_appendix(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
