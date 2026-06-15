import os
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

def add_cover(doc, is_sub_cover=False):
    # Cover text
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("TRƯỜNG ĐẠI HỌC NGOẠI NGỮ-TIN HỌC TP. HỒ CHÍ MINH\nKHOA CNTT\n")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("_______________________________\n\n\n\n")
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("KHÓA LUẬN TỐT NGHIỆP\n\n\n\n")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(18)
    run.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("DỰ ĐOÁN KẾT QUẢ HỌC TẬP CỦA SINH VIÊN BẰNG CÁC MÔ HÌNH HỌC MÁY\n\n\n\n\n\n\n")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(24)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.left_indent = Inches(2.5)
    
    run = p.add_run("GIẢNG VIÊN HƯỚNG DẪN: TS. Nguyễn Văn A\n")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    
    run = p.add_run("SINH VIÊN THỰC HIỆN: Nguyễn Thái H – 20DH11xxxx\n")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    for _ in range(7):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("TP. HỒ CHÍ MINH – THÁNG 06 NĂM 2026")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    
    doc.add_page_break()

def create_report():
    doc = Document()
    
    # Page setup (A4, Margins: Top 3cm, Bottom 3cm, Left 3.5cm, Right 2cm)
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(3.0)

    # Set normal style font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    # Cover Pages
    add_cover(doc, is_sub_cover=False)
    add_cover(doc, is_sub_cover=True)
    
    # Helper to add chapter and filler
    def add_chapter(title, content, page_breaks=1):
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(15)
        
        doc.add_paragraph(content)
        for _ in range(page_breaks):
            doc.add_page_break()

    add_chapter("LỜI CẢM ƠN", "[Nội dung lời cảm ơn...]", 1)
    add_chapter("LỜI CAM ĐOAN", "[Nội dung lời cam đoan...]", 1)
    add_chapter("MỤC LỤC", "[Nội dung mục lục...]", 1)
    add_chapter("DANH MỤC CÁC KÝ HIỆU, CHỮ VIẾT TẮT VÀ THUẬT NGỮ", "[Nội dung...]", 1)
    add_chapter("DANH MỤC CÁC BẢNG", "[Danh mục bảng...]", 1)
    add_chapter("DANH MỤC CÁC HÌNH VẼ, ĐỒ THỊ", "[Danh mục hình vẽ...]", 1)

    # Chapter 1
    add_chapter("CHƯƠNG 1. MỞ ĐẦU", 
        "1.1. Lý do chọn đề tài\n"
        "[Phần này trình bày lý do chọn đề tài dự đoán kết quả học tập...]\n\n"
        "1.2. Mục tiêu nghiên cứu\n"
        "[Mục tiêu nghiên cứu...]\n\n"
        "1.3. Đối tượng và phạm vi nghiên cứu\n"
        "[Đối tượng và phạm vi...]\n\n"
        "1.4. Ý nghĩa khoa học và thực tiễn\n"
        "[Ý nghĩa...]\n\n", 2)
        
    # Chapter 2
    add_chapter("CHƯƠNG 2. TỔNG QUAN", 
        "2.1. Tổng quan về các phương pháp học máy trong giáo dục\n"
        "[Nội dung tổng quan...]\n\n"
        "2.2. Các nghiên cứu liên quan\n"
        "[Các nghiên cứu trước đây...]\n\n", 2)
        
    # Chapter 3 - MÔ HÌNH
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("CHƯƠNG 3. MÔ HÌNH DỰ ĐOÁN VÀ KHUYẾN NGHỊ LỘ TRÌNH HỌC TẬP")
    run.bold = True
    run.font.size = Pt(15)
    
    doc.add_paragraph("3.1. Kiến trúc tổng thể của hệ thống")
    doc.add_paragraph("Kiến trúc hệ thống dự đoán kết quả học tập của sinh viên được thiết kế với nhiều thành phần tích hợp chặt chẽ. Đầu tiên, hệ thống thu thập dữ liệu từ nhiều nguồn khác nhau, bao gồm thông tin nhân khẩu học (demographics), điểm số (grades) và các hoạt động của sinh viên trong quá trình học tập (activity). Dưới đây là hình ảnh tổng quan về kiến trúc của hệ thống:")
    
    # Add Image
    try:
        img_path = r"C:\Users\THPhu\.gemini\antigravity\brain\45358a05-baaa-4e75-92f1-9c5c6264125e\student_prediction_model_1781372236144.png"
        doc.add_picture(img_path, width=Inches(5.5))
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.add_run("Hình 3.1: Kiến trúc mô hình dự đoán kết quả học tập").italic = True
    except Exception as e:
        doc.add_paragraph(f"[Lỗi chèn hình ảnh: {e}]")

    doc.add_paragraph("\n3.2. Tiền xử lý dữ liệu (Data Preprocessing)")
    doc.add_paragraph("Dữ liệu thu thập thường chưa chuẩn hóa, có thể chứa các giá trị thiếu hoặc bị nhiễu. Bước tiền xử lý dữ liệu sẽ giúp làm sạch và chuẩn bị dữ liệu tốt nhất cho việc huấn luyện mô hình. Các phương pháp như Imputation để điền khuyết, Normalization/Standardization để chuẩn hóa thang đo được áp dụng.")

    doc.add_paragraph("\n3.3. Trích xuất đặc trưng (Feature Engineering)")
    doc.add_paragraph("Việc xây dựng các đặc trưng mới từ dữ liệu thô đóng vai trò quan trọng trong việc cải thiện hiệu suất của mô hình học máy. Ví dụ, tỷ lệ vắng mặt hoặc xu hướng thay đổi điểm số qua các kỳ học.")

    doc.add_paragraph("\n3.4. Mô hình Học máy (Machine Learning Model)")
    doc.add_paragraph("Nhiều mô hình học máy đã được thử nghiệm và đánh giá. Trong đó, mô hình học sâu lai (CNN + BiLSTM) được cấu hình và tinh chỉnh làm xương sống (backbone) cho bài toán dự đoán hiệu suất học tập. Chúng tôi cũng tiến hành tinh chỉnh siêu tham số (Hyperparameter Tuning) để tối ưu độ chính xác.")

    doc.add_paragraph("\n3.5. Hệ thống Khuyến nghị Lộ trình Học tập Hỗn hợp Thích ứng Rủi ro (RA-HLPR)")
    doc.add_paragraph("Hệ thống Khuyến nghị Lộ trình Học tập Hỗn hợp Thích ứng Rủi ro (Risk-Aware Hybrid Learning Path Recommender - RA-HLPR) hoạt động như một mô-đun hạ nguồn (downstream) độc lập, nhận đầu vào từ kết quả phân loại của mô hình chính và các đặc trưng của người học.")
    
    doc.add_paragraph("\n3.5.1. Đầu chẩn đoán rủi ro (Risk Diagnosis Head)")
    doc.add_paragraph("Đầu chẩn đoán rủi ro (Risk Diagnosis Head) là một mạng thần kinh MLP 3 lớp, nhận đầu vào là các đặc trưng của sinh viên kết hợp với phân phối xác suất dự đoán của mô hình phân loại chính. Thành phần này chẩn đoán các nguy cơ học thuật cụ thể dưới dạng các xác suất rủi ro. Số lượng đầu ra rủi ro được điều chỉnh tự động tùy thuộc vào bộ dữ liệu (6 rủi ro cho dữ liệu học sinh student-mat/por, và 3 rủi ro cho dữ liệu xapi).")
    
    doc.add_paragraph("\n3.5.2. Cơ sở tri thức can thiệp (Intervention Knowledge Base)")
    doc.add_paragraph("Cơ sở tri thức (Intervention Knowledge Base) lưu trữ các biện pháp can thiệp học thuật được chuẩn hóa trong file 'intervention_catalog.csv'. Mỗi can thiệp được định nghĩa bằng các thuộc tính như: mã can thiệp (item_id), tên biện pháp, mô tả chi tiết, nhóm rủi ro hướng tới (target_risks), độ khó sư phạm (difficulty_level), số giờ tự học ước tính hàng tuần (estimated_hours_per_week), giai đoạn đề xuất (recommended_phase), hiệu năng kỳ vọng (expected_effect) và yêu cầu kiến thức tiên quyết (prerequisite_level).")
    
    doc.add_paragraph("\n3.5.3. Chiến lược gán nhãn yếu (Weak Labeling Strategy)")
    doc.add_paragraph("Do dữ liệu thực tế không có sẵn nhãn rủi ro cụ thể của từng học sinh, phương pháp gán nhãn yếu (Weak Labeling) dựa trên tri thức chuyên gia được áp dụng để sinh nhãn huấn luyện cho đầu chẩn đoán rủi ro. Các quy tắc gán nhãn yếu được thiết kế chặt chẽ theo nguyên tắc 'Không dùng risk không có feature' nhằm tránh thiên kiến học máy. Cụ thể, đối với dữ liệu student, toàn bộ 6 rủi ro được ánh xạ thông qua các thuộc tính hiện có (failures, G1/G2, absences, freetime/goout, studytime). Đối với dữ liệu xapi, các rủi ro không có đặc trưng tương ứng (như lịch sử trượt môn, điểm số lịch sử, thời gian tự học) sẽ được loại bỏ, chỉ thực hiện gán nhãn yếu cho 3 rủi ro có dữ liệu hỗ trợ (nghỉ học, mức độ tương tác LMS, rủi ro học lực yếu). Việc huấn luyện đầu chẩn đoán được thực hiện bằng cách sử dụng hàm lỗi BCEWithLogitsLoss có trọng số pos_weight để cân bằng nhãn.")
    
    doc.add_paragraph("\n3.5.4. Bộ chấm điểm hỗn hợp (Hybrid Scorer) và Bộ lọc ứng viên (Candidate Generator)")
    doc.add_paragraph("Bộ chấm điểm hỗn hợp (Hybrid Scorer) tính điểm ưu tiên cho từng biện pháp can thiệp dựa trên công thức đa tiêu chí tối ưu: score = 0.3 * risk_match + 0.2 * performance_need + 0.15 * difficulty_fit + 0.15 * time_fit + 0.1 * prerequisite_fit + 0.1 * expected_effect. Trước khi chấm điểm, Bộ lọc ứng viên (Candidate Generator) sẽ lọc bớt các can thiệp không phù hợp với mức độ rủi ro hiện tại (xác suất rủi ro hướng tới phải từ 0.3 trở lên) và lớp học lực dự đoán để tối ưu hóa hiệu suất tính toán và tăng độ tập trung sư phạm.")
    
    doc.add_paragraph("\n3.5.5. Bộ lập lộ trình học tập (Learning Path Planner)")
    doc.add_paragraph("Bộ lập lộ trình học tập (Learning Path Planner) phân bổ các biện pháp can thiệp đã được chấm điểm vào một lộ trình 4 tuần tuần tự theo các chủ đề sư phạm tăng tiến: Tuần 1: Ổn định (Stabilize - giải quyết rào cản khẩn cấp), Tuần 2: Thực hành (Practice - bù đắp hổng kiến thức), Tuần 3: Củng cố (Reinforce - tăng tương tác học tập), Tuần 4: Đánh giá & Điều chỉnh (Evaluate & Adjust - đánh giá lại hoặc thách thức nâng cao). Đồng thời, hệ thống tự động sinh ra các diễn giải tiếng Việt thân thiện giải thích lý do cụ thể đề xuất các can thiệp này dựa trên hồ sơ của từng học sinh.")
    
    doc.add_paragraph("\n* Hạn chế của phương pháp đánh giá:")
    doc.add_paragraph("Mặc dù các chỉ số đo lường lộ trình học tập (độ phủ rủi ro, tính tăng tiến khó dần, độ vi phạm tiên quyết, độ cân bằng tải) đều đạt kết quả tốt trên dữ liệu mô phỏng, phương pháp này vẫn tồn tại hạn chế lớn là thiếu kiểm chứng thực nghiệm thực tế (longitudinal validation/A-B Testing) trên người học thực tế trong thời gian dài để chứng minh hiệu quả nâng cao kết quả học tập cuối cùng.")

    doc.add_page_break()

    # Chapter 4
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("CHƯƠNG 4. KẾT QUẢ VÀ THẢO LUẬN")
    run.bold = True
    run.font.size = Pt(15)
    
    doc.add_paragraph("4.1. Môi trường thực nghiệm\n[Trình bày về cấu hình phần cứng, phần mềm...]\n")
    doc.add_paragraph("4.2. Bộ dữ liệu thử nghiệm\n[Mô tả về bộ dữ liệu...]\n")
    doc.add_paragraph("4.3. Kết quả đánh giá mô hình học lực\n[Các bảng và biểu đồ kết quả dự đoán học lực của CNN-BiLSTM...]\n")
    
    doc.add_paragraph("4.4. Kết quả đánh giá hệ thống khuyến nghị RA-HLPR")
    doc.add_paragraph("Hệ thống RA-HLPR được đánh giá độc lập qua 3 khía cạnh: khả năng chẩn đoán rủi ro (Risk Diagnosis), hiệu suất xếp hạng can thiệp (Ranking Metrics), và chất lượng lộ trình được tạo ra (Path Quality Metrics).")
    
    # Load evaluation JSON files
    recommendations_dir = Path("C:/Huflit/kltn/outputs/recommender")
    eval_files = {
        "student-mat": recommendations_dir / "student-mat" / "recommender_metrics.json",
        "student-por": recommendations_dir / "student-por" / "recommender_metrics.json",
        "xapi": recommendations_dir / "xapi" / "recommender_metrics.json"
    }
    
    eval_data_all = {}
    for dataset_name, filepath in eval_files.items():
        if filepath.exists():
            with open(filepath, "r", encoding="utf-8") as f:
                eval_data_all[dataset_name] = json.load(f)
                
    def format_decimal(val, digits=4):
        if val is None:
            return "N/A"
        return f"{val:.{digits}f}".replace(".", ",")
        
    # Add table 4.1: Risk Diagnosis & Ranking metrics
    p_title = doc.add_paragraph()
    p_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run_title = p_title.add_run("Bảng 4.1: Kết quả chẩn đoán rủi ro và xếp hạng can thiệp của RA-HLPR")
    run_title.bold = True
    
    headers = ["Bộ dữ liệu", "Micro F1", "Macro F1", "Precision@3", "NDCG@3", "Catalog Coverage"]
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    # Header formatting
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.bold = True
                
    if eval_data_all:
        for dataset_key in ["student-mat", "student-por", "xapi"]:
            eval_data = eval_data_all.get(dataset_key)
            if not eval_data: continue
            dataset = eval_data.get("dataset", dataset_key)
            risk = eval_data.get("risk_diagnosis", {})
            ranking = eval_data.get("ranking", {})
            
            row_cells = table.add_row().cells
            row_cells[0].text = dataset
            row_cells[1].text = format_decimal(risk.get("f1_micro"))
            row_cells[2].text = format_decimal(risk.get("f1_macro"))
            row_cells[3].text = format_decimal(ranking.get("precision_at_3"))
            row_cells[4].text = format_decimal(ranking.get("ndcg_at_3"))
            row_cells[5].text = format_decimal(ranking.get("coverage_at_3"))
            
            for col_idx, cell in enumerate(row_cells):
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
                        
    doc.add_paragraph() # Spacing
    
    # Add table 4.2: Path Quality metrics
    p_title2 = doc.add_paragraph()
    p_title2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run_title2 = p_title2.add_run("Bảng 4.2: Đánh giá chất lượng lộ trình học tập 4 tuần")
    run_title2.bold = True
    
    headers2 = ["Bộ dữ liệu", "Độ phủ Rủi ro", "Độ khó Tăng tiến", "Vi phạm Tiên quyết", "Cân bằng Tải học tập"]
    table2 = doc.add_table(rows=1, cols=5)
    table2.style = 'Table Grid'
    
    hdr_cells2 = table2.rows[0].cells
    for i, header in enumerate(headers2):
        hdr_cells2[i].text = header
        for paragraph in hdr_cells2[i].paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.bold = True
                
    if eval_data_all:
        for dataset_key in ["student-mat", "student-por", "xapi"]:
            eval_data = eval_data_all.get(dataset_key)
            if not eval_data: continue
            dataset = eval_data.get("dataset", dataset_key)
            pq = eval_data.get("path_quality", {})
            
            row_cells = table2.add_row().cells
            row_cells[0].text = dataset
            row_cells[1].text = format_decimal(pq.get("risk_coverage_rate"))
            row_cells[2].text = format_decimal(pq.get("difficulty_progression_rate"))
            row_cells[3].text = format_decimal(pq.get("prerequisite_violation_rate"))
            row_cells[4].text = format_decimal(pq.get("workload_balance_std"))
            
            for col_idx, cell in enumerate(row_cells):
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)

    doc.add_paragraph() # Spacing
    doc.add_paragraph("4.5. Thảo luận\n[Phân tích và thảo luận kết quả nghiên cứu...]\n")
    doc.add_page_break()

    # Chapter 5
    add_chapter("CHƯƠNG 5. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", 
        "5.1. Kết luận\n"
        "[Tóm tắt những kết quả đạt được...]\n\n"
        "5.2. Hướng phát triển\n"
        "Hệ thống RA-HLPR đã chứng minh được tính khả thi trong việc cá nhân hóa quá trình học. Tuy nhiên, việc đánh giá chất lượng lộ trình hiện tại thuần túy dùng thuật toán đo lường tính hợp lý logic. Trong tương lai, việc thực hiện A-B Testing thực tế trên sinh viên hoặc tích hợp Human Evaluation từ các chuyên gia tâm lý giáo dục sẽ giúp hoàn thiện hệ thống hơn.\n\n", 1)

    # References
    add_chapter("TÀI LIỆU THAM KHẢO", 
        "[1] Nguyễn Văn A, Tên sách, Nhà xuất bản, Năm.\n"
        "[2] Tên Tác Giả, \"Tên bài báo\", Tên Tạp chí, Số, Trang, Năm.\n", 1)

    # Appendix
    add_chapter("PHỤ LỤC", "[Các bảng phụ lục, mã nguồn...]", 1)

    doc.save(r"C:\Huflit\kltn\Bao_cao_cuoi_cung.docx")
    print("Done")

if __name__ == '__main__':
    create_report()
