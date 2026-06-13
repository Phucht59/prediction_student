from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"C:\Huflit\kltn")
OUT = ROOT / "reports" / "final" / "Luan_van_du_doan_va_khuyen_nghi_thanh_tich_hoc_tap_cap_nhat.docx"
FIG_DIR = ROOT / "reports" / "final" / "figures" / "thesis_revised"

FONT = "Times New Roman"
NAVY = "000000"
BLUE = "000000"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D9D9D9"
TEXT_GRAY = "595959"
WHITE = "FFFFFF"
RED = "9C0006"
CONTENT_WIDTH_DXA = 9072


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent_dxa: int = 120) -> None:
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {CONTENT_WIDTH_DXA}: {widths}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=13, bold=False, italic=False, color="000000") -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Trang ")
    set_run_font(run, size=10, color=TEXT_GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_field(paragraph, instruction: str, display_text: str = "Nhấn F9 để cập nhật") -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display_text
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])


def set_page_number_format(section, fmt: str, start: int = 1) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:fmt"), fmt)
    pg_num.set(qn("w:start"), str(start))


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(13)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.first_line_indent = Cm(1)

    for style_name, size, color, before, after in (
        ("Heading 1", 17, NAVY, 14, 8),
        ("Heading 2", 15, BLUE, 12, 6),
        ("Heading 3", 13.5, NAVY, 10, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.first_line_indent = Cm(0)

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(13)
        style.paragraph_format.left_indent = Cm(0.75)
        style.paragraph_format.first_line_indent = Cm(-0.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.35

    caption = doc.styles["Caption"]
    caption.font.name = FONT
    caption._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    caption.font.size = Pt(12)
    caption.font.bold = True
    caption.font.italic = False
    caption.font.color.rgb = RGBColor.from_string("000000")
    for custom_name in ("Figure Caption", "Table Caption"):
        if custom_name not in doc.styles:
            custom = doc.styles.add_style(custom_name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            custom = doc.styles[custom_name]
        custom.font.name = FONT
        custom._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        custom._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        custom._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        custom.font.size = Pt(12)
        custom.font.bold = True
        custom.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        custom.paragraph_format.first_line_indent = Cm(0)
        custom.paragraph_format.space_before = Pt(3)
        custom.paragraph_format.space_after = Pt(3)


def configure_sections(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(3.0)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(2.25)
    section.footer_distance = Cm(1.5)


def finalize_sections(doc: Document) -> None:
    settings = doc.settings._element
    for item in settings.findall(".//" + qn("w:compatSetting")):
        if item.get(qn("w:name")) == "compatibilityMode":
            item.set(qn("w:val"), "15")

    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(3.0)
        section.left_margin = Cm(3.5)
        section.right_margin = Cm(2.0)
        section.header_distance = Cm(2.25)
        section.footer_distance = Cm(1.5)
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False

    cover = doc.sections[0]
    cover.header.paragraphs[0].clear()
    cover.footer.paragraphs[0].clear()

    prelim = doc.sections[1]
    prelim.header.paragraphs[0].clear()
    prelim.footer.paragraphs[0].clear()
    add_page_number(prelim.footer.paragraphs[0])
    set_page_number_format(prelim, "lowerRoman", 1)

    main = doc.sections[2]
    hp = main.header.paragraphs[0]
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    r = hp.add_run()
    set_run_font(r, size=11, color=TEXT_GRAY)
    add_field(hp, ' STYLEREF "Heading 1" ', "TÊN CHƯƠNG")
    fp = main.footer.paragraphs[0]
    fp.clear()
    add_page_number(fp)
    set_page_number_format(main, "decimal", 1)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("TRƯỜNG ĐẠI HỌC NGOẠI NGỮ - TIN HỌC TP. HỒ CHÍ MINH")
    set_run_font(r, size=14, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(65)
    r = p.add_run("KHOA CÔNG NGHỆ THÔNG TIN")
    set_run_font(r, size=13, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("LUẬN VĂN TỐT NGHIỆP")
    set_run_font(r, size=18, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(
        "DỰ ĐOÁN VÀ KHUYẾN NGHỊ THÀNH TÍCH HỌC TẬP CỦA SINH VIÊN "
        "BẰNG KIẾN TRÚC LAI CNN-BiLSTM VÀ CONTEXT MLP"
    )
    set_run_font(r, size=22, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(55)
    r = p.add_run("Predicting and Recommending Student Academic Performance using a Hybrid CNN-BiLSTM + Context MLP Architecture")
    set_run_font(r, size=12, italic=True, color=TEXT_GRAY)

    metadata = [
        ("Sinh viên thực hiện", "[HỌ VÀ TÊN SINH VIÊN]"),
        ("Mã số sinh viên", "[MSSV]"),
        ("Giảng viên hướng dẫn", "[HỌ VÀ TÊN GIẢNG VIÊN]"),
        ("Ngành / Chuyên ngành", "[TÊN NGÀNH / CHUYÊN NGÀNH]"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(5)
        r1 = p.add_run(f"{label}: ")
        set_run_font(r1, size=12, bold=True, color=NAVY)
        r2 = p.add_run(value)
        set_run_font(r2, size=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(65)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Thành phố Hồ Chí Minh, 2026")
    set_run_font(r, size=13, bold=True, color=NAVY)
    doc.add_page_break()


def add_subcover_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("TRƯỜNG ĐẠI HỌC NGOẠI NGỮ - TIN HỌC TP. HỒ CHÍ MINH")
    set_run_font(r, size=14, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(60)
    r = p.add_run("KHOA CÔNG NGHỆ THÔNG TIN")
    set_run_font(r, size=13, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(15)
    r = p.add_run("KHÓA LUẬN TỐT NGHIỆP")
    set_run_font(r, size=18, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(60)
    r = p.add_run("DỰ ĐOÁN VÀ KHUYẾN NGHỊ THÀNH TÍCH HỌC TẬP CỦA SINH VIÊN BẰNG KIẾN TRÚC LAI CNN-BiLSTM VÀ CONTEXT MLP")
    set_run_font(r, size=20, bold=True, color=NAVY)
    for label, value in [
        ("Giảng viên hướng dẫn", "[HỌ VÀ TÊN GIẢNG VIÊN]"),
        ("Sinh viên thực hiện", "[HỌ VÀ TÊN SINH VIÊN]"),
        ("Mã số sinh viên", "[MSSV]"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(6)
        r1 = p.add_run(f"{label}: ")
        set_run_font(r1, size=13, bold=True)
        r2 = p.add_run(value)
        set_run_font(r2, size=13)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(100)
    r = p.add_run("TP. HỒ CHÍ MINH - 2026")
    set_run_font(r, size=13, bold=True)
    doc.add_section(WD_SECTION.NEW_PAGE)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text)
    set_run_font(r, size={1: 17, 2: 15, 3: 13.5}[level], bold=True, color=NAVY if level != 2 else BLUE)


def add_para(doc: Document, text: str, *, bold_lead: str | None = None, italic=False, align=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
        p.paragraph_format.first_line_indent = Cm(0)
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, italic=italic)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.first_line_indent = Cm(-0.5)
    r = p.add_run(text)
    set_run_font(r)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.first_line_indent = Cm(-0.5)
    r = p.add_run(text)
    set_run_font(r)


def add_placeholder(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5)
    clean_text = text.replace("[REQUIRES EMPIRICAL DATA - DO NOT HALLUCINATE]", "Giới hạn dữ liệu:").strip()
    r = p.add_run(clean_text)
    set_run_font(r, size=11.5, italic=True, color=TEXT_GRAY)


def add_figure(doc: Document, filename: str, caption: str, source: str = "Nguồn: Tác giả tổng hợp từ kết quả thực nghiệm của hệ thống.", width_cm: float = 15.0) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    shape = run.add_picture(str(FIG_DIR / filename), width=Cm(width_cm))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", caption)
    doc_pr.set("title", caption)
    cap = doc.add_paragraph()
    cap.style = doc.styles["Figure Caption"]
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.paragraph_format.space_after = Pt(2)
    r = cap.add_run(caption)
    set_run_font(r, size=12, bold=True)
    src = doc.add_paragraph()
    src.alignment = WD_ALIGN_PARAGRAPH.CENTER
    src.paragraph_format.first_line_indent = Cm(0)
    src.paragraph_format.space_after = Pt(8)
    r = src.add_run(source)
    set_run_font(r, size=10.5, italic=True, color=TEXT_GRAY)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], caption: str) -> None:
    p = doc.add_paragraph(style="Table Caption")
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(caption)
    set_run_font(r, size=11.5, bold=True, color=NAVY)

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for j, header in enumerate(headers):
        cell = table.cell(0, j)
        set_cell_shading(cell, "D9D9D9")
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=10.2, bold=True, color="000000")

    for i, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        for j, value in enumerate(row):
            cell = cells[j]
            if i % 2 == 0:
                set_cell_shading(cell, "F8FAFC")
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            set_run_font(r, size=10.2)
    set_table_geometry(table, widths)


def chapter_break(doc: Document) -> None:
    doc.add_page_break()


def build_document() -> Document:
    doc = Document()
    configure_styles(doc)
    configure_sections(doc)
    doc.core_properties.title = "Dự đoán và khuyến nghị thành tích học tập bằng CNN-BiLSTM + Context MLP"
    doc.core_properties.subject = "Luận văn khoa học về dự đoán thành tích học tập"
    doc.core_properties.author = "[HỌ VÀ TÊN SINH VIÊN]"
    doc.core_properties.keywords = "Educational Data Mining, CNN, BiLSTM, MLP, SMOTE, ADASYN, Optuna"

    add_title_page(doc)
    add_subcover_page(doc)

    add_heading(doc, "LỜI CẢM ƠN", 1)
    add_para(
        doc,
        "Trong quá trình thực hiện khóa luận, tác giả đã nhận được sự hướng dẫn, hỗ trợ và góp ý từ giảng viên, nhà trường, gia đình và bạn bè. Tác giả xin bày tỏ lòng biết ơn sâu sắc đến giảng viên hướng dẫn [HỌ VÀ TÊN GIẢNG VIÊN], người đã định hướng phương pháp nghiên cứu, góp ý về thiết kế thực nghiệm và hỗ trợ tác giả hoàn thiện báo cáo theo chuẩn mực khoa học."
    )
    add_para(
        doc,
        "Tác giả trân trọng cảm ơn Trường Đại học Ngoại ngữ - Tin học Thành phố Hồ Chí Minh và Khoa Công nghệ Thông tin đã cung cấp môi trường học tập, nền tảng kiến thức và điều kiện cần thiết để triển khai đề tài. Tác giả cũng cảm ơn gia đình và bạn bè đã động viên trong suốt quá trình thu thập tài liệu, xây dựng hệ thống, thực hiện thí nghiệm và biên soạn khóa luận."
    )
    add_para(
        doc,
        "Mặc dù đã nỗ lực kiểm tra nội dung và kết quả, khóa luận khó tránh khỏi những hạn chế. Tác giả kính mong nhận được các ý kiến phản biện để tiếp tục hoàn thiện hệ thống và phương pháp nghiên cứu trong tương lai."
    )

    chapter_break(doc)
    add_heading(doc, "LỜI CAM ĐOAN", 1)
    add_para(
        doc,
        "Tác giả cam đoan khóa luận này là kết quả nghiên cứu và triển khai của cá nhân dưới sự hướng dẫn của giảng viên hướng dẫn. Các kết quả thực nghiệm được trích xuất từ mã nguồn, tệp metric, tệp dự đoán và nhật ký huấn luyện trong dự án; không có số liệu đánh giá nào được tự ý tạo ra hoặc sửa đổi nhằm làm tăng kết quả."
    )
    add_para(
        doc,
        "Các tài liệu, công trình và dữ liệu tham khảo được ghi nguồn theo chuẩn IEEE. Những nội dung kế thừa từ nghiên cứu trước được diễn giải và trích dẫn rõ ràng. Tác giả chịu trách nhiệm về tính trung thực của nội dung, tính tái lập của quy trình thực nghiệm và việc sử dụng kết quả đúng mục đích học thuật."
    )
    add_para(doc, "TP. Hồ Chí Minh, tháng ... năm 2026", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_para(doc, "Sinh viên thực hiện\n[HỌ VÀ TÊN SINH VIÊN]", align=WD_ALIGN_PARAGRAPH.RIGHT)

    chapter_break(doc)
    add_heading(doc, "TÓM TẮT", 1)
    add_para(
        doc,
        "Nghiên cứu xây dựng hệ thống phân loại thành tích học tập theo ba mức Low, Medium và High, sau đó sinh lộ trình hỗ trợ theo luật. Mô hình gồm hai nhánh. Nhánh CNN-BiLSTM xử lý G1-G2 của hai tập Student hoặc bốn chỉ báo tương tác của xAPI. Nhánh Context MLP nhận các biến số và embedding của biến phân loại. Hai vector được ghép trước tầng dự đoán. Student-Mat và Student-Por dùng đầu ra softmax ba lớp; xAPI dùng hai logit ordinal và chuyển chúng thành xác suất của ba mức. Locked test 20% được tách trước khi tối ưu. Trong huấn luyện, validation được tách trước resampling; Student dùng ADASYN, còn xAPI dùng SMOTENC."
    )
    add_para(
        doc,
        "Trên locked test, F1-Macro đạt 0,8690 với Student-Mat, 0,8156 với Student-Por và 0,7850 với xAPI. Accuracy lần lượt là 0,8608; 0,8462; và 0,7813. So với điểm CV tốt nhất, F1 giảm 0,0345; 0,0648; và 0,0382. Student-Por có chênh lệch CV-test lớn nhất và lớp Low còn khó nhận diện. xAPI cải thiện so với lần chạy trước, nhưng Recall của lớp High chỉ đạt 0,6786. Các kết quả này được xem là kết quả của một lần phân hoạch locked test; chưa có khoảng tin cậy để kết luận về độ ổn định giữa nhiều phép chia dữ liệu."
    )
    add_para(
        doc,
        "Sau bước dự đoán, bộ luật đọc các dấu hiệu như chuyên cần, kết quả giữa kỳ, thời gian học và mức tương tác để tạo kế hoạch theo tuần. PostgreSQL lưu phiên chạy, xác suất, metric và khuyến nghị. Phần khuyến nghị mới dừng ở mức hỗ trợ cố vấn; nghiên cứu chưa đo được tác động của lộ trình đối với kết quả học tập thực tế."
    )
    add_para(doc, "Từ khóa: Educational Data Mining; dự đoán thành tích học tập; CNN-BiLSTM; Context MLP; F1-Macro; SMOTE; ADASYN; Optuna; learning path.", italic=True)

    chapter_break(doc)
    add_heading(doc, "ABSTRACT", 1)
    add_para(
        doc,
        "This thesis develops a three-class student-performance prediction system and a rule-based learning-path module. The model combines a CNN-BiLSTM sequence branch with a context MLP using numerical variables and categorical embeddings. Student-Mat and Student-Por use a three-class softmax output, whereas xAPI uses two ordinal logits. A locked 20% test set is isolated before optimization, and validation is split before resampling. The locked-test F1-Macro scores are 0.8690 for Student-Mat, 0.8156 for Student-Por and 0.7850 for xAPI. The corresponding best-CV gaps are 0.0345, 0.0648 and 0.0382. These results describe the current experimental split and do not establish statistical significance. A deterministic rule engine produces staged recommendations, while PostgreSQL stores predictions, metrics and recommendation records."
    )

    chapter_break(doc)
    add_heading(doc, "MỤC LỤC", 1)
    toc = doc.add_paragraph()
    toc.paragraph_format.first_line_indent = Cm(0)
    add_field(toc, ' TOC \\o "1-3" \\h \\z \\u ', "Mục lục sẽ được cập nhật tự động trong Microsoft Word.")

    chapter_break(doc)
    add_heading(doc, "DANH MỤC BẢNG", 1)
    lot = doc.add_paragraph()
    lot.paragraph_format.first_line_indent = Cm(0)
    add_field(lot, ' TOC \\h \\z \\t "Table Caption,1" ', "Danh mục bảng sẽ được cập nhật tự động trong Microsoft Word.")

    chapter_break(doc)
    add_heading(doc, "DANH MỤC HÌNH", 1)
    lof = doc.add_paragraph()
    lof.paragraph_format.first_line_indent = Cm(0)
    add_field(lof, ' TOC \\h \\z \\t "Figure Caption,1" ', "Danh mục hình sẽ được cập nhật tự động trong Microsoft Word.")

    add_heading(doc, "DANH MỤC TỪ VIẾT TẮT", 1)
    add_table(
        doc,
        ["Từ viết tắt", "Diễn giải"],
        [
            ["CNN", "Convolutional Neural Network"],
            ["BiLSTM", "Bidirectional Long Short-Term Memory"],
            ["MLP", "Multilayer Perceptron"],
            ["EDM", "Educational Data Mining"],
            ["SMOTE", "Synthetic Minority Over-sampling Technique"],
            ["ADASYN", "Adaptive Synthetic Sampling"],
            ["CV", "Cross-validation"],
            ["TPE", "Tree-structured Parzen Estimator"],
            ["RMSE", "Root Mean Squared Error"],
            ["R²", "Coefficient of Determination"],
            ["LMS", "Learning Management System"],
        ],
        [2100, 6972],
        "Bảng 0.1. Danh mục từ viết tắt",
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "CHƯƠNG 1. GIỚI THIỆU", 1)
    add_heading(doc, "1.1. Bối cảnh nghiên cứu", 2)
    add_para(
        doc,
        "Chuyển đổi số trong giáo dục tạo ra lượng lớn dữ liệu về điểm số, chuyên cần, hoạt động trên hệ thống quản lý học tập, đặc điểm xã hội và mức độ tham gia của người học. Educational Data Mining và Learning Analytics sử dụng các nguồn dữ liệu này để mô tả hành vi, phát hiện rủi ro và hỗ trợ can thiệp sớm. Trong thực tế, giá trị của hệ thống không chỉ nằm ở việc dự đoán đúng nhãn thành tích, mà còn ở khả năng giải thích tín hiệu rủi ro, duy trì khả năng khái quát hóa khi chuyển sang người học chưa quan sát, và chuyển kết quả dự đoán thành hành động giáo dục có thể kiểm tra."
    )
    add_para(
        doc,
        "Dữ liệu giáo dục thường đồng thời chứa hai cấu trúc. Thứ nhất là cấu trúc có thứ tự hoặc gần thời gian, chẳng hạn điểm giai đoạn G1-G2 hoặc các chỉ báo tương tác được sắp theo một quy ước nhất định. Thứ hai là dữ liệu bảng mô tả bối cảnh, chẳng hạn thời gian học, số lần trượt, hỗ trợ gia đình, loại trường, chủ đề, mức độ hài lòng và chuyên cần. Một mô hình chỉ dùng MLP có thể bỏ qua quan hệ theo thứ tự; ngược lại, một mô hình tuần tự thuần túy khó tận dụng đầy đủ các biến bối cảnh dị thể. Vì vậy, kiến trúc lai hai nhánh là lựa chọn hợp lý về mặt thiết kế."
    )

    add_heading(doc, "1.2. Phát biểu bài toán", 2)
    add_para(
        doc,
        "Bài toán chính là phân loại mỗi người học vào một trong ba mức thành tích Low, Medium hoặc High. Đối với Student-Mat và Student-Por, điểm cuối kỳ G3 được rời rạc hóa thành Low = 0-9, Medium = 10-14 và High = 15-20. Đối với xAPI, nhãn L, M và H có sẵn được ánh xạ thành 0, 1 và 2. Đầu ra của mô hình bao gồm nhãn dự đoán, vector xác suất ba lớp và độ tin cậy. Đầu ra của mô-đun khuyến nghị là một lộ trình theo giai đoạn, được tạo từ tổ hợp giữa lớp dự đoán, độ tin cậy và các yếu tố rủi ro quan sát được."
    )
    add_para(doc, "Các thách thức khoa học và kỹ thuật trọng yếu gồm:")
    add_bullet(doc, "Mất cân bằng lớp làm Accuracy có thể che khuất chất lượng ở lớp thiểu số; do đó F1-Macro là chỉ số trung tâm.")
    add_bullet(doc, "Dữ liệu nhỏ làm mô hình sâu nhạy với seed, lựa chọn siêu tham số và cách chia tập.")
    add_bullet(doc, "Dữ liệu hỗn hợp số-phân loại gây khó khăn cho nội suy tổng hợp của SMOTE/ADASYN.")
    add_bullet(doc, "Chuỗi Student chỉ dài hai bước G1-G2; chuỗi xAPI gồm bốn chỉ báo tổng hợp, không phải nhật ký theo timestamp.")
    add_bullet(doc, "Kết quả CV có thể lạc quan do lựa chọn cấu hình tốt nhất trong nhiều trial; locked test phải được cô lập hoàn toàn.")
    add_bullet(doc, "Khuyến nghị theo luật cần được phân biệt với can thiệp đã được kiểm chứng nhân quả.")

    add_heading(doc, "1.3. Mục tiêu nghiên cứu", 2)
    objectives = [
        "Xây dựng mô hình lai CNN-BiLSTM + Context MLP cho phân loại thành tích học tập ba lớp.",
        "Thiết kế quy trình dữ liệu chống rò rỉ, bao gồm locked test, fit tiền xử lý trên train, chọn đặc trưng trong train và resampling chỉ trên train.",
        "Tối ưu siêu tham số bằng Optuna và Stratified 5-fold CV, sau đó huấn luyện ensemble nhiều seed.",
        "Đánh giá bằng Accuracy, Precision-Macro, Recall-Macro, F1-Macro, RMSE và R² trên locked test.",
        "Phân tích chênh lệch CV-test, mất cân bằng lớp, ma trận nhầm lẫn và giới hạn năng lực biểu diễn.",
        "Sinh learning path theo luật và tích hợp PostgreSQL để theo dõi toàn bộ vòng đời dự đoán.",
    ]
    for item in objectives:
        add_number(doc, item)

    add_heading(doc, "1.4. Câu hỏi nghiên cứu", 2)
    add_number(doc, "Kiến trúc lai có đạt F1-Macro ổn định trên ba bộ dữ liệu có cấu trúc khác nhau hay không?")
    add_number(doc, "Khoảng cách giữa CV tốt nhất và locked test phản ánh mức độ khái quát hóa như thế nào?")
    add_number(doc, "Mất cân bằng lớp được phản ánh ra sao trong Precision, Recall, F1 theo lớp và ma trận nhầm lẫn?")
    add_number(doc, "Những giới hạn biểu diễn nào xuất hiện khi áp dụng CNN-BiLSTM cho chuỗi rất ngắn hoặc chỉ báo tổng hợp xAPI?")
    add_number(doc, "Làm thế nào chuyển dự đoán thành lộ trình học tập có thể truy vết mà không tuyên bố quá mức về hiệu quả can thiệp?")

    add_heading(doc, "1.5. Phạm vi và đóng góp", 2)
    add_para(
        doc,
        "Nghiên cứu giới hạn ở ba bộ dữ liệu công khai Student-Mat, Student-Por và xAPI-Edu-Data; nhiệm vụ đầu ra là phân loại ba lớp. Kiến trúc được giới hạn ở CNN, BiLSTM, attention pooling, MLP và fusion dense. Hệ thống không sử dụng Transformer, DeepFM, DCN-V2 hoặc tìm kiếm phản thực tế trong mô hình chính. Đóng góp thực hành nằm ở một pipeline end-to-end tái lập được, tách locked test, lưu metric và dự đoán, sinh khuyến nghị theo luật và cung cấp schema PostgreSQL. Đóng góp khoa học nằm ở việc đánh giá nghiêm ngặt chênh lệch CV-test và công khai các giới hạn của dữ liệu, biểu diễn và giao thức."
    )

    chapter_break(doc)
    add_heading(doc, "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT VÀ BỐI CẢNH NGHIÊN CỨU", 1)
    add_heading(doc, "2.1. Educational Data Mining và dự đoán thành tích", 2)
    add_para(
        doc,
        "Educational Data Mining nghiên cứu các phương pháp khai phá mẫu từ dữ liệu giáo dục để hỗ trợ hiểu biết về người học và quá trình học. Bộ Student Performance của UCI được xây dựng từ dữ liệu trường trung học Bồ Đào Nha và thường được dùng để dự đoán điểm G3 từ đặc điểm cá nhân, gia đình, nhà trường và điểm các giai đoạn trước [1]. xAPI-Edu-Data được thu thập từ môi trường Kalboard 360, nhấn mạnh các thuộc tính hành vi như raisedhands, VisITedResources, AnnouncementsView và Discussion [2]. Hai họ dữ liệu đại diện cho hai tình huống: dự đoán dựa mạnh vào lịch sử điểm và dự đoán dựa trên tương tác LMS cùng bối cảnh."
    )

    add_heading(doc, "2.2. CNN cho quan hệ cục bộ", 2)
    add_para(
        doc,
        "Convolution một chiều áp dụng bộ lọc trượt trên trục thứ tự để học các mẫu cục bộ. Với đầu vào X thuộc R^(T×C), một kênh đầu ra tại vị trí t có thể được mô tả khái quát bởi h_t = sigma(W * X_(t:t+k-1) + b). Trong hệ thống, Conv1D được theo sau bởi BatchNorm, ReLU và Dropout. Vai trò dự kiến là tái mã hóa sự kết hợp cục bộ giữa các mốc điểm hoặc chỉ báo tương tác trước khi đưa vào BiLSTM. Tuy nhiên, khi T chỉ bằng 2 hoặc 4, không gian mẫu cục bộ rất hạn chế; đây là một giới hạn cần được đánh giá thay vì mặc định rằng CNN luôn mang lại lợi ích."
    )

    add_heading(doc, "2.3. LSTM và BiLSTM", 2)
    add_para(
        doc,
        "LSTM sử dụng các cổng vào, quên và đầu ra để điều khiển dòng thông tin qua trạng thái ô nhớ, qua đó giảm khó khăn gradient suy giảm trong RNN truyền thống [3]. BiLSTM mở rộng bằng hai hướng xử lý: hướng thuận và hướng nghịch; biểu diễn tại mỗi vị trí là phép ghép hai trạng thái [4]. Trong bài toán này, BiLSTM không được dùng để dự báo tương lai theo thời gian thực, mà để mã hóa toàn bộ chuỗi đầu vào đã quan sát trước thời điểm dự đoán. Vì vậy, điều kiện triển khai là mọi phần tử của chuỗi phải sẵn có tại thời điểm suy luận, tránh rò rỉ thông tin tương lai."
    )

    add_heading(doc, "2.4. Context MLP và fusion", 2)
    add_para(
        doc,
        "Nhánh Context MLP nhận các biến số sau chuẩn hóa Min-Max và các biến phân loại sau Label Encoding. Trong triển khai hiện tại, biến phân loại được chuẩn hóa theo cardinality thay vì embedding; sau đó vector context đi qua hai tầng tuyến tính với ReLU và Dropout. Vector chuỗi từ attention pooling được ghép với vector context, qua tầng fusion và classifier ba lớp. Cách thiết kế này đơn giản, phù hợp dữ liệu nhỏ, nhưng việc ánh xạ danh mục thành số có thể tạo quan hệ thứ tự giả giữa các giá trị không có thứ tự."
    )

    add_heading(doc, "2.5. Xử lý mất cân bằng", 2)
    add_para(
        doc,
        "SMOTE tạo mẫu tổng hợp bằng nội suy giữa một quan sát lớp thiểu số và các láng giềng gần [5]. ADASYN phân bổ nhiều mẫu tổng hợp hơn cho các vùng khó phân loại [6]. Hệ thống đồng thời sử dụng trọng số lớp trong CrossEntropyLoss. Hai cơ chế cùng hướng đến việc giảm thiên lệch về lớp phổ biến, nhưng cũng có thể làm thay đổi ranh giới quyết định quá mức nếu không có ablation. Đặc biệt, nội suy trực tiếp trên mã số của biến phân loại không bảo toàn ngữ nghĩa danh mục; lựa chọn phù hợp hơn trong tương lai là SMOTENC hoặc resampling tách biệt theo loại biến."
    )

    add_heading(doc, "2.6. Optuna và tối ưu siêu tham số", 2)
    add_para(
        doc,
        "Optuna dùng TPE và MedianPruner để tìm siêu tham số [7]. Phiên bản hiện tại đánh giá mỗi trial bằng Repeated Stratified 5-fold CV với ba lần lặp trên train pool. Điểm tốt nhất dùng để chọn cấu hình, còn locked test chỉ được dùng sau khi cấu hình đã cố định."
    )

    chapter_break(doc)
    add_heading(doc, "CHƯƠNG 3. PHƯƠNG PHÁP NGHIÊN CỨU VÀ KIẾN TRÚC HỆ THỐNG", 1)
    add_heading(doc, "3.1. Quy trình nghiên cứu tổng thể", 2)
    add_number(doc, "Nạp dữ liệu gốc và chuyển nhãn về bài toán ba lớp.")
    add_number(doc, "Tách stratified locked test 20% với random seed 42 trước tối ưu.")
    add_number(doc, "Trong từng fold CV: feature engineering, fit encoder/scaler trên train fold, resampling train fold, chọn đặc trưng trên train fold và transform validation fold.")
    add_number(doc, "Optuna tối đa hóa F1-Macro của repeated stratified 5-fold CV với ba lần lặp.")
    add_number(doc, "Tách validation trước resampling trong từng seed, huấn luyện ensemble và trung bình hóa xác suất.")
    add_number(doc, "Đánh giá locked test một lần, xuất metric, dự đoán, feature importance và learning path.")
    add_number(doc, "Ghi phiên chạy, dự đoán, metric và khuyến nghị vào PostgreSQL trong một giao dịch.")
    add_figure(doc, "01_system_pipeline.png", "Hình 3.1. Quy trình tổng thể của hệ thống dự đoán và khuyến nghị", width_cm=15.2)

    add_heading(doc, "3.2. Kiến trúc hai nhánh", 2)
    add_table(
        doc,
        ["Khối", "Đầu vào", "Phép biến đổi", "Đầu ra"],
        [
            ["Nhánh tuần tự", "G1-G2 hoặc 4 chỉ báo xAPI", "Conv1D → BatchNorm → ReLU → Dropout → BiLSTM → Attention Pooling", "Vector chuỗi 2h"],
            ["Nhánh bối cảnh", "Biến số và biến phân loại", "Min-Max cho biến số; embedding cho biến phân loại; MLP hai tầng", "Vector context"],
            ["Fusion", "Vector chuỗi + context", "Concatenate → Linear → ReLU → Dropout", "Vector hợp nhất"],
            ["Đầu ra", "Vector hợp nhất", "Student: 3 logits; xAPI: 2 logits ordinal", "P(Low), P(Medium), P(High)"],
        ],
        [1500, 2200, 3472, 1900],
        "Bảng 3.1. Cấu trúc chức năng của mô hình lai",
    )
    add_figure(doc, "02_hybrid_architecture.png", "Hình 3.2. Kiến trúc nội bộ CNN-BiLSTM + Context MLP", width_cm=15.2)

    add_heading(doc, "3.3. Nhánh CNN-BiLSTM", 2)
    add_para(
        doc,
        "Tensor đầu vào tuần tự có dạng batch × T × 1. Trước Conv1D, tensor được chuyển thành batch × 1 × T. Conv1D dùng padding bằng floor(kernel/2), vì vậy giữ gần nguyên chiều dài đối với kernel lẻ. Đầu ra được chuyển về batch × T × channels và đưa qua LSTM hai chiều. Nếu kích thước ẩn mỗi hướng là h, đầu ra tại mỗi bước có kích thước 2h. AttentionPooling1D học một điểm số vô hướng cho từng bước, chuẩn hóa bằng softmax và tính tổng có trọng số để thu vector chuỗi."
    )
    add_para(
        doc,
        "Với Student-Mat và Student-Por, T = 2 tương ứng G1 và G2. Với xAPI, T = 4 tương ứng raisedhands, VisITedResources, AnnouncementsView và Discussion. Cần nhấn mạnh rằng bốn biến xAPI là các tổng hợp hành vi theo thuộc tính, không phải chuỗi sự kiện có timestamp. Do đó, thuật ngữ 'sequential branch' mô tả cách tensor được xử lý trong mô hình hơn là khẳng định dữ liệu xAPI là time series chuẩn."
    )

    add_heading(doc, "3.4. Nhánh Context MLP", 2)
    add_para(
        doc,
        "Nhánh bối cảnh loại các cột đã đưa vào nhánh tuần tự. Biến số được chuẩn hóa Min-Max. Biến phân loại được Label Encoding để tạo chỉ số, sau đó đi qua embedding riêng cho từng cột. Các embedding được ghép với biến số và đưa qua MLP hai tầng. Cách này tránh xem mã danh mục như một đại lượng liên tục."
    )

    add_heading(doc, "3.5. Hàm mất mát và ensemble", 2)
    add_para(
        doc,
        "Student-Mat và Student-Por dùng Weighted Cross-Entropy. xAPI được mô hình hóa theo thứ tự Low < Medium < High bằng hai ngưỡng và BCEWithLogitsLoss. Hai xác suất vượt ngưỡng được đổi thành ba xác suất lớp trước khi lấy argmax. Adam, ReduceLROnPlateau, early stopping và SWA được dùng trong quá trình huấn luyện. Kết quả cuối là trung bình xác suất của các mô hình theo seed; thư mục checkpoint hiện có 5 seed cho mỗi tập Student và 11 seed cho xAPI."
    )

    add_heading(doc, "3.6. Rule-based Learning Path Engine", 2)
    add_para(
        doc,
        "Bộ máy khuyến nghị không thay đổi dự đoán của mạng nơ-ron. Nó đọc feature gốc, lớp dự đoán và confidence để phát hiện RiskFactor có mã, tiêu đề, bằng chứng và mức ưu tiên. Với dữ liệu Student, luật xét số buổi vắng, tỷ lệ vắng/thời gian học, số lần trượt, G1-G2, studytime, Dalc+Walc và goout. Với xAPI, luật xét StudentAbsenceDays, VisITedResources, raisedhands, Discussion, AnnouncementsView, ParentAnsweringSurvey và ParentschoolSatisfaction."
    )
    add_figure(doc, "04_rule_engine.png", "Hình 3.3. Luồng xử lý của Rule-based Learning Path Engine", width_cm=15.0)
    add_table(
        doc,
        ["Nhóm rủi ro", "Ví dụ điều kiện", "Hành động điển hình", "Giai đoạn"],
        [
            ["Chuyên cần", "absences ≥ 10; hoặc Above-7", "Xác minh nguyên nhân, học bù, theo dõi chuyên cần", "Tuần 1"],
            ["Lỗ hổng kiến thức", "failures > 0; G2 < 10; G2 < G1", "Bài chẩn đoán, học lại chủ đề yếu, bài luyện có phản hồi", "Tuần 1-2"],
            ["Sử dụng học liệu", "VisITedResources < 40", "Truy cập LMS ≥ 4 ngày/tuần, hoàn thành tài nguyên trọng tâm", "Tuần 1-2"],
            ["Tương tác lớp", "raisedhands < 30 hoặc Discussion < 30", "Đặt câu hỏi, phản hồi, tham gia thảo luận", "Tuần 2-4"],
            ["Quản lý thời gian", "studytime thấp; goout cao", "Tăng giờ tự học có kế hoạch, điều chỉnh lịch", "Tuần 2-4"],
            ["Gia đình-nhà trường", "Khảo sát phụ huynh = No; hài lòng = Bad", "Báo cáo tiến độ và thống nhất mục tiêu tuần", "Trong 2 tuần"],
        ],
        [1600, 2300, 3472, 1700],
        "Bảng 3.2. Ví dụ ánh xạ luật sang learning path",
    )
    add_para(
        doc,
        "Về bản chất, learning path là khuyến nghị có thể kiểm toán: mỗi hành động liên kết với một tín hiệu đầu vào cụ thể. Tuy nhiên, các ngưỡng là heuristic miền ứng dụng; chưa có thử nghiệm ngẫu nhiên hoặc nghiên cứu người dùng chứng minh rằng lộ trình làm tăng điểm số. Vì vậy, hệ thống phải được triển khai như công cụ hỗ trợ cố vấn, không phải cơ chế tự động ra quyết định kỷ luật hoặc phân luồng có hệ quả cao."
    )

    chapter_break(doc)
    add_heading(doc, "CHƯƠNG 4. THIẾT LẬP THỰC NGHIỆM VÀ TIỀN XỬ LÝ DỮ LIỆU", 1)
    add_heading(doc, "4.1. Mô tả bộ dữ liệu", 2)
    add_table(
        doc,
        ["Bộ dữ liệu", "Số mẫu", "Số cột gốc", "Đích", "Chuỗi mô hình", "Nguồn thông tin chính"],
        [
            ["Student-Mat", "395", "33", "G3 → 3 lớp", "G1, G2", "Điểm, cá nhân, gia đình, trường học"],
            ["Student-Por", "649", "33", "G3 → 3 lớp", "G1, G2", "Điểm, cá nhân, gia đình, trường học"],
            ["xAPI-Edu-Data", "480", "17", "Class L/M/H", "4 chỉ báo hành vi", "Nhân khẩu, học phần, phụ huynh, tương tác LMS"],
        ],
        [1700, 1000, 1100, 1450, 1700, 2122],
        "Bảng 4.1. So sánh ba bộ dữ liệu",
    )
    add_para(
        doc,
        "Student-Mat và Student-Por có cùng schema 33 cột, khác môn học. G3 là điểm cuối kỳ trên thang 0-20. xAPI có 17 cột, trong đó Class là nhãn; các biến hành vi chính có thang đếm 0-100. Không có tuyên bố rằng ba tập đại diện cho cùng quần thể. Vì vậy, so sánh metric giữa tập chỉ mang tính mô tả năng lực của pipeline trên ba bối cảnh, không phải kiểm định trực tiếp bộ dữ liệu nào 'dễ' hay 'khó' nếu chưa kiểm soát phân bố và nhiễu nhãn."
    )
    add_figure(doc, "06_class_distribution.png", "Hình 4.1. Phân bố lớp trong ba tập locked test", width_cm=15.0)

    add_heading(doc, "4.2. Tách dữ liệu và chống rò rỉ", 2)
    add_para(
        doc,
        "Mỗi tập được tách stratified thành train pool 80% và locked test 20% bằng seed 42. Số mẫu locked test quan sát trong tệp dự đoán cuối là 79, 130 và 96. Locked test không tham gia Optuna, fit scaler/encoder, chọn đặc trưng hoặc chọn epoch. Trong từng fold Optuna, feature engineering được áp dụng riêng cho train và validation; DataPreprocessor chỉ fit trên train fold; SMOTE/ADASYN chỉ tác động train fold; FeatureSelector fit trên train fold rồi transform validation fold. Đây là cấu trúc đúng để giảm leakage ở CV."
    )
    add_figure(doc, "03_leakage_control.png", "Hình 4.2. Giao thức tách dữ liệu và kiểm soát rò rỉ", width_cm=15.0)
    add_para(
        doc,
        "Trong giai đoạn ensemble cuối, toàn bộ train pool được tiền xử lý và resampling trước khi tách validation nội bộ 15% theo seed. Locked test vẫn hoàn toàn độc lập, nên metric cuối không bị fit trực tiếp. Tuy nhiên, validation nội bộ sau resampling có thể chứa mẫu tổng hợp liên quan gần với mẫu train và làm lựa chọn epoch lạc quan hơn. Đây là hạn chế giao thức cần sửa trong phiên bản tiếp theo bằng cách tách validation trước resampling."
    )

    add_heading(doc, "4.3. Feature engineering", 2)
    add_para(doc, "Các đặc trưng dẫn xuất được tạo theo quy tắc xác định trước:")
    add_bullet(doc, "Student: grade_growth = G2−G1; grade_avg = (G1+G2)/2; absence_study_ratio; failure_risk; alcohol_risk; social_risk.")
    add_bullet(doc, "xAPI: engagement_score = tổng raisedhands, VisITedResources, AnnouncementsView, Discussion; absence_risk; parent_support_signal.")
    add_para(
        doc,
        "Một lưu ý khoa học quan trọng là grade_avg và grade_growth được suy ra trực tiếp từ G1-G2, vốn đồng thời được đưa vào nhánh tuần tự. Điều này làm tăng tính lặp thông tin giữa hai nhánh, có thể giúp tối ưu nhưng cũng làm khó tách riêng đóng góp của CNN-BiLSTM và MLP. Cần ablation để xác định lợi ích độc lập của từng nhánh."
    )

    add_heading(doc, "4.4. Chuẩn hóa, mã hóa và chọn đặc trưng", 2)
    add_para(
        doc,
        "Biến số được Min-Max scaling. Biến phân loại được Label Encoding; giá trị chưa thấy ở validation/test ánh xạ về 0. FeatureSelector dùng Pearson cho biến số và Chi-square cho biến phân loại với p-value < 0,1. Các biến chuỗi bắt buộc được giữ lại. Cách chọn đơn biến giúp giảm chiều nhưng không nắm bắt tương tác đa biến và không hiệu chỉnh multiple comparisons. Do đó, p-value chỉ là tiêu chí kỹ thuật chọn đặc trưng, không được diễn giải như bằng chứng nhân quả."
    )

    add_heading(doc, "4.5. SMOTE/ADASYN và Weighted Cross-Entropy", 2)
    add_para(
        doc,
        "Sampling strategy đặt số mẫu mục tiêu của mỗi lớp thiểu số bằng max(số hiện có, tỷ lệ × số lớp đa số). Số láng giềng được chặn bởi số mẫu nhỏ nhất trừ một để tránh lỗi ở fold nhỏ. Nếu resampling thất bại, pipeline ghi cảnh báo và quay về không oversampling. Song song, trọng số lớp được tính trên nhãn train gốc. Thiết kế này ưu tiên Recall lớp thiểu số, nhưng hiệu quả riêng của oversampling và class weighting chưa được cô lập."
    )
    add_placeholder(doc, "[REQUIRES EMPIRICAL DATA - DO NOT HALLUCINATE] Cần bảng ablation so sánh: không xử lý mất cân bằng; chỉ class weight; chỉ SMOTE; chỉ ADASYN; class weight + SMOTE; class weight + ADASYN.")

    add_heading(doc, "4.6. Optuna và cấu hình huấn luyện", 2)
    add_para(
        doc,
        "Student-Mat và Student-Por chạy 50 trial theo cấu hình mặc định. xAPI chạy đủ 150 trial; log ghi nhận trial 149 hoàn tất và trial tốt nhất là trial 55 với CV F1-Macro 0,812144. Mỗi trial dùng Stratified 5-fold CV. xAPI có không gian tìm kiếm rộng hơn, gồm kernel size 2/3/4, LSTM hidden đến 128, context/fusion hidden đến 256, dropout riêng theo nhánh, learning rate 5×10^-5 đến 5×10^-2 và weight decay 10^-8 đến 10^-2."
    )
    add_table(
        doc,
        ["Thành phần", "Student-Mat", "Student-Por", "xAPI"],
        [
            ["Số trial", "50", "50", "150"],
            ["CV", "Stratified 5-fold", "Stratified 5-fold", "Stratified 5-fold"],
            ["Epoch tối đa/trial", "50", "50", "80"],
            ["Early stopping patience", "15", "15", "25"],
            ["Sampler", "TPE", "TPE", "TPE đa biến"],
            ["Pruner", "MedianPruner", "MedianPruner", "MedianPruner"],
            ["Mục tiêu", "F1-Macro trung bình", "F1-Macro trung bình", "F1-Macro trung bình"],
        ],
        [2300, 2257, 2257, 2258],
        "Bảng 4.2. Thiết lập tối ưu siêu tham số",
    )
    add_table(
        doc,
        ["Siêu tham số", "Student-Mat", "Student-Por", "xAPI"],
        [
            ["Learning rate", "0,00363377", "0,00495091", "0,00589508"],
            ["Weight decay", "1,1783×10^-5", "5,1938×10^-5", "8,1862×10^-4"],
            ["Batch size", "16", "64", "64"],
            ["Resampling", "ADASYN", "ADASYN", "SMOTE"],
            ["Sampling ratio", "0,5152", "0,4678", "0,6139"],
            ["k-neighbors", "5", "5", "2"],
            ["CNN channels", "64", "64", "16"],
            ["CNN kernel", "3", "3", "3"],
            ["BiLSTM hidden", "64", "64", "128"],
            ["Context hidden", "32", "32", "256"],
            ["Fusion hidden", "64", "32", "256"],
            ["Dropout", "0,1626 chung", "0,1847 chung", "0,1461 / 0,1248 / 0,2315"],
        ],
        [2300, 2257, 2257, 2258],
        "Bảng 4.3. Siêu tham số tốt nhất được dùng trong báo cáo chính thức",
    )

    add_heading(doc, "4.7. Chỉ số đánh giá", 2)
    add_para(
        doc,
        "Accuracy là tỷ lệ dự đoán đúng. Precision_k = TP_k/(TP_k+FP_k), Recall_k = TP_k/(TP_k+FN_k), và F1_k là trung bình điều hòa của Precision_k và Recall_k. F1-Macro là trung bình không trọng số của F1 trên ba lớp, nên mỗi lớp có đóng góp ngang nhau bất kể support. RMSE được tính trên mã lớp thứ tự 0-1-2; R² so sánh sai số bình phương với phương sai nhãn. Vì phân loại là nhiệm vụ chính, F1-Macro là metric quyết định; RMSE và R² chỉ được dùng như chỉ số phụ phản ánh độ xa thứ tự giữa nhãn dự đoán và nhãn thật."
    )

    chapter_break(doc)
    add_heading(doc, "CHƯƠNG 5. ĐÁNH GIÁ MÔ HÌNH TRÊN CƠ SỞ KHOA HỌC", 1)
    add_heading(doc, "5.1. Nguyên tắc đánh giá", 2)
    add_para(
        doc,
        "Chương này chỉ sử dụng dữ liệu thực nghiệm có trong các tệp locked-test metrics, báo cáo classification và predictions của pipeline chính thức. Mọi kết luận về overfitting, độ ổn định, calibration, ý nghĩa thống kê hoặc hiệu quả can thiệp vượt quá dữ liệu hiện có đều được giới hạn hoặc thay bằng placeholder. Không sử dụng kết quả v27 vì yêu cầu phân tích xAPI CV xấp xỉ 0,81 và test xấp xỉ 0,77 tương ứng với lần chạy 150-trial và bộ báo cáo chính thức không hậu tố v27."
    )

    add_heading(doc, "5.2. Kết quả locked test tổng hợp", 2)
    add_table(
        doc,
        ["Dataset", "Accuracy", "Precision-Macro", "Recall-Macro", "F1-Macro", "RMSE", "R²"],
        [
            ["Student-Mat", "0,8861", "0,8781", "0,9170", "0,8905", "0,3375", "0,7720"],
            ["Student-Por", "0,8615", "0,8110", "0,8816", "0,8394", "0,3721", "0,6063"],
            ["xAPI", "0,7604", "0,7661", "0,7711", "0,7663", "0,4895", "0,5737"],
        ],
        [1500, 1100, 1400, 1300, 1200, 1100, 1472],
        "Bảng 5.1. Metric locked-test cuối cùng",
    )
    add_para(
        doc,
        "Student-Mat đạt kết quả tốt nhất trên toàn bộ metric: F1-Macro 0,8905, Recall-Macro 0,9170, RMSE thấp nhất 0,3375 và R² cao nhất 0,7720. Student-Por đạt F1-Macro 0,8394 dù phân bố lớp test lệch mạnh về Medium. xAPI đạt F1-Macro 0,7663 và RMSE 0,4895, cho thấy sai số phân loại và độ lệch thứ tự lớn hơn. Precision-Macro và Recall-Macro của xAPI gần nhau, gợi ý không có sự đánh đổi cực đoan giữa hai đại lượng ở mức macro, nhưng phân tích theo lớp vẫn cần thiết."
    )
    add_figure(doc, "07_locked_metrics.png", "Hình 5.1. So sánh các metric phân loại trên locked test", width_cm=15.0)

    add_heading(doc, "5.3. CV so với locked test", 2)
    add_table(
        doc,
        ["Dataset", "Optuna best CV F1", "Locked-test F1", "Chênh lệch CV-Test", "Tỷ lệ giữ lại"],
        [
            ["Student-Mat", "0,9112", "0,8905", "0,0207", "97,73%"],
            ["Student-Por", "0,8898", "0,8394", "0,0504", "94,33%"],
            ["xAPI", "0,8121", "0,7663", "0,0458", "94,36%"],
        ],
        [1700, 1700, 1700, 2000, 1972],
        "Bảng 5.2. So sánh F1-Macro CV và locked test",
    )
    add_para(
        doc,
        "Cả ba locked-test score đều thấp hơn best CV score, phù hợp với kỳ vọng rằng điểm CV tốt nhất sau nhiều trial có thiên lệch lựa chọn. Student-Mat có khoảng cách nhỏ nhất 0,0207 và giữ lại 97,73% điểm CV, là bằng chứng mạnh nhất trong ba tập cho khả năng khái quát hóa tương đối ổn định. Student-Por có khoảng cách lớn nhất 0,0504, nhỉnh hơn xAPI 0,0458; do đó không chính xác nếu khẳng định xAPI có generalization gap lớn nhất. Tuy nhiên, xAPI có mức F1 tuyệt đối thấp nhất, vì vậy giới hạn chính của xAPI nằm ở chất lượng dự đoán tổng thể và biểu diễn hơn là chỉ ở độ lớn gap."
    )
    add_para(
        doc,
        "Các khoảng cách 2,1-5,0 điểm phần trăm không cho thấy sự sụp đổ hiệu năng trên locked test. Điều này hỗ trợ nhận định rằng mô hình không bị overfit nghiêm trọng đến mức mất khả năng tổng quát hóa. Tuy nhiên, không thể 'chứng minh không overfit' chỉ từ một phép chia locked test, bởi best CV là cực đại qua nhiều trial và không có phân phối lặp của score. Kết luận khoa học phù hợp là: dữ liệu hiện có cung cấp bằng chứng về robustness ở mức thực nghiệm, nhưng chưa đủ để loại trừ hoàn toàn overfitting hoặc selection bias."
    )
    add_figure(doc, "08_cv_test_gap.png", "Hình 5.2. So sánh F1-Macro giữa cross-validation và locked test", width_cm=15.0)
    add_placeholder(doc, "[REQUIRES EMPIRICAL DATA - DO NOT HALLUCINATE] Cần mean ± standard deviation của 5 fold cho cấu hình tốt nhất, kết quả repeated stratified CV hoặc bootstrap confidence interval của locked-test F1 để định lượng bất định.")
    add_placeholder(doc, "[REQUIRES EMPIRICAL DATA - DO NOT HALLUCINATE] Cần train/validation loss và F1 theo epoch của từng seed để đánh giá trực tiếp divergence, thời điểm early stopping và overfitting động.")

    add_heading(doc, "5.4. Phân tích theo lớp và ma trận nhầm lẫn", 2)
    add_table(
        doc,
        ["Dataset", "Lớp", "Precision", "Recall", "F1", "Support"],
        [
            ["Student-Mat", "Low", "0,83", "0,96", "0,89", "26"],
            ["Student-Mat", "Medium", "0,97", "0,79", "0,87", "38"],
            ["Student-Mat", "High", "0,83", "1,00", "0,91", "15"],
            ["Student-Por", "Low", "0,65", "0,85", "0,74", "20"],
            ["Student-Por", "Medium", "0,95", "0,83", "0,89", "84"],
            ["Student-Por", "High", "0,83", "0,96", "0,89", "26"],
            ["xAPI", "Low", "0,77", "0,88", "0,82", "26"],
            ["xAPI", "Medium", "0,73", "0,71", "0,72", "42"],
            ["xAPI", "High", "0,80", "0,71", "0,75", "28"],
        ],
        [1500, 1200, 1400, 1400, 1300, 2272],
        "Bảng 5.3. Metric theo lớp từ classification report",
    )
    add_para(
        doc,
        "Student-Mat nhận diện toàn bộ 15 mẫu High và 25/26 mẫu Low, nhưng 5 mẫu Medium bị dự đoán thành Low và 3 mẫu Medium thành High. Điều này giải thích Recall Medium 0,79 thấp hơn hai lớp biên. Student-Por có 84 mẫu Medium, chiếm 64,6% locked test; mô hình nhận đúng 70 mẫu. Recall Low 0,85 nhưng Precision Low chỉ 0,65 do 9 mẫu Medium bị hạ xuống Low. Macro-F1 vì vậy thấp hơn Accuracy khoảng 2,2 điểm phần trăm, cho thấy Accuracy được hưởng lợi từ lớp Medium đông."
    )
    add_para(
        doc,
        "xAPI cho thấy cấu trúc lỗi rõ ở lớp lân cận: 7 Medium bị dự đoán Low, 5 Medium bị dự đoán High và 8 High bị dự đoán Medium. Không có lỗi nhảy trực tiếp giữa Low và High trong locked test, tức mô hình chủ yếu nhầm giữa các mức kề nhau. Điều này phù hợp với việc RMSE trên mã lớp không quá lớn như trường hợp nhảy hai bậc, nhưng Recall Medium và High đều chỉ 0,71."
    )
    add_table(
        doc,
        ["Thật / Dự đoán", "Low", "Medium", "High"],
        [["Low", "25", "1", "0"], ["Medium", "5", "30", "3"], ["High", "0", "0", "15"]],
        [3000, 2024, 2024, 2024],
        "Bảng 5.4. Ma trận nhầm lẫn Student-Mat",
    )
    add_table(
        doc,
        ["Thật / Dự đoán", "Low", "Medium", "High"],
        [["Low", "17", "3", "0"], ["Medium", "9", "70", "5"], ["High", "0", "1", "25"]],
        [3000, 2024, 2024, 2024],
        "Bảng 5.5. Ma trận nhầm lẫn Student-Por",
    )
    add_table(
        doc,
        ["Thật / Dự đoán", "Low", "Medium", "High"],
        [["Low", "23", "3", "0"], ["Medium", "7", "30", "5"], ["High", "0", "8", "20"]],
        [3000, 2024, 2024, 2024],
        "Bảng 5.6. Ma trận nhầm lẫn xAPI",
    )
    add_figure(doc, "09_confusion_matrices.png", "Hình 5.3. Ma trận nhầm lẫn của ba bộ dữ liệu", width_cm=15.0)
    add_figure(doc, "10_per_class_f1.png", "Hình 5.4. F1-score theo từng lớp Low, Medium và High", width_cm=15.0)

    add_heading(doc, "5.5. Đánh giá xử lý mất cân bằng", 2)
    add_para(
        doc,
        "Phân bố locked test lần lượt là Student-Mat 26/38/15, Student-Por 20/84/26 và xAPI 26/42/28. Student-Por mất cân bằng mạnh nhất. Việc Recall Low = 0,85 và Recall High = 0,96 trên Student-Por cho thấy mô hình không bỏ qua hai lớp ít mẫu. F1-Macro 0,8394, thấp hơn weighted-F1 0,86 trong classification report, phản ánh đúng mức khó ở Low. Kết quả hỗ trợ hiệu quả thực dụng của class weighting và ADASYN, nhưng không đủ để quy kết nhân quả cho từng thành phần vì thiếu ablation."
    )
    add_para(
        doc,
        "Student-Mat dùng ADASYN và đạt Recall-Macro cao 0,9170; xAPI dùng SMOTE và đạt các macro metric quanh 0,77. Không thể kết luận ADASYN tốt hơn SMOTE bằng so sánh chéo dataset, vì dữ liệu, không gian siêu tham số, cỡ mẫu và phân bố khác nhau. Cần thí nghiệm trong cùng dataset với cùng seed và cấu hình để so sánh công bằng."
    )

    add_heading(doc, "5.6. Giới hạn năng lực và biểu diễn trên xAPI", 2)
    add_para(
        doc,
        "xAPI đạt best CV F1 0,8121 nhưng locked-test F1 0,7663. Mặc dù không phải gap lớn nhất, xAPI có F1 thấp nhất, Accuracy thấp nhất và RMSE cao nhất. Optuna đã chọn BiLSTM hidden 128, Context MLP 256 và Fusion 256, lớn hơn cấu hình hai tập Student, trong khi CNN chỉ có 16 channels. Việc tăng chiều context và fusion nhưng hiệu năng vẫn dừng quanh 0,77 gợi ý rằng giới hạn không chỉ do thiếu số tham số."
    )
    add_para(
        doc,
        "Ba nguyên nhân hợp lý có thể được rút ra từ cấu trúc dữ liệu và mã nguồn. Thứ nhất, bốn chỉ báo hành vi là thuộc tính tổng hợp không có timestamp; áp đặt thứ tự raisedhands → resources → announcements → discussion không tương đương chuỗi thời gian. Thứ hai, biến phân loại được Label Encoding và chuẩn hóa như số liên tục, làm xuất hiện quan hệ thứ tự giả. Thứ ba, SMOTE nội suy trên toàn bộ vector đã mã hóa, sau đó categorical values được ép về số nguyên, có thể tạo mẫu danh mục không tự nhiên. Các yếu tố này là giới hạn của biểu diễn/tiền xử lý, không đơn thuần là 'capacity' theo nghĩa số lượng tham số."
    )
    add_para(
        doc,
        "Ma trận nhầm lẫn xAPI cho thấy ranh giới Medium-High là điểm yếu chính, với 8/28 mẫu High bị hạ xuống Medium và 5/42 mẫu Medium bị nâng lên High. Khi nhãn được hình thành từ mức điểm rộng, các nhóm gần biên có thể chồng lấn về hành vi. Nếu không có điểm số liên tục hoặc timestamp tương tác, mô hình khó học ranh giới tinh hơn."
    )
    add_figure(doc, "11_feature_importance.png", "Hình 5.5. Permutation feature importance trên ba bộ dữ liệu", width_cm=15.0)
    add_placeholder(doc, "[REQUIRES EMPIRICAL DATA - DO NOT HALLUCINATE] Cần ablation xAPI: Context MLP-only; CNN-BiLSTM-only; hybrid; thay Label Encoding bằng embedding; thay SMOTE bằng SMOTENC; dùng interaction logs theo thời gian nếu có.")

    add_heading(doc, "5.7. Diễn giải RMSE và R²", 2)
    add_para(
        doc,
        "R² dương trên cả ba tập cho thấy dự đoán mã lớp giảm sai số bình phương so với baseline dự đoán giá trị trung bình của nhãn mã hóa. Student-Mat có R² 0,7720; Student-Por 0,6063; xAPI 0,5737. Tuy nhiên, R² và RMSE phụ thuộc giả định khoảng cách Low-Medium bằng Medium-High. Đây là quy ước ordinal hợp lý nhưng không phải thang đo khoảng được xác nhận. Vì vậy, hai metric không thay thế F1-Macro và không nên được dùng một mình để tuyên bố chất lượng phân loại."
    )

    add_heading(doc, "5.8. Những đánh giá chưa thể kết luận", 2)
    add_figure(doc, "12_confidence_distribution.png", "Hình 5.6. Phân bố độ tin cậy của dự đoán ensemble", width_cm=15.0)
    add_placeholder(doc, "[REQUIRES EMPIRICAL DATA - DO NOT HALLUCINATE] Không có confidence interval hoặc kiểm định paired bootstrap để kết luận khác biệt giữa ba dataset/model có ý nghĩa thống kê.")
    add_placeholder(doc, "[REQUIRES EMPIRICAL DATA - DO NOT HALLUCINATE] Không có Brier score, Expected Calibration Error hoặc reliability diagram để đánh giá xác suất/confidence có được hiệu chỉnh hay không.")
    add_placeholder(doc, "[REQUIRES EMPIRICAL DATA - DO NOT HALLUCINATE] Không có metric fairness theo giới tính, trường, quốc tịch hoặc nhóm xã hội; không được tuyên bố mô hình công bằng.")
    add_placeholder(doc, "[REQUIRES EMPIRICAL DATA - DO NOT HALLUCINATE] Không có external validation trên cơ sở giáo dục khác; robustness hiện chỉ được chứng minh trong ba phép chia nội bộ của ba dataset công khai.")
    add_placeholder(doc, "[REQUIRES EMPIRICAL DATA - DO NOT HALLUCINATE] Không có thử nghiệm người dùng hoặc A/B test để đánh giá learning path cải thiện kết quả học tập.")

    add_heading(doc, "5.9. Kết luận khoa học của chương", 2)
    add_para(
        doc,
        "Bằng chứng locked test cho thấy mô hình hoạt động tốt trên Student-Mat, khá trên Student-Por và trung bình-khá trên xAPI theo F1-Macro. Gap CV-test không lớn đến mức cho thấy sụp đổ tổng quát hóa, nhưng Student-Por và xAPI giảm khoảng 4,6-5,0 điểm phần trăm, cần được coi là selection/generalization gap có ý nghĩa thực hành. Mất cân bằng được xử lý đủ để duy trì Recall cho lớp ít mẫu, đặc biệt trên Student-Por, song chưa thể tách đóng góp của ADASYN/SMOTE và class weight. xAPI bộc lộ giới hạn biểu diễn rõ nhất; cải tiến nên tập trung vào dữ liệu tuần tự thực, encoding danh mục và resampling hỗn hợp thay vì chỉ tăng kích thước mạng."
    )

    chapter_break(doc)
    add_heading(doc, "CHƯƠNG 6. THẢO LUẬN VÀ TRIỂN KHAI THỰC TIỄN", 1)
    add_heading(doc, "6.1. Ý nghĩa thực tiễn của dự đoán", 2)
    add_para(
        doc,
        "Một hệ thống dự đoán giáo dục hữu ích phải hỗ trợ ưu tiên nguồn lực chứ không thay thế đánh giá chuyên môn. Người học dự đoán Low hoặc có risk factor ưu tiên 1 có thể được đưa vào danh sách cố vấn xem xét. Người học Medium có confidence thấp cần được xem là trường hợp không chắc chắn, không phải nhãn cố định. Người học High vẫn có thể nhận khuyến nghị nếu xuất hiện rủi ro chuyên cần hoặc suy giảm G2 so với G1."
    )

    add_heading(doc, "6.2. Vận hành Learning Path", 2)
    add_number(doc, "Mô hình tạo nhãn, xác suất và confidence cho từng người học.")
    add_number(doc, "Rule engine đọc feature snapshot và tạo danh sách risk factor có bằng chứng.")
    add_number(doc, "Risk band được xác định từ lớp dự đoán và mức ưu tiên rủi ro.")
    add_number(doc, "Các hành động được sắp theo Tuần 1, Tuần 1-2, Tuần 2-4 và kiểm tra cuối tuần.")
    add_number(doc, "Cố vấn xác nhận tính phù hợp, điều chỉnh mục tiêu và ghi nhận việc thực hiện.")
    add_number(doc, "Dữ liệu theo dõi mới được đưa vào lần đánh giá sau để so sánh xu hướng.")
    add_para(
        doc,
        "Trong vận hành, mỗi khuyến nghị cần có trạng thái proposed, accepted, modified, completed hoặc rejected; người thực hiện; ngày bắt đầu; ngày đánh giá; và outcome. Schema hiện tại lưu recommended_learning_path dưới dạng JSONB nhưng chưa có bảng sự kiện thực thi chi tiết. Đây là hướng mở rộng cần thiết để biến hệ thống từ máy sinh khuyến nghị thành nền tảng theo dõi can thiệp."
    )
    add_figure(doc, "13_risk_band_distribution.png", "Hình 6.1. Phân bố risk band do Learning Path Engine tạo ra", width_cm=15.0)

    add_heading(doc, "6.3. Tích hợp PostgreSQL", 2)
    add_table(
        doc,
        ["Bảng", "Mục đích", "Trường quan trọng"],
        [
            ["paper_runs", "Quản lý phiên chạy", "generated_at, status, run_payload"],
            ["paper_predictions", "Lưu dự đoán locked test hoặc production", "true/pred label, probability, confidence, original_features, seed"],
            ["paper_evaluation_metrics", "Lưu metric theo dataset và protocol", "accuracy, precision_macro, recall_macro, f1_macro, rmse, r2"],
            ["paper_learning_recommendations", "Lưu risk và learning path", "risk_band, feature_snapshot, recommended_learning_path"],
            ["students / student_grades", "Liên kết hồ sơ và điểm nguồn", "dataset_name, source_row_index, G1, G2, G3, target_class"],
        ],
        [2200, 2800, 4072],
        "Bảng 6.1. Vai trò các bảng PostgreSQL",
    )
    add_para(
        doc,
        "Việc lưu original_features và probability cho phép kiểm toán sau này: có thể tái tạo đầu vào, kiểm tra drift, xem confidence và đối chiếu learning path. Giao dịch ghi dữ liệu cần bảo đảm atomicity để tránh tồn tại prediction mà thiếu metric hoặc recommendation. Trong môi trường thật, cần bổ sung model_version, preprocessing_version, feature_schema_version, consent status, retention policy và audit log."
    )
    add_figure(doc, "05_postgresql_schema.png", "Hình 6.2. Quan hệ giữa các bảng lưu trữ PostgreSQL", width_cm=15.0)

    add_heading(doc, "6.4. Kiến trúc triển khai đề xuất", 2)
    add_bullet(doc, "Batch scoring định kỳ cho Student-Mat/Por khi có điểm G1-G2 mới.")
    add_bullet(doc, "API scoring cho LMS khi chỉ báo xAPI được cập nhật.")
    add_bullet(doc, "Model registry lưu trọng số, tham số, code commit và metric phê duyệt.")
    add_bullet(doc, "Dashboard cố vấn hiển thị xác suất, risk factor và lộ trình; không chỉ hiển thị nhãn.")
    add_bullet(doc, "Monitoring theo dõi phân bố feature, tỷ lệ lớp dự đoán, calibration và metric khi có nhãn thật.")
    add_bullet(doc, "Cơ chế human-in-the-loop bắt buộc đối với mọi quyết định có ảnh hưởng đến người học.")

    add_heading(doc, "6.5. Rủi ro đạo đức và quản trị", 2)
    add_para(
        doc,
        "Dữ liệu học tập chứa thông tin nhạy cảm về kết quả, hành vi và bối cảnh gia đình. Hệ thống cần tuân thủ nguyên tắc tối thiểu hóa dữ liệu, phân quyền theo vai trò, mã hóa khi lưu và truyền, giới hạn thời gian lưu, và ghi log truy cập. Người học cần được thông báo mục đích sử dụng dữ liệu và có cơ chế phản hồi khi dự đoán sai. Không nên dùng nhãn Low để từ chối cơ hội học tập hoặc tạo kỳ vọng tiêu cực."
    )
    add_para(
        doc,
        "Các thuộc tính như giới tính, quốc tịch, trường hoặc quan hệ phụ huynh có thể trở thành proxy cho bất bình đẳng xã hội. Do chưa có fairness audit, triển khai thực tế phải tạm giới hạn vai trò của mô hình ở sàng lọc hỗ trợ, đồng thời đo performance theo nhóm trước khi mở rộng. Rule engine cũng cần được hội đồng chuyên môn phê duyệt ngưỡng và nội dung khuyến nghị."
    )

    add_heading(doc, "6.6. Hạn chế của nghiên cứu", 2)
    limitations = [
        "Cỡ mẫu nhỏ làm metric nhạy với cách chia tập và seed.",
        "Chuỗi đầu vào quá ngắn; xAPI không có nhật ký thời gian thực trong dữ liệu hiện dùng.",
        "Label Encoding và chuẩn hóa danh mục tạo thứ tự giả; SMOTE/ADASYN trên mã danh mục có thể tạo mẫu phi tự nhiên.",
        "Validation nội bộ ensemble được tách sau resampling, có nguy cơ lạc quan khi chọn epoch.",
        "Thiếu ablation định lượng đóng góp của CNN, BiLSTM, attention, MLP, feature engineering và xử lý mất cân bằng.",
        "Thiếu khoảng tin cậy, calibration, fairness, external validation và kiểm định drift.",
        "Learning path dựa trên luật chuyên gia, chưa có bằng chứng nhân quả hoặc đánh giá người dùng.",
        "RMSE/R² trên mã lớp chỉ có ý nghĩa phụ thuộc giả định ordinal.",
    ]
    for item in limitations:
        add_bullet(doc, item)

    chapter_break(doc)
    add_heading(doc, "CHƯƠNG 7. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", 1)
    add_heading(doc, "7.1. Kết luận", 2)
    add_para(
        doc,
        "Luận văn đã xây dựng một pipeline hoàn chỉnh cho dự đoán và khuyến nghị thành tích học tập dựa trên kiến trúc lai CNN-BiLSTM + Context MLP. Thiết kế hai nhánh cho phép kết hợp chuỗi điểm/chỉ báo hành vi với thông tin bối cảnh; attention pooling tạo biểu diễn chuỗi cố định; fusion dense tạo dự đoán ba lớp. Quy trình thực nghiệm tách locked test trước tối ưu, sử dụng fold-wise preprocessing trong Optuna, xử lý mất cân bằng và ensemble nhiều seed."
    )
    add_para(
        doc,
        "Trên locked test, F1-Macro đạt 0,8905 cho Student-Mat, 0,8394 cho Student-Por và 0,7663 cho xAPI. Kết quả cho thấy khả năng khái quát hóa tốt nhất trên Student-Mat. Student-Por duy trì Recall cao ở lớp Low và High dù lớp Medium chiếm đa số. xAPI là trường hợp hạn chế nhất về hiệu năng tuyệt đối; phân tích cho thấy nguyên nhân có khả năng nằm ở bản chất chỉ báo tổng hợp, encoding danh mục và resampling hỗn hợp, không chỉ ở kích thước mạng."
    )
    add_para(
        doc,
        "Rule-based Learning Path Engine chuyển risk factor thành kế hoạch theo tuần và PostgreSQL cung cấp khả năng lưu vết. Hai thành phần này làm hệ thống gần với triển khai thực tế hơn một mô hình phân loại đơn lẻ. Tuy vậy, khuyến nghị chưa được kiểm chứng về hiệu quả can thiệp và mô hình chưa được đánh giá fairness, calibration hoặc external validity."
    )

    add_heading(doc, "7.2. Hướng phát triển", 2)
    future = [
        "Tách validation trước resampling trong giai đoạn ensemble và chạy nested CV hoặc repeated CV.",
        "Báo cáo bootstrap confidence interval, fold standard deviation và kiểm định paired khi so sánh cấu hình.",
        "Thay Label Encoding liên tục bằng embedding hoặc one-hot phù hợp; dùng SMOTENC cho dữ liệu hỗn hợp.",
        "Thu thập event log có timestamp để hình thành chuỗi xAPI thực, kèm masking cho chuỗi độ dài khác nhau.",
        "Thực hiện ablation MLP-only, sequence-only, hybrid, không attention, không feature engineering và các chiến lược imbalance.",
        "Đánh giá calibration bằng Brier score/ECE và áp dụng temperature scaling nếu cần.",
        "Thực hiện fairness audit theo nhóm với tiêu chí phù hợp và quy trình phê duyệt đạo đức.",
        "Thiết kế nghiên cứu can thiệp có đối chứng hoặc stepped-wedge để đo hiệu quả learning path.",
        "Bổ sung drift monitoring, model registry, versioning và quy trình rollback trong PostgreSQL/MLOps.",
    ]
    for item in future:
        add_number(doc, item)

    add_heading(doc, "7.3. Tuyên bố cuối cùng", 2)
    add_para(
        doc,
        "Giá trị chính của nghiên cứu không chỉ là đạt metric cao, mà là xây dựng một quy trình có thể kiểm tra, thừa nhận bất định và tách bạch rõ dự đoán với khuyến nghị. Kết quả hiện tại đủ để chứng minh tính khả thi kỹ thuật của hệ thống trên ba bộ dữ liệu công khai, nhưng chưa đủ để coi mô hình là công cụ tự động ra quyết định giáo dục. Bước tiếp theo phải ưu tiên validation nghiêm ngặt, biểu diễn dữ liệu phù hợp và đánh giá tác động thực tế."
    )

    chapter_break(doc)
    add_heading(doc, "TÀI LIỆU THAM KHẢO", 1)
    references = [
        "[1] Cortez, P., & Silva, A. M. G. (2008). Using Data Mining to Predict Secondary School Student Performance. UCI Machine Learning Repository, Student Performance dataset. https://archive.ics.uci.edu/dataset/320/student+performance",
        "[2] Amrieh, E. A., Hamtini, T., & Aljarah, I. (2016). Mining Educational Data to Predict Student's Academic Performance Using Ensemble Methods. International Journal of Database Theory and Application, 9(8), 119-136. https://doi.org/10.14257/ijdta.2016.9.8.13",
        "[3] Hochreiter, S., & Schmidhuber, J. (1997). Long Short-Term Memory. Neural Computation, 9(8), 1735-1780. https://doi.org/10.1162/neco.1997.9.8.1735",
        "[4] Schuster, M., & Paliwal, K. K. (1997). Bidirectional Recurrent Neural Networks. IEEE Transactions on Signal Processing, 45(11), 2673-2681. https://doi.org/10.1109/78.650093",
        "[5] Chawla, N. V., Bowyer, K. W., Hall, L. O., & Kegelmeyer, W. P. (2002). SMOTE: Synthetic Minority Over-sampling Technique. Journal of Artificial Intelligence Research, 16, 321-357. https://doi.org/10.1613/jair.953",
        "[6] He, H., Bai, Y., Garcia, E. A., & Li, S. (2008). ADASYN: Adaptive Synthetic Sampling Approach for Imbalanced Learning. IEEE International Joint Conference on Neural Networks, 1322-1328. https://doi.org/10.1109/IJCNN.2008.4633969",
        "[7] Akiba, T., Sano, S., Yanase, T., Ohta, T., & Koyama, M. (2019). Optuna: A Next-generation Hyperparameter Optimization Framework. Proceedings of KDD 2019. https://doi.org/10.1145/3292500.3330701",
        "[8] LeCun, Y., Bottou, L., Bengio, Y., & Haffner, P. (1998). Gradient-Based Learning Applied to Document Recognition. Proceedings of the IEEE, 86(11), 2278-2324. https://doi.org/10.1109/5.726791",
        "[9] Mã nguồn dự án: src/models.py, src/data_pipeline.py, src/train_pipeline.py, src/explainability.py, src/evaluation.py và scripts/run_pipeline.py.",
        "[10] Kết quả thực nghiệm dự án: reports/final/metrics/*_locked_test_metrics.json; reports/final/*_final_report.txt; reports/final/predictions/*_predictions.csv; logs/xapi_150_stderr.log.",
    ]
    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(-0.75)
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(ref)
        set_run_font(r, size=11.5)

    chapter_break(doc)
    add_heading(doc, "PHỤ LỤC A. NGUỒN SỐ LIỆU THỰC NGHIỆM", 1)
    add_table(
        doc,
        ["Nội dung", "Tệp nguồn"],
        [
            ["Metric Student-Mat", "reports/final/metrics/student-mat_3class_locked_test_metrics.json"],
            ["Metric Student-Por", "reports/final/metrics/student-por_3class_locked_test_metrics.json"],
            ["Metric xAPI", "reports/final/metrics/xapi_3class_locked_test_metrics.json"],
            ["CV và siêu tham số", "reports/final/*_3class_final_report.txt"],
            ["Ma trận nhầm lẫn", "reports/final/predictions/*_3class_predictions.csv"],
            ["Xác nhận 150 trial xAPI", "logs/xapi_150_stderr.log"],
            ["Kiến trúc", "src/models.py"],
            ["Tiền xử lý", "src/data_pipeline.py"],
            ["Khuyến nghị", "src/explainability.py"],
            ["PostgreSQL", "database/schema.sql"],
        ],
        [3000, 6072],
        "Bảng A.1. Truy vết nguồn dữ liệu dùng trong luận văn",
    )

    add_heading(doc, "PHỤ LỤC B. CHECKLIST TÁI LẬP", 1)
    checklist = [
        "Cố định seed và lưu phiên bản thư viện/môi trường.",
        "Tạo locked split một lần; không dùng locked test trong Optuna hoặc chọn epoch.",
        "Fit feature engineering state, scaler, encoder và selector chỉ trên train tương ứng.",
        "Resampling chỉ trên train; ghi lại method, ratio và k-neighbors.",
        "Lưu toàn bộ best params, fold score, seed score và checkpoint.",
        "Tính metric từ prediction CSV và đối chiếu classification report.",
        "Ghi model version, preprocessing version và schema version khi đưa vào PostgreSQL.",
        "Không tuyên bố hiệu quả khuyến nghị nếu chưa có dữ liệu can thiệp.",
    ]
    for item in checklist:
        add_bullet(doc, item)

    add_heading(doc, "PHỤ LỤC C. CÁC PHÂN TÍCH CẦN BỔ SUNG KHI CÓ DỮ LIỆU", 1)
    add_placeholder(doc, "[REQUIRES EMPIRICAL DATA - DO NOT HALLUCINATE] Cần bổ sung learning curves, biểu đồ calibration, khoảng tin cậy, đánh giá fairness và biểu đồ ablation khi có dữ liệu thực nghiệm tương ứng.")

    finalize_sections(doc)
    return doc


if __name__ == "__main__":
    OUT.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUT)
    print(OUT)
