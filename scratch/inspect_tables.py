import docx

doc = docx.Document("Bao_cao_cuoi_cung.docx")
print(f"Total tables: {len(doc.tables)}")
for i, table in enumerate(doc.tables):
    print(f"\n--- Table {i} ---")
    for r_idx, row in enumerate(table.rows):
        # Avoid printing massive rows if they exist
        try:
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            # Deduplicate adjacent identical cells due to merged cells
            dedup_cells = []
            for cell in cells:
                if not dedup_cells or cell != dedup_cells[-1]:
                    dedup_cells.append(cell)
            print(f"Row {r_idx}: {' | '.join(dedup_cells[:8])}")
        except Exception as e:
            print(f"Row {r_idx} Error: {e}")
