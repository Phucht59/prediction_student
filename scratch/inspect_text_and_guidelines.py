import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc = docx.Document("scratch/HUFLIT_Mau_Hinh_Thuc_KLTN.docx")
print(f"Total paragraphs in template: {len(doc.paragraphs)}")

# Print paragraphs that contain formatting keywords
keywords = ["lề", "giãn", "cách", "font", "chữ", "size", "bìa", "bảng", "hình"]
for idx, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    lower_text = text.lower()
    if any(kw in lower_text for kw in keywords):
        print(f"[{idx}]: {text}")
