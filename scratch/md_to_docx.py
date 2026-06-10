import docx
from docx.shared import Pt

def convert_md_to_docx(md_path, docx_path):
    doc = docx.Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(16)
    h1.font.bold = True
    
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(14)
    h2.font.bold = True

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            doc.add_paragraph(line, style='List Bullet')
        elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
            doc.add_paragraph(line, style='List Number')
        else:
            doc.add_paragraph(line)
            
    doc.save(docx_path)
    print(f"Saved to {docx_path}")

if __name__ == '__main__':
    convert_md_to_docx('C:/Huflit/kltn/reports/KhoaLuan_ChinhThuc.md', 'C:/Huflit/kltn/reports/KhoaLuan_ChinhThuc.docx')
