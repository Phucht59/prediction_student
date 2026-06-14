import json
from pathlib import Path
import docx
from docx.shared import Cm, Pt

# Helper to check margin equality in centimeters
def close_to_cm(val_emu, expected_cm):
    if val_emu is None:
        return False
    val_cm = val_emu.cm
    return abs(val_cm - expected_cm) < 0.01

def main():
    report = {
        "document_path": "Bao_cao_cuoi_cung.docx",
        "page_setup": {},
        "table_0_fidelity": [],
        "table_1_llm_judge": [],
        "errors": []
    }
    
    # 1. Load document
    try:
        doc = docx.Document("Bao_cao_cuoi_cung.docx")
    except Exception as e:
        report["errors"].append(f"Failed to load docx: {e}")
        print(json.dumps(report, indent=2, ensure_ascii=False))
        return

    # 2. Verify Page Setup
    section = doc.sections[0]
    
    # A4 dimensions: 21cm x 29.7cm
    w_ok = close_to_cm(section.page_width, 21.0)
    h_ok = close_to_cm(section.page_height, 29.7)
    left_ok = close_to_cm(section.left_margin, 3.5)
    right_ok = close_to_cm(section.right_margin, 2.0)
    top_ok = close_to_cm(section.top_margin, 3.0)
    bottom_ok = close_to_cm(section.bottom_margin, 3.0)
    
    report["page_setup"] = {
        "page_width_cm": section.page_width.cm if section.page_width else None,
        "page_height_cm": section.page_height.cm if section.page_height else None,
        "left_margin_cm": section.left_margin.cm if section.left_margin else None,
        "right_margin_cm": section.right_margin.cm if section.right_margin else None,
        "top_margin_cm": section.top_margin.cm if section.top_margin else None,
        "bottom_margin_cm": section.bottom_margin.cm if section.bottom_margin else None,
        "a4_dimensions_match": w_ok and h_ok,
        "margins_match": left_ok and right_ok and top_ok and bottom_ok
    }

    # 3. Load JSON sources
    recommendations_dir = Path("reports/final/recommendations")
    eval_files = {
        "student-mat": recommendations_dir / "student_mat_evaluation.json",
        "student-por": recommendations_dir / "student_por_evaluation.json",
        "xapi": recommendations_dir / "xapi_evaluation.json"
    }
    
    eval_data = {}
    for dataset_name, filepath in eval_files.items():
        if filepath.exists():
            with open(filepath, "r", encoding="utf-8") as f:
                eval_data[dataset_name] = json.load(f)
        else:
            report["errors"].append(f"Source JSON file not found: {filepath}")

    def format_decimal(val, digits=4):
        if val is None:
            return "N/A"
        return f"{val:.{digits}f}".replace(".", ",")

    # 4. Verify Table 0 (Fidelity Table)
    if len(doc.tables) < 1:
        report["errors"].append("Table 0 (Fidelity table) not found.")
    else:
        table_0 = doc.tables[0]
        # Inspect rows
        rows = table_0.rows
        # Expect 10 rows (1 header + 9 data)
        if len(rows) != 10:
            report["errors"].append(f"Table 0 should have 10 rows, but has {len(rows)} rows.")
        
        # Verify Headers
        headers = [c.text.strip() for c in rows[0].cells]
        expected_headers = [
            "Bộ dữ liệu (Dataset)",
            "K",
            "Độ chính xác (Precision@K)",
            "Độ phủ (Recall@K)",
            "Điểm NDCG (NDCG@K)"
        ]
        # De-duplicate cell texts (docx might report merged cell text in multiple cells)
        clean_headers = []
        for h in headers:
            if not clean_headers or h != clean_headers[-1]:
                clean_headers.append(h)
        # Note: clean_headers could contain fewer items if columns are merged, but here headers are simple.
        report["table_0_headers"] = headers
        
        # Map row index to expected dataset and K
        row_mapping = [
            # (row_index, dataset_key, display_name, K)
            (1, "student-mat", "Student-Mat", 1),
            (2, "student-mat", "", 3),
            (3, "student-mat", "", 5),
            (4, "student-por", "Student-Por", 1),
            (5, "student-por", "", 3),
            (6, "student-por", "", 5),
            (7, "xapi", "xAPI", 1),
            (8, "xapi", "", 3),
            (9, "xapi", "", 5)
        ]
        
        for r_idx, dataset_key, display_name, k in row_mapping:
            if r_idx >= len(rows):
                report["errors"].append(f"Table 0 has fewer rows than expected. Missing row {r_idx}")
                continue
            cells = [c.text.strip() for c in rows[r_idx].cells]
            
            # For merged cells, row_cells[0].text might be Student-Mat, Student-Por or xAPI even for K=3, 5
            # but docx generate code does: row_cells[0].text = display_name if idx == 0 else ""
            actual_dataset = cells[0]
            actual_k = cells[1]
            actual_p = cells[2]
            actual_r = cells[3]
            actual_ndcg = cells[4]
            
            # Fetch from JSON
            js_data = eval_data.get(dataset_key, {})
            ranking = js_data.get("ranking", {})
            exp_p = format_decimal(ranking.get(f"precision_at_{k}"))
            exp_r = format_decimal(ranking.get(f"recall_at_{k}"))
            exp_ndcg = format_decimal(ranking.get(f"ndcg_at_{k}"))
            
            p_match = (actual_p == exp_p)
            r_match = (actual_r == exp_r)
            ndcg_match = (actual_ndcg == exp_ndcg)
            
            k_match = (actual_k == str(k))
            dataset_match = (actual_dataset == display_name)
            
            row_report = {
                "row_index": r_idx,
                "dataset": dataset_key,
                "k": k,
                "values": {
                    "dataset_display": {"actual": actual_dataset, "expected": display_name, "match": dataset_match},
                    "k": {"actual": actual_k, "expected": str(k), "match": k_match},
                    "precision": {"actual": actual_p, "expected": exp_p, "match": p_match},
                    "recall": {"actual": actual_r, "expected": exp_r, "match": r_match},
                    "ndcg": {"actual": actual_ndcg, "expected": exp_ndcg, "match": ndcg_match}
                },
                "all_match": p_match and r_match and ndcg_match and k_match and dataset_match
            }
            report["table_0_fidelity"].append(row_report)

    # 5. Verify Table 1 (LLM Judge Table)
    if len(doc.tables) < 2:
        report["errors"].append("Table 1 (LLM Judge table) not found.")
    else:
        table_1 = doc.tables[1]
        rows = table_1.rows
        if len(rows) != 4:
            report["errors"].append(f"Table 1 should have 4 rows, but has {len(rows)} rows.")
            
        expected_headers_1 = ["Bộ dữ liệu (Dataset)", "Trạng thái đánh giá", "Điểm số LLM", "Lý do / Mô tả chi tiết"]
        headers_1 = [c.text.strip() for c in rows[0].cells]
        report["table_1_headers"] = headers_1
        
        row_mapping_1 = [
            (1, "student-mat", "Student-Mat"),
            (2, "student-por", "Student-Por"),
            (3, "xapi", "xAPI")
        ]
        
        for r_idx, dataset_key, display_name in row_mapping_1:
            if r_idx >= len(rows):
                report["errors"].append(f"Table 1 has fewer rows than expected. Missing row {r_idx}")
                continue
            cells = [c.text.strip() for c in rows[r_idx].cells]
            
            actual_dataset = cells[0]
            actual_status = cells[1]
            actual_score = cells[2]
            actual_reason = cells[3]
            
            # Fetch from JSON
            js_data = eval_data.get(dataset_key, {})
            judge = js_data.get("llm_judge", {})
            
            status = judge.get("status", "N/A")
            exp_status = "Chưa thực hiện" if status == "not_run" else status
            
            score = judge.get("score")
            exp_score = format_decimal(score) if score is not None else "N/A"
            
            reason = judge.get("reason", "")
            exp_reason = "Không có dữ liệu đánh giá từ LLM bên ngoài hoặc tập gán nhãn thủ công." if "No external LLM annotations" in reason else reason
            
            dataset_match = (actual_dataset == display_name)
            status_match = (actual_status == exp_status)
            score_match = (actual_score == exp_score)
            reason_match = (actual_reason == exp_reason)
            
            row_report = {
                "row_index": r_idx,
                "dataset": dataset_key,
                "values": {
                    "dataset_display": {"actual": actual_dataset, "expected": display_name, "match": dataset_match},
                    "status": {"actual": actual_status, "expected": exp_status, "match": status_match},
                    "score": {"actual": actual_score, "expected": exp_score, "match": score_match},
                    "reason": {"actual": actual_reason, "expected": exp_reason, "match": reason_match}
                },
                "all_match": dataset_match and status_match and score_match and reason_match
            }
            report["table_1_llm_judge"].append(row_report)

    # 6. Verify Font Styling, alignments
    # Check default style font settings
    try:
        normal_font = doc.styles['Normal'].font
        report["font_style"] = {
            "default_font_name": normal_font.name,
            "default_font_size_pt": normal_font.size.pt if normal_font.size else None
        }
    except Exception as e:
        report["font_style"] = {"error": str(e)}

    # Inspect paragraph properties and inline fonts of key elements
    headings_checked = []
    for p in doc.paragraphs:
        # Check alignment and fonts if paragraph has text and runs
        text = p.text.strip()
        if text.startswith("CHƯƠNG") or text in ["LỜI CẢM ƠN", "LỜI CAM ĐOAN", "MỤC LỤC", "TÀI LIỆU THAM KHẢO", "PHỤ LỤC"]:
            # Heading styling checks
            font_names = []
            font_sizes = []
            is_bold = []
            for r in p.runs:
                if r.font.name:
                    font_names.append(r.font.name)
                if r.font.size:
                    font_sizes.append(r.font.size.pt)
                if r.bold:
                    is_bold.append(r.bold)
            
            headings_checked.append({
                "heading": text[:50],
                "alignment": str(p.alignment) if p.alignment else "None (default LEFT)",
                "fonts": list(set(font_names)),
                "sizes": list(set(font_sizes)),
                "bold": any(is_bold)
            })
    report["headings_style_audit"] = headings_checked

    # Audit table cells styling
    table_cell_styles = []
    for t_idx, table in enumerate(doc.tables):
        sample_cells = []
        # Sample first row and second row cells
        for r_idx in [0, 1]:
            if r_idx < len(table.rows):
                for c_idx in range(min(len(table.rows[r_idx].cells), 3)):
                    cell = table.rows[r_idx].cells[c_idx]
                    p = cell.paragraphs[0] if cell.paragraphs else None
                    font_names = []
                    font_sizes = []
                    is_bold = []
                    if p:
                        alignment = str(p.alignment) if p.alignment else "None (default LEFT)"
                        for r in p.runs:
                            if r.font.name:
                                font_names.append(r.font.name)
                            if r.font.size:
                                font_sizes.append(r.font.size.pt)
                            if r.bold:
                                is_bold.append(r.bold)
                    else:
                        alignment = "No paragraph"
                        
                    sample_cells.append({
                        "row": r_idx,
                        "col": c_idx,
                        "text_preview": cell.text[:30],
                        "alignment": alignment,
                        "fonts": list(set(font_names)),
                        "sizes": list(set(font_sizes)),
                        "bold": any(is_bold) if is_bold else False
                    })
        table_cell_styles.append({
            "table_index": t_idx,
            "samples": sample_cells
        })
    report["table_cells_style_audit"] = table_cell_styles

    # Write report file
    with open("scratch/document_verification_results.json", "w", encoding="utf-8") as f:
        json.dump(report, f, indent=2, ensure_ascii=False)
    print("Verification report written successfully.")

if __name__ == '__main__':
    main()
