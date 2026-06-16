"""Update the current HUFLIT thesis DOCX with verified final artifacts."""

from __future__ import annotations

import json
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph
from sklearn.metrics import classification_report, confusion_matrix


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports" / "final"
SOURCE = ROOT / "scratch" / "report_sources" / "Luan_van_du_doan_va_khuyen_nghi_thanh_tich_hoc_tap.docx"
OUTPUT = REPORTS / "LUAN_VAN_HOAN_CHINH_FINAL.docx"
FIGURES = REPORTS / "figures" / "current"
DATASETS = ("student-mat", "student-por", "xapi")


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def decimal(value: float, digits: int = 4) -> str:
    return f"{value:.{digits}f}".replace(".", ",")


def replace_text(paragraph, text: str) -> None:
    paragraph.text = text


def find_paragraph(doc: Document, prefix: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise KeyError(prefix)


def insert_paragraph_after(paragraph, text: str = "", style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    result = Paragraph(new_p, paragraph._parent)
    if style is not None:
        result.style = style
    result.add_run(text)
    return result


def insert_table_after(document: Document, paragraph, rows: list[list[str]]):
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_index, values in enumerate(rows):
        for col_index, value in enumerate(values):
            cell = table.cell(row_index, col_index)
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_index else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
                    run.bold = row_index == 0
    paragraph._p.addnext(table._tbl)
    return table


def insert_paragraph_after_table(table, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    table._tbl.addnext(new_p)
    result = Paragraph(new_p, table._parent)
    result.add_run(text)
    return result


def replace_picture_before_caption(doc: Document, caption_prefix: str, image_path: Path, width: Cm) -> None:
    candidates = [paragraph for paragraph in doc.paragraphs if paragraph.text.strip().startswith(caption_prefix)]
    for caption in reversed(candidates):
        sibling = caption._p.getprevious()
        checked = 0
        while sibling is not None and checked < 6:
            if sibling.tag.endswith("}p"):
                picture_paragraph = Paragraph(sibling, caption._parent)
                if picture_paragraph._p.xpath(".//w:drawing"):
                    for child in list(picture_paragraph._p):
                        picture_paragraph._p.remove(child)
                    picture_paragraph.add_run().add_picture(str(image_path), width=width)
                    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    return
            sibling = sibling.getprevious()
            checked += 1
    raise RuntimeError(f"No picture found before caption: {caption_prefix}")


def update_table(table, rows: list[list[str]]) -> None:
    while len(table.rows) < len(rows):
        table.add_row()
    while len(table.rows) > len(rows):
        table._tbl.remove(table.rows[-1]._tr)
    for row_index, values in enumerate(rows):
        for col_index, value in enumerate(values):
            cell = table.cell(row_index, col_index)
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_index else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(9.5)
                    run.bold = row_index == 0


def set_table_font(table, size: float) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(size)


def restart_numbered_list(document: Document, paragraphs: list[Paragraph]) -> None:
    style_num_pr = paragraphs[0].style.element.pPr.numPr
    base_num_id = int(style_num_pr.numId.val)
    numbering = document.part.numbering_part.element
    base_num = next(node for node in numbering.findall(qn("w:num")) if int(node.get(qn("w:numId"))) == base_num_id)
    abstract_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    new_num_id = max(int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))) + 1

    new_num = OxmlElement("w:num")
    new_num.set(qn("w:numId"), str(new_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    new_num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    new_num.append(override)
    numbering.append(new_num)

    for paragraph in paragraphs:
        p_pr = paragraph._p.get_or_add_pPr()
        old = p_pr.find(qn("w:numPr"))
        if old is not None:
            p_pr.remove(old)
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id = OxmlElement("w:numId")
        num_id.set(qn("w:val"), str(new_num_id))
        num_pr.extend([ilvl, num_id])
        p_pr.append(num_pr)


def main() -> None:
    doc = Document(SOURCE)
    metrics = {ds: load_json(REPORTS / "metrics" / f"{ds}_3class_locked_test_metrics.json") for ds in DATASETS}
    cv = {ds: load_json(REPORTS / "metrics" / f"{ds}_3class_optuna_cv.json")["f1_macro_best"] for ds in DATASETS}
    recommendation = {ds: load_json(REPORTS / "recommendations" / f"{ds.replace('-', '_')}_evaluation.json") for ds in DATASETS}
    params = {ds: load_json(ROOT / "models" / "saved" / "final" / f"{ds}_3class_best_params.json") for ds in DATASETS}

    paragraph_updates = {
        32: "Nghiên cứu xây dựng hệ thống phân loại thành tích học tập theo ba mức Low, Medium và High, đồng thời sinh lộ trình hỗ trợ cá nhân hóa. Thành phần dự đoán sử dụng kiến trúc hai nhánh: CNN-BiLSTM xử lý chuỗi điểm hoặc chỉ báo tương tác, còn Context MLP xử lý các biến số và biến phân loại mô tả bối cảnh học tập. Hai biểu diễn được hợp nhất để tạo xác suất ba lớp. Quy trình thực nghiệm tách locked test 20% trước tối ưu, fit tiền xử lý và chọn đặc trưng trên dữ liệu huấn luyện, xử lý mất cân bằng bằng Weighted Cross-Entropy và SMOTENC cho dữ liệu hỗn hợp, rồi đánh giá bằng Stratified 5-fold CV và ensemble 11 seed.",
        33: f"Trên locked test, F1-Macro đạt {decimal(metrics['student-mat']['F1-Macro'])} đối với Student-Mat, {decimal(metrics['student-por']['F1-Macro'])} đối với Student-Por và {decimal(metrics['xapi']['F1-Macro'])} đối với xAPI. Accuracy tương ứng là {decimal(metrics['student-mat']['Accuracy'])}, {decimal(metrics['student-por']['Accuracy'])} và {decimal(metrics['xapi']['Accuracy'])}. Chênh lệch giữa CV cố định và locked test lần lượt là {decimal(cv['student-mat']-metrics['student-mat']['F1-Macro'])}, {decimal(cv['student-por']-metrics['student-por']['F1-Macro'])} và {decimal(cv['xapi']-metrics['xapi']['F1-Macro'])}. Kết quả hỗ trợ khả năng khái quát hóa nhưng không đủ để khẳng định tuyệt đối rằng mô hình không overfit.",
        34: "Sau dự đoán, một MLP riêng xếp hạng sáu yếu tố rủi ro và chuyển các rủi ro được chọn thành lộ trình theo giai đoạn bằng thư viện mẫu hành động có cấu trúc. MLP chỉ được huấn luyện trên train pool; locked test chỉ dùng để đánh giá độ trung thành với bộ tiêu chí chuyên môn. PostgreSQL lưu phiên chạy, dự đoán, xác suất, metric và khuyến nghị. Các khuyến nghị chưa phải bằng chứng nhân quả về hiệu quả can thiệp.",
        38: f"This thesis develops a three-class student-performance predictor and an MLP-based risk-ranking module for staged learning-path recommendations. The predictor combines a CNN-BiLSTM branch for short ordered academic or interaction indicators with a context MLP for numerical and categorical attributes. A locked 20% test set is isolated before model development. The verified locked-test F1-Macro scores are {metrics['student-mat']['F1-Macro']:.4f} for Student-Mat, {metrics['student-por']['F1-Macro']:.4f} for Student-Por and {metrics['xapi']['F1-Macro']:.4f} for xAPI. A separate 64-32-6 MLP ranks six observable risk dimensions; its offline metrics measure fidelity to explicit domain criteria, not causal improvement in academic outcomes. PostgreSQL provides operational traceability.",
        178: "SMOTE tạo mẫu tổng hợp bằng nội suy giữa các quan sát lớp thiểu số [5], trong khi ADASYN ưu tiên các vùng khó phân loại [6]. Tuy nhiên, nội suy trực tiếp mã số của biến phân loại không bảo toàn ngữ nghĩa danh mục. Vì vậy, khi cấu hình yêu cầu SMOTE hoặc ADASYN trên dữ liệu hỗn hợp, pipeline hiện tại sử dụng SMOTENC để giữ nguyên miền giá trị phân loại; Weighted Cross-Entropy vẫn được dùng để giảm thiên lệch về lớp phổ biến.",
        180: "Optuna cung cấp không gian tìm kiếm define-by-run, TPE sampler và pruning [7]. Cấu hình mặc định của pipeline là 150 trial cho mỗi bộ dữ liệu, với Stratified 5-fold CV trên train pool. Sau khi sửa resampling hỗn hợp, khóa luận không tuyên bố đã tái tối ưu đủ 150 trial; thay vào đó, các siêu tham số đã khóa được đánh giá lại bằng 5-fold CV và ensemble 11 seed. Locked test tiếp tục được cô lập hoàn toàn.",
        207: "Mô-đun khuyến nghị sử dụng MLP đầu vào 8 đặc trưng đối với Student-Mat/Student-Por và 7 đặc trưng đối với xAPI. Hai tầng ẩn có 64 và 32 nút, đầu ra gồm sáu logit rủi ro. Đặc trưng được chuẩn hóa bằng thống kê của train pool. Mô hình được huấn luyện bằng BCEWithLogitsLoss trên nhãn weak supervision từ bộ tiêu chí chuyên môn có thể kiểm toán; locked test không tham gia huấn luyện.",
        212: "MLP quyết định thứ tự ưu tiên của các yếu tố rủi ro. Phần diễn đạt hành động dùng mẫu nội dung theo từng nhóm rủi ro để bảo đảm câu chữ ổn định và có thể kiểm tra. Do nhãn huấn luyện bắt nguồn từ tiêu chí chuyên môn, metric offline phản ánh khả năng mô phỏng chính sách tham chiếu, không phản ánh tác động thực tế lên điểm số. LLM-Judge chưa được thực hiện vì chưa có bộ chấm độc lập hoặc đánh giá con người được cung cấp.",
        222: "Mỗi tập được tách stratified thành train pool 80% và locked test 20% bằng seed 42. Số mẫu locked test là 79, 130 và 96. Locked test không tham gia fit scaler/encoder, chọn đặc trưng, resampling, chọn epoch hoặc huấn luyện MLP khuyến nghị. Trong từng fold CV, toàn bộ tiền xử lý chỉ fit trên train fold; validation fold chỉ được transform.",
        234: "4.5. ADASYN, SMOTENC và Weighted Loss",
        238: "Pipeline cấu hình 150 Optuna trial và Stratified 5-fold CV. Bộ kết quả cuối trong khóa luận là lần chạy cố định sau khi sửa resampling: dùng siêu tham số đã khóa, đánh giá lại bằng 5-fold CV trên train pool và huấn luyện ensemble 11 seed. Việc tái chạy toàn bộ 150 trial sau thay đổi này chưa được thực hiện; do đó báo cáo không gọi các giá trị CV hiện tại là Optuna-best mới.",
        246: "Chương này chỉ sử dụng metric, prediction và evaluation JSON được tạo lại sau khi sửa resampling và mô-đun khuyến nghị. Các artifact hậu tố v27 và các con số trong bản báo cáo cũ không được dùng. Mọi kết luận vượt quá dữ liệu hiện có tiếp tục được giới hạn hoặc đánh dấu cần bổ sung thực nghiệm.",
        249: f"Student-Mat đạt F1-Macro {decimal(metrics['student-mat']['F1-Macro'])}, Student-Por đạt {decimal(metrics['student-por']['F1-Macro'])}, còn xAPI đạt {decimal(metrics['xapi']['F1-Macro'])}. Student-Mat có RMSE thấp nhất ({decimal(metrics['student-mat']['RMSE'])}) và R² cao nhất ({decimal(metrics['student-mat']['R2'])}). xAPI có F1-Macro và R² thấp nhất, cho thấy giới hạn rõ hơn khi biểu diễn bốn chỉ báo tổng hợp như một chuỗi ngắn.",
        255: f"Cả ba locked-test score thấp hơn CV cố định. Chênh lệch là {decimal(cv['student-mat']-metrics['student-mat']['F1-Macro'])} ở Student-Mat, {decimal(cv['student-por']-metrics['student-por']['F1-Macro'])} ở Student-Por và {decimal(cv['xapi']-metrics['xapi']['F1-Macro'])} ở xAPI. Student-Por có gap lớn nhất; xAPI có gap nhỏ nhất nhưng hiệu năng tuyệt đối thấp nhất. Vì CV này đánh giá cấu hình đã khóa thay vì chọn cấu hình tốt nhất trong 150 trial mới, kết quả chỉ được dùng như phép kiểm tra nhất quán nội bộ.",
        256: "Không bộ dữ liệu nào sụp đổ hoàn toàn trên locked test, nhưng khoảng cách 2,4-5,5 điểm phần trăm vẫn đáng lưu ý. Một phép chia locked test duy nhất không thể chứng minh không overfit. Kết luận phù hợp là mô hình giữ được phần lớn hiệu năng CV, trong khi Student-Por cho thấy độ nhạy khái quát hóa cao nhất trong lần chạy hiện tại.",
        276: "Student-Por có phân bố locked test lệch mạnh về lớp Medium (20/84/26). Ma trận nhầm lẫn mới cho thấy lớp High vẫn được nhận diện tốt, nhưng nhiều mẫu Medium bị chuyển sang Low hoặc High. F1-Macro thấp hơn Accuracy, vì vậy Accuracy đơn lẻ sẽ đánh giá lạc quan chất lượng trên các lớp ít mẫu.",
        277: "Student-Mat và Student-Por yêu cầu ADASYN trong bộ siêu tham số cũ, nhưng pipeline thực thi SMOTENC vì dữ liệu có nhiều biến phân loại. xAPI cũng sử dụng SMOTENC khi cấu hình yêu cầu SMOTE. So sánh chéo dataset không cho phép kết luận phương pháp resampling nào tốt hơn; cần ablation trong từng dataset với cùng seed và kiến trúc.",
        279: f"xAPI đạt CV F1 {decimal(cv['xapi'])} và locked-test F1 {decimal(metrics['xapi']['F1-Macro'])}, chênh {decimal(cv['xapi']-metrics['xapi']['F1-Macro'])}. Dù gap không lớn nhất, xAPI vẫn có F1 tuyệt đối thấp nhất. Cấu hình hiện tại dùng CNN 32 kênh, BiLSTM hidden 96, Context MLP 256 và Fusion 128; việc tăng chiều nhánh context chưa giải quyết được hạn chế của chuỗi bốn chỉ báo tổng hợp.",
        280: "Ba nguyên nhân hợp lý được rút ra từ dữ liệu và mã nguồn. Thứ nhất, bốn chỉ báo xAPI không có timestamp nên thứ tự đưa vào CNN-BiLSTM không tương đương chuỗi thời gian thật. Thứ hai, Label Encoding vẫn tạo quan hệ thứ tự trong biểu diễn danh mục, dù SMOTENC đã ngăn nội suy thành mã phân loại không hợp lệ. Thứ ba, cỡ mẫu nhỏ làm mô hình sâu nhạy với seed và cách chia dữ liệu.",
        298: f"Kết quả locked test hiện tại cho thấy mô hình hoạt động tốt nhất trên Student-Mat (F1-Macro {decimal(metrics['student-mat']['F1-Macro'])}), tiếp theo là Student-Por ({decimal(metrics['student-por']['F1-Macro'])}) và xAPI ({decimal(metrics['xapi']['F1-Macro'])}). Student-Por có CV-test gap lớn nhất. xAPI có hiệu năng tuyệt đối thấp nhất và tiếp tục bộc lộ giới hạn biểu diễn. Các kết luận về lợi ích riêng của resampling hoặc từng nhánh kiến trúc vẫn cần ablation.",
        305: "MLP khuyến nghị đọc feature snapshot, chuẩn hóa theo train pool và trả về xác suất cho sáu yếu tố rủi ro.",
        310: "Trong vận hành, mỗi khuyến nghị cần có trạng thái proposed, accepted, modified, completed hoặc rejected; người thực hiện; ngày bắt đầu; ngày đánh giá; và outcome. Schema hiện tại lưu learning path dưới dạng JSONB nhưng chưa có bảng sự kiện thực thi chi tiết. Đây là phần cần mở rộng để đánh giá hiệu quả can thiệp thay vì chỉ đánh giá độ trung thành với tiêu chí tham chiếu.",
        329: "Các thuộc tính nhạy cảm có thể trở thành proxy cho bất bình đẳng xã hội. Do chưa có fairness audit, mô hình chỉ nên dùng để sàng lọc hỗ trợ. Bộ tiêu chí weak supervision và thư viện mẫu hành động của mô-đun khuyến nghị cần được hội đồng chuyên môn phê duyệt trước khi triển khai.",
        333: "Siêu tham số chưa được tối ưu lại đầy đủ sau khi chuyển sang SMOTENC; Label Encoding vẫn có thể tạo quan hệ thứ tự giả giữa các danh mục.",
        337: "MLP khuyến nghị học từ nhãn weak supervision dựa trên tiêu chí chuyên môn; chưa có đánh giá người dùng, đánh giá LLM độc lập hoặc bằng chứng nhân quả về hiệu quả can thiệp.",
        343: f"Trên locked test, F1-Macro đạt {decimal(metrics['student-mat']['F1-Macro'])} cho Student-Mat, {decimal(metrics['student-por']['F1-Macro'])} cho Student-Por và {decimal(metrics['xapi']['F1-Macro'])} cho xAPI. Student-Mat giữ hiệu năng tốt nhất; Student-Por có generalization gap lớn nhất; xAPI là trường hợp hạn chế nhất về hiệu năng tuyệt đối.",
        344: "MLP Learning Path Engine xếp hạng sáu yếu tố rủi ro và kết hợp chúng với thư viện mẫu hành động theo giai đoạn. PostgreSQL cung cấp khả năng lưu vết. Khuyến nghị vẫn cần human-in-the-loop vì metric hiện tại chỉ đo fidelity với bộ tiêu chí tham chiếu.",
        346: "Tái tối ưu đủ 150 Optuna trial sau thay đổi SMOTENC; bổ sung nested CV hoặc repeated CV để định lượng selection bias.",
        348: "So sánh Label Encoding với embedding hoặc one-hot; thực hiện ablation không resampling, class weight, SMOTENC và các tỷ lệ lấy mẫu khác nhau.",
    }
    paragraph_updates.update({
        32: "Nghiên cứu xây dựng hệ thống phân loại thành tích học tập theo ba mức Low, Medium và High, đồng thời sinh lộ trình hỗ trợ cá nhân hóa. Thành phần dự đoán sử dụng kiến trúc hai nhánh: CNN-BiLSTM xử lý chuỗi điểm hoặc chỉ báo tương tác, còn Context MLP xử lý các biến số và biến phân loại mô tả bối cảnh học tập. Quy trình thực nghiệm tách locked test 20% trước tối ưu, tách validation trước resampling và huấn luyện ensemble 11 seed. Student-Mat và Student-Por sử dụng ADASYN; xAPI sử dụng SMOTENC.",
        33: f"Trên locked test, F1-Macro đạt {decimal(metrics['student-mat']['F1-Macro'])} đối với Student-Mat, {decimal(metrics['student-por']['F1-Macro'])} đối với Student-Por và {decimal(metrics['xapi']['F1-Macro'])} đối với xAPI. Accuracy tương ứng là {decimal(metrics['student-mat']['Accuracy'])}, {decimal(metrics['student-por']['Accuracy'])} và {decimal(metrics['xapi']['Accuracy'])}. Chênh lệch giữa Optuna best CV và locked test lần lượt là {decimal(cv['student-mat']-metrics['student-mat']['F1-Macro'])}, {decimal(cv['student-por']-metrics['student-por']['F1-Macro'])} và {decimal(cv['xapi']-metrics['xapi']['F1-Macro'])}. Kết quả hỗ trợ khả năng khái quát hóa nhưng không đủ để khẳng định tuyệt đối rằng mô hình không overfit.",
        178: "SMOTE tạo mẫu tổng hợp bằng nội suy giữa các quan sát lớp thiểu số [5], trong khi ADASYN ưu tiên các vùng khó phân loại [6]. Final model v1 giữ ADASYN cho Student-Mat và Student-Por; xAPI yêu cầu SMOTE và được thực thi bằng SMOTENC do có các biến phân loại. Weighted loss tiếp tục giảm thiên lệch về lớp phổ biến.",
        180: "Optuna cung cấp không gian tìm kiếm define-by-run, TPE sampler và pruning [7]. Final model v1 sử dụng Optuna best CV để chọn cấu hình, sau đó huấn luyện ensemble 11 seed. Student-Mat và Student-Por có best CV F1 lần lượt 0,9035 và 0,8804; study xAPI lưu trong SQLite có best CV F1 0,8233. Locked test được cô lập khỏi quá trình chọn siêu tham số.",
        238: "Bộ kết quả trong khóa luận thuộc phiên bản final model v1: chọn siêu tham số bằng Optuna, tách validation trước resampling và huấn luyện ensemble 11 seed. Báo cáo sử dụng đúng artifact của phiên bản này, không trộn kết quả từ lần chạy lại với hàm loss hoặc resampling khác.",
        246: "Chương này chỉ sử dụng locked-test metric, prediction, Optuna study và evaluation JSON của final model v1. Các artifact hậu tố v27 và kết quả từ các lần chạy khác cấu hình không được dùng. Mọi kết luận vượt quá dữ liệu hiện có tiếp tục được giới hạn hoặc đánh dấu cần bổ sung thực nghiệm.",
        255: f"Cả ba locked-test score thấp hơn Optuna best CV. Chênh lệch là {decimal(cv['student-mat']-metrics['student-mat']['F1-Macro'])} ở Student-Mat, {decimal(cv['student-por']-metrics['student-por']['F1-Macro'])} ở Student-Por và {decimal(cv['xapi']-metrics['xapi']['F1-Macro'])} ở xAPI. Student-Por có gap lớn nhất; Student-Mat có gap nhỏ nhất. Điểm CV tốt nhất sau tối ưu có thể mang selection bias, nên locked test mới là số liệu chính để báo cáo hiệu năng cuối.",
        277: "Student-Mat và Student-Por sử dụng ADASYN theo siêu tham số được chọn; xAPI sử dụng SMOTENC do không gian siêu tham số chọn SMOTE trên dữ liệu hỗn hợp. So sánh chéo dataset không cho phép kết luận phương pháp resampling nào tốt hơn; cần ablation trong từng dataset với cùng seed và kiến trúc.",
        333: "Việc nội suy các mã danh mục trong ADASYN là một giới hạn phương pháp cần được kiểm tra bằng ablation; thay đổi resampling có thể làm thay đổi đáng kể kết quả và phải được tái tối ưu riêng.",
        346: "Bổ sung nested CV hoặc repeated holdout để định lượng selection bias và độ ổn định của locked-test score.",
        348: "So sánh Label Encoding với embedding hoặc one-hot; thực hiện ablation không resampling, class weight, ADASYN, SMOTENC và các tỷ lệ lấy mẫu khác nhau.",
    })
    for index, text in paragraph_updates.items():
        replace_text(doc.paragraphs[index], text)

    global_replacements = {
        "Rule-based Learning Path Engine": "MLP Learning Path Engine",
        "rule-based learning-path recommendation system": "MLP-based learning-path recommendation system",
        "Rule-based": "MLP-based",
        "rule engine": "MLP risk-ranking engine",
        "Rule engine": "MLP risk-ranking engine",
        "theo luật": "dựa trên MLP",
        "sinh khuyến nghị theo luật": "sinh khuyến nghị từ xếp hạng rủi ro của MLP",
        "Ví dụ ánh xạ luật sang learning path": "Từ điểm rủi ro MLP đến learning path",
        "4.5. SMOTENC và Weighted Cross-Entropy": "4.5. ADASYN, SMOTENC và Weighted Loss",
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text
        for old, new in global_replacements.items():
            text = text.replace(old, new)
        if text != paragraph.text:
            replace_text(paragraph, text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    for old, new in global_replacements.items():
                        text = text.replace(old, new)
                    if text != paragraph.text:
                        replace_text(paragraph, text)

    update_table(
        doc.tables[2],
        [
            ["Đầu ra MLP", "Ý nghĩa", "Mẫu hành động", "Giai đoạn"],
            ["attendance", "Rủi ro chuyên cần", "Xác minh nguyên nhân, học bù, theo dõi chuyên cần", "Tuần 1"],
            ["failure_history / grade_gap", "Lỗ hổng kiến thức", "Chẩn đoán chủ đề yếu và luyện tập có phản hồi", "Tuần 1-2"],
            ["resource_usage / course_updates", "Sử dụng học liệu thấp", "Tăng truy cập LMS và hoàn thành tài nguyên trọng tâm", "Tuần 1-2"],
            ["class_engagement", "Tương tác lớp học thấp", "Đặt câu hỏi, phản hồi và tham gia thảo luận", "Tuần 2-4"],
            ["study_time / time_management", "Nếp tự học chưa ổn định", "Tăng giờ tự học có kế hoạch và điều chỉnh lịch", "Tuần 2-4"],
            ["parent_support / school_support", "Thiếu phối hợp hỗ trợ", "Thống nhất mục tiêu và lịch kiểm tra hằng tuần", "Trong 2 tuần"],
        ],
    )

    update_table(
        doc.tables[4],
        [
            ["Thành phần", "Student-Mat", "Student-Por", "xAPI"],
            ["Số trial ghi nhận", "50", "50", "100"],
            ["CV", "Repeated Stratified 5-fold ×3", "Repeated Stratified 5-fold ×3", "Repeated Stratified 5-fold ×3"],
            ["Epoch tối đa/trial", "50", "50", "80"],
            ["Early stopping patience", "15", "15", "25"],
            ["Sampler", "TPE", "TPE", "TPE đa biến"],
            ["Pruner", "MedianPruner", "MedianPruner", "MedianPruner"],
            ["Mục tiêu", "F1-Macro trung bình", "F1-Macro trung bình", "F1-Macro trung bình"],
        ],
    )

    update_table(
        doc.tables[5],
        [
            ["Siêu tham số", "Student-Mat", "Student-Por", "xAPI"],
            ["Learning rate", decimal(params["student-mat"]["learning_rate"], 8), decimal(params["student-por"]["learning_rate"], 8), decimal(params["xapi"]["learning_rate"], 8)],
            ["Weight decay", f"{params['student-mat']['weight_decay']:.4e}", f"{params['student-por']['weight_decay']:.4e}", f"{params['xapi']['weight_decay']:.4e}"],
            ["Batch size", params["student-mat"]["batch_size"], params["student-por"]["batch_size"], params["xapi"]["batch_size"]],
            ["Resampling thực thi", "ADASYN", "ADASYN", "SMOTENC"],
            ["Sampling ratio", decimal(params["student-mat"]["smote_ratio"]), decimal(params["student-por"]["smote_ratio"]), decimal(params["xapi"]["smote_ratio"])],
            ["k-neighbors", params["student-mat"].get("resampling_k_neighbors", 5), params["student-por"].get("resampling_k_neighbors", 5), params["xapi"].get("resampling_k_neighbors", 5)],
            ["CNN channels", params["student-mat"]["cnn_channels"], params["student-por"]["cnn_channels"], params["xapi"]["cnn_channels"]],
            ["CNN kernel", params["student-mat"].get("cnn_kernel_size", 3), params["student-por"].get("cnn_kernel_size", 3), params["xapi"].get("cnn_kernel_size", 3)],
            ["BiLSTM hidden", params["student-mat"]["lstm_hidden_dim"], params["student-por"]["lstm_hidden_dim"], params["xapi"]["lstm_hidden_dim"]],
            ["Context hidden", params["student-mat"]["context_hidden_dim"], params["student-por"]["context_hidden_dim"], params["xapi"]["context_hidden_dim"]],
            ["Fusion hidden", params["student-mat"]["fusion_hidden_dim"], params["student-por"]["fusion_hidden_dim"], params["xapi"]["fusion_hidden_dim"]],
            ["Dropout", decimal(params["student-mat"]["dropout"]), decimal(params["student-por"]["dropout"]), f"{decimal(params['xapi']['sequence_dropout'])} / {decimal(params['xapi']['context_dropout'])} / {decimal(params['xapi']['fusion_dropout'])}"],
        ],
    )

    update_table(
        doc.tables[6],
        [["Dataset", "Acc.", "Prec.", "Recall", "F1", "RMSE", "R²"]]
        + [[{"student-mat": "Mat", "student-por": "Por", "xapi": "xAPI"}[ds], decimal(metrics[ds]["Accuracy"], 3), decimal(metrics[ds]["Precision-Macro"], 3), decimal(metrics[ds]["Recall-Macro"], 3), decimal(metrics[ds]["F1-Macro"], 3), decimal(metrics[ds]["RMSE"], 3), decimal(metrics[ds]["R2"], 3)] for ds in DATASETS],
    )

    update_table(
        doc.tables[7],
        [["Dataset", "Optuna best CV F1", "Locked-test F1", "Chênh lệch", "Tỷ lệ giữ lại"]]
        + [[ds.replace("student-", "Student-").replace("xapi", "xAPI"), decimal(cv[ds]), decimal(metrics[ds]["F1-Macro"]), decimal(cv[ds] - metrics[ds]["F1-Macro"]), f"{100 * metrics[ds]['F1-Macro'] / cv[ds]:.2f}%".replace(".", ",")] for ds in DATASETS],
    )
    set_table_font(doc.tables[6], 9.0)
    set_table_font(doc.tables[7], 8.5)

    per_class_rows = [["Dataset", "Lớp", "Precision", "Recall", "F1", "Support"]]
    confusion_tables = []
    for ds in DATASETS:
        pred = pd.read_csv(REPORTS / "predictions" / f"{ds}_3class_predictions.csv")
        report = classification_report(pred["True_Label"], pred["Pred_Label"], labels=[0, 1, 2], output_dict=True, zero_division=0)
        display = ds.replace("student-", "Student-").replace("xapi", "xAPI")
        for label_index, label in enumerate(("Low", "Medium", "High")):
            values = report[str(label_index)]
            per_class_rows.append([display, label, decimal(values["precision"], 2), decimal(values["recall"], 2), decimal(values["f1-score"], 2), int(values["support"])])
        cm = confusion_matrix(pred["True_Label"], pred["Pred_Label"], labels=[0, 1, 2])
        confusion_tables.append([["Thật / Dự đoán", "Low", "Medium", "High"]] + [[label, *map(int, cm[index])] for index, label in enumerate(("Low", "Medium", "High"))])
    update_table(doc.tables[8], per_class_rows)
    for table_index, rows in zip((9, 10, 11), confusion_tables):
        update_table(doc.tables[table_index], rows)

    list_index = next(index for index, paragraph in enumerate(doc.paragraphs) if paragraph.text.startswith("Mô hình tạo nhãn, xác suất"))
    restart_numbered_list(doc, doc.paragraphs[list_index:list_index + 6])

    update_table(
        doc.tables[13],
        [
            ["Nội dung", "Tệp nguồn"],
            ["Metric Student-Mat", "reports/final/metrics/student-mat_3class_locked_test_metrics.json"],
            ["Metric Student-Por", "reports/final/metrics/student-por_3class_locked_test_metrics.json"],
            ["Metric xAPI", "reports/final/metrics/xapi_3class_locked_test_metrics.json"],
            ["Optuna best CV", "reports/final/metrics/*_3class_optuna_cv.json"],
            ["Ma trận nhầm lẫn", "reports/final/predictions/*_3class_predictions.csv"],
            ["Đánh giá khuyến nghị", "reports/final/recommendations/*_evaluation.json"],
            ["Kiến trúc dự đoán", "src/models.py"],
            ["Tiền xử lý", "src/data_pipeline.py"],
            ["MLP khuyến nghị", "src/recommendation.py"],
            ["PostgreSQL", "database/schema.sql"],
        ],
    )

    replace_picture_before_caption(doc, "Hình 3.3.", FIGURES / "04_recommendation_mlp.png", Cm(15.5))
    replace_picture_before_caption(doc, "Hình 5.1.", FIGURES / "07_locked_metrics.png", Cm(14.5))
    replace_picture_before_caption(doc, "Hình 5.2.", FIGURES / "08_cv_test_gap.png", Cm(14.0))
    replace_picture_before_caption(doc, "Hình 5.3.", FIGURES / "09_confusion_matrices.png", Cm(16.0))
    replace_picture_before_caption(doc, "Hình 5.4.", FIGURES / "10_per_class_f1.png", Cm(14.0))
    replace_picture_before_caption(doc, "Hình 5.5.", FIGURES / "11_feature_importance.png", Cm(16.0))
    replace_picture_before_caption(doc, "Hình 5.6.", FIGURES / "12_confidence_distribution.png", Cm(14.0))
    replace_picture_before_caption(doc, "Hình 6.1.", FIGURES / "13_risk_band_distribution.png", Cm(14.0))

    anchor = next(
        paragraph
        for paragraph in reversed(doc.paragraphs)
        if paragraph.text.startswith("Kết quả locked test hiện tại")
    )
    anchor = insert_paragraph_after(anchor, "5.10. Đánh giá offline mô-đun khuyến nghị", style="Heading 2")
    anchor.paragraph_format.page_break_before = True
    anchor = insert_paragraph_after(
        anchor,
        "Mô-đun MLP được đánh giá trên locked test bằng Precision@K, Recall@K và NDCG@K so với nhãn tham chiếu từ bộ tiêu chí chuyên môn. Kết quả cao cần được hiểu là độ trung thành với chính sách tham chiếu, không phải hiệu quả can thiệp. LLM-Judge không được báo cáo vì chưa có đánh giá độc lập.",
    )
    caption = insert_paragraph_after(anchor, "Bảng 5.7. Kết quả đánh giá MLP khuyến nghị trên locked test")
    caption.style = "Table Caption"
    rows = [["Dataset", "F1 đa nhãn", "P@3", "R@3", "NDCG@3", "LLM-Judge"]]
    for ds in DATASETS:
        values = recommendation[ds]
        rows.append(
            [
                {"student-mat": "Mat", "student-por": "Por", "xapi": "xAPI"}[ds],
                decimal(values["multilabel"]["f1_macro"], 3),
                decimal(values["ranking"]["precision_at_3"], 3),
                decimal(values["ranking"]["recall_at_3"], 3),
                decimal(values["ranking"]["ndcg_at_3"], 3),
                "Chưa chạy",
            ]
        )
    table = insert_table_after(doc, caption, rows)
    set_table_font(table, 8.5)
    image_paragraph = insert_paragraph_after_table(table)
    image_paragraph.add_run().add_picture(str(FIGURES / "14_recommendation_ranking.png"), width=Cm(14.0))
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure_caption = insert_paragraph_after(image_paragraph, "Hình 5.7. Precision@3, Recall@3 và NDCG@3 của MLP khuyến nghị")
    figure_caption.style = "Figure Caption"
    figure_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.name is None:
                run.font.name = "Times New Roman"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
