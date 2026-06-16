import docx
import sys
from docx.shared import Pt, Cm, Inches

sys.stdout.reconfigure(encoding='utf-8')

def main():
    doc = docx.Document("Bao_cao_cuoi_cung.docx")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    
    # Check Section settings
    print("\n--- Sections Audit ---")
    for idx, s in enumerate(doc.sections):
        print(f"Section {idx}:")
        print(f"  Page Width: {s.page_width.cm:.2f} cm (Expected: 21.00 cm)")
        print(f"  Page Height: {s.page_height.cm:.2f} cm (Expected: 29.70 cm)")
        print(f"  Top Margin: {s.top_margin.cm:.2f} cm (Expected: 3.00 cm)")
        print(f"  Bottom Margin: {s.bottom_margin.cm:.2f} cm (Expected: 3.00 cm)")
        print(f"  Left Margin: {s.left_margin.cm:.2f} cm (Expected: 3.50 cm)")
        print(f"  Right Margin: {s.right_margin.cm:.2f} cm (Expected: 2.00 cm)")
        print(f"  Header Distance: {s.header_distance.cm if s.header_distance else 'N/A'} cm (Expected: 2.25 cm)")
        print(f"  Footer Distance: {s.footer_distance.cm if s.footer_distance else 'N/A'} cm")
        
        # Check Header / Footer
        print(f"  Has Header: {len(s.header.paragraphs) > 0 and any(p.text for p in s.header.paragraphs)}")
        if len(s.header.paragraphs) > 0:
            for p_idx, p in enumerate(s.header.paragraphs):
                if p.text.strip():
                    print(f"    Header {p_idx}: text='{p.text}', align={p.alignment}")
                    for r in p.runs:
                        print(f"      Run: font={r.font.name}, size={r.font.size.pt if r.font.size else 'N/A'}, bold={r.bold}")
                        
        print(f"  Has Footer: {len(s.footer.paragraphs) > 0 and any(p.text for p in s.footer.paragraphs)}")
        if len(s.footer.paragraphs) > 0:
            for p_idx, p in enumerate(s.footer.paragraphs):
                if p.text.strip():
                    print(f"    Footer {p_idx}: text='{p.text}', align={p.alignment}")
                    for r in p.runs:
                        print(f"      Run: font={r.font.name}, size={r.font.size.pt if r.font.size else 'N/A'}, bold={r.bold}")

    # Check Paragraph Line Spacing and Font
    print("\n--- Paragraph Line Spacing and Font Audit (Sample of first 30 non-empty paragraphs) ---")
    count = 0
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        count += 1
        if count > 30:
            break
            
        print(f"\nParagraph {idx} (style={p.style.name}, align={p.alignment}):")
        print(f"  Text: '{text[:80]}...'")
        p_format = p.paragraph_format
        print(f"  Line Spacing: {p_format.line_spacing} (Expected: 1.5)")
        print(f"  Space Before: {p_format.space_before.pt if p_format.space_before else 'N/A'} pt")
        print(f"  Space After: {p_format.space_after.pt if p_format.space_after else 'N/A'} pt")
        print(f"  Left Indent: {p_format.left_indent.inches if p_format.left_indent else 'N/A'} inches")
        
        # Check Runs
        if len(p.runs) == 0:
            print("  No runs (inherits from style)")
        else:
            for r_idx, r in enumerate(p.runs[:5]):
                font_name = r.font.name
                font_size = r.font.size.pt if r.font.size else 'N/A'
                bold = r.bold
                print(f"    Run {r_idx}: text='{r.text.strip().replace(chr(10), ' ')[:30]}...', font={font_name}, size={font_size}, bold={bold}")

    print("\n--- Table Cells Line Spacing and Font Audit ---")
    for t_idx, table in enumerate(doc.tables):
        print(f"\nTable {t_idx}:")
        row = table.rows[1] if len(table.rows) > 1 else table.rows[0]
        for c_idx, cell in enumerate(row.cells[:3]):
            print(f"  Cell {c_idx} first paragraph:")
            p = cell.paragraphs[0]
            p_format = p.paragraph_format
            print(f"    Text: '{p.text.strip()}'")
            print(f"    Line Spacing: {p_format.line_spacing}")
            print(f"    Space Before: {p_format.space_before}")
            print(f"    Space After: {p_format.space_after}")
            for r in p.runs[:2]:
                print(f"      Run: font={r.font.name}, size={r.font.size.pt if r.font.size else 'N/A'}, bold={r.bold}")

if __name__ == "__main__":
    main()
