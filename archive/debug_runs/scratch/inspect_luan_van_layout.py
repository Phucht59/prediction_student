import docx
import sys
from docx.shared import Pt, Cm, Inches

sys.stdout.reconfigure(encoding='utf-8')

def main():
    doc = docx.Document("reports/final/LUAN_VAN_HOAN_CHINH_FINAL.docx")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    
    print("\n--- Sections Audit ---")
    for idx, s in enumerate(doc.sections):
        print(f"Section {idx}:")
        print(f"  Page Width: {s.page_width.cm:.2f} cm")
        print(f"  Page Height: {s.page_height.cm:.2f} cm")
        print(f"  Margins: Top={s.top_margin.cm:.2f} cm, Bottom={s.bottom_margin.cm:.2f} cm, Left={s.left_margin.cm:.2f} cm, Right={s.right_margin.cm:.2f} cm")
        print(f"  Header Distance: {s.header_distance.cm if s.header_distance else 'N/A'} cm")
        print(f"  Footer Distance: {s.footer_distance.cm if s.footer_distance else 'N/A'} cm")
        
        # Check Header / Footer
        has_header = len(s.header.paragraphs) > 0 and any(p.text.strip() for p in s.header.paragraphs)
        print(f"  Has Header: {has_header}")
        if has_header:
            for p_idx, p in enumerate(s.header.paragraphs):
                if p.text.strip():
                    print(f"    Header {p_idx}: text='{p.text}', align={p.alignment}")
                    for r in p.runs:
                        print(f"      Run: font={r.font.name}, size={r.font.size.pt if r.font.size else 'N/A'}, bold={r.bold}")
                        
        has_footer = len(s.footer.paragraphs) > 0 and any(p.text.strip() for p in s.footer.paragraphs)
        print(f"  Has Footer: {has_footer}")
        if has_footer:
            for p_idx, p in enumerate(s.footer.paragraphs):
                if p.text.strip():
                    print(f"    Footer {p_idx}: text='{p.text}', align={p.alignment}")
                    for r in p.runs:
                        print(f"      Run: font={r.font.name}, size={r.font.size.pt if r.font.size else 'N/A'}, bold={r.bold}")

    # Check Normal Style
    if 'Normal' in doc.styles:
        normal = doc.styles['Normal']
        print("\nNormal Style:")
        print(f"  Font Name: {normal.font.name}")
        print(f"  Font Size: {normal.font.size.pt if normal.font.size else 'N/A'} pt")
        print(f"  Line Spacing: {normal.paragraph_format.line_spacing}")
        print(f"  Space After: {normal.paragraph_format.space_after.pt if normal.paragraph_format.space_after else 'N/A'} pt")
        print(f"  Space Before: {normal.paragraph_format.space_before.pt if normal.paragraph_format.space_before else 'N/A'} pt")

if __name__ == "__main__":
    main()
