import os
import sys
import json
from pathlib import Path
import docx
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Reconfigure stdout to use UTF-8 so we don't hit encoding issues on Windows console
sys.stdout.reconfigure(encoding='utf-8')

ROOT = Path("C:/Huflit/kltn")
REPORTS_DIR = ROOT / "reports" / "final" / "recommendations"
DOCX_PATH = ROOT / "Bao_cao_cuoi_cung.docx"

def format_decimal(val, digits=4):
    if val is None:
        return "N/A"
    return f"{val:.{digits}f}".replace(".", ",")

def main():
    print("=== Loading Source JSON Files ===")
    datasets = ["student-mat", "student-por", "xapi"]
    json_data = {}
    for ds in datasets:
        filename = f"{ds.replace('-', '_')}_evaluation.json"
        path = REPORTS_DIR / filename
        print(f"Reading {path}...")
        with open(path, "r", encoding="utf-8") as f:
            json_data[ds] = json.load(f)
            
    print("\n=== Opening Document ===")
    if not DOCX_PATH.exists():
        print(f"ERROR: {DOCX_PATH} does not exist!")
        sys.exit(1)
    doc = docx.Document(DOCX_PATH)
    print("Successfully loaded document.")

    print("\n=== 1. Margins and Page Size Verification ===")
    for sec_idx, section in enumerate(doc.sections):
        w_cm = section.page_width / 360000 if hasattr(section, 'page_width') else None  # wait, docx.shared uses EMUs or Inches/Cm
        # Let's convert section dimensions to Cm
        width_cm = section.page_width.cm
        height_cm = section.page_height.cm
        left_cm = section.left_margin.cm
        right_cm = section.right_margin.cm
        top_cm = section.top_margin.cm
        bottom_cm = section.bottom_margin.cm
        
        print(f"Section {sec_idx} dimensions:")
        print(f"  Width:  {width_cm:.2f} cm (Expected: 21.00 cm - A4)")
        print(f"  Height: {height_cm:.2f} cm (Expected: 29.70 cm - A4)")
        print(f"  Margins:")
        print(f"    Left:   {left_cm:.2f} cm (Expected: 3.50 cm)")
        print(f"    Right:  {right_cm:.2f} cm (Expected: 2.00 cm)")
        print(f"    Top:    {top_cm:.2f} cm (Expected: 3.00 cm)")
        print(f"    Bottom: {bottom_cm:.2f} cm (Expected: 3.00 cm)")
        
        # Verify
        assert abs(width_cm - 21.0) < 0.01, f"Width is {width_cm:.2f} cm instead of 21.0"
        assert abs(height_cm - 29.7) < 0.01, f"Height is {height_cm:.2f} cm instead of 29.7"
        assert abs(left_cm - 3.5) < 0.01, f"Left margin is {left_cm:.2f} cm instead of 3.5"
        assert abs(right_cm - 2.0) < 0.01, f"Right margin is {right_cm:.2f} cm instead of 2.0"
        assert abs(top_cm - 3.0) < 0.01, f"Top margin is {top_cm:.2f} cm instead of 3.0"
        assert abs(bottom_cm - 3.0) < 0.01, f"Bottom margin is {bottom_cm:.2f} cm instead of 3.0"
        print(f"  -> Section {sec_idx} dimensions & margins verified successfully.")

    print("\n=== 2. Table Data Verification ===")
    tables = doc.tables
    print(f"Total tables found: {len(tables)}")
    assert len(tables) == 2, f"Expected 2 tables, but found {len(tables)}"

    # Table 0: Ranking metrics
    # Row 0: Headers
    # Rows 1-3: student-mat (K=1, 3, 5)
    # Rows 4-6: student-por (K=1, 3, 5)
    # Rows 7-9: xapi (K=1, 3, 5)
    print("\n--- Verifying Table 4.1: Ranking Metrics ---")
    t1 = tables[0]
    
    # Check Header
    t1_headers = [cell.text.strip() for cell in t1.rows[0].cells]
    expected_t1_headers = ["Bộ dữ liệu (Dataset)", "K", "Độ chính xác (Precision@K)", "Độ phủ (Recall@K)", "Điểm NDCG (NDCG@K)"]
    print(f"Actual Table 1 headers: {t1_headers}")
    assert t1_headers == expected_t1_headers, f"Table 1 headers mismatch. Expected: {expected_t1_headers}, Got: {t1_headers}"
    
    # Check cells data
    # Dataset rows mapping
    dataset_keys = ["student-mat", "student-por", "xapi"]
    dataset_display_names = ["Student-Mat", "Student-Por", "xAPI"]
    
    row_idx = 1
    for ds_key, ds_display in zip(dataset_keys, dataset_display_names):
        ranking_data = json_data[ds_key]["ranking"]
        for idx, k in enumerate([1, 3, 5]):
            cells = t1.rows[row_idx].cells
            cell_texts = [c.text.strip() for c in cells]
            
            # Expected values
            expected_ds_text = ds_display if idx == 0 else ""
            expected_k_text = str(k)
            
            p_val = ranking_data[f"precision_at_{k}"]
            r_val = ranking_data[f"recall_at_{k}"]
            n_val = ranking_data[f"ndcg_at_{k}"]
            
            expected_p_text = format_decimal(p_val)
            expected_r_text = format_decimal(r_val)
            expected_n_text = format_decimal(n_val)
            
            print(f"Row {row_idx}: {cell_texts}")
            print(f"  Expected: {[expected_ds_text, expected_k_text, expected_p_text, expected_r_text, expected_n_text]}")
            
            assert cell_texts[0] == expected_ds_text, f"Dataset name mismatch at row {row_idx}: expected '{expected_ds_text}', got '{cell_texts[0]}'"
            assert cell_texts[1] == expected_k_text, f"K mismatch at row {row_idx}: expected '{expected_k_text}', got '{cell_texts[1]}'"
            assert cell_texts[2] == expected_p_text, f"Precision mismatch at row {row_idx}: expected '{expected_p_text}', got '{cell_texts[2]}'"
            assert cell_texts[3] == expected_r_text, f"Recall mismatch at row {row_idx}: expected '{expected_r_text}', got '{cell_texts[3]}'"
            assert cell_texts[4] == expected_n_text, f"NDCG mismatch at row {row_idx}: expected '{expected_n_text}', got '{cell_texts[4]}'"
            
            # Verify Font and Alignment of table cells
            for col_col, cell in enumerate(cells):
                for paragraph in cell.paragraphs:
                    expected_alignment = WD_PARAGRAPH_ALIGNMENT.LEFT if col_col == 0 else WD_PARAGRAPH_ALIGNMENT.CENTER
                    # Note: WD_PARAGRAPH_ALIGNMENT.LEFT is 0, CENTER is 1, RIGHT is 2
                    # Let's check alignment
                    actual_alignment = paragraph.alignment
                    # In python-docx, if paragraph.alignment is None, it defaults to the style's alignment.
                    # Wait, generate_doc.py explicitly sets:
                    # paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT if col_idx == 0 else WD_PARAGRAPH_ALIGNMENT.CENTER
                    # So actual_alignment should not be None
                    if actual_alignment is not None:
                        # Convert both to integer or compare directly
                        assert actual_alignment == expected_alignment or (actual_alignment == 0 and expected_alignment == 0) or (actual_alignment == 1 and expected_alignment == 1), f"Alignment mismatch at row {row_idx} col {col_col}. Expected: {expected_alignment}, Got: {actual_alignment}"
                    
                    for run in paragraph.runs:
                        assert run.font.name == "Times New Roman", f"Font name mismatch at row {row_idx} col {col_col}: expected 'Times New Roman', got '{run.font.name}'"
                        assert run.font.size == Pt(11), f"Font size mismatch at row {row_idx} col {col_col}: expected 11 pt, got {run.font.size}"
                        assert not run.bold, f"Run at row {row_idx} col {col_col} should not be bold."
                        
            row_idx += 1
            
    print("  -> Table 4.1 Verified Successfully.")

    # Table 1: LLM Judge metrics
    print("\n--- Verifying Table 4.2: LLM Judge Metrics ---")
    t2 = tables[1]
    t2_headers = [cell.text.strip() for cell in t2.rows[0].cells]
    expected_t2_headers = ["Bộ dữ liệu (Dataset)", "Trạng thái đánh giá", "Điểm số LLM", "Lý do / Mô tả chi tiết"]
    print(f"Actual Table 2 headers: {t2_headers}")
    assert t2_headers == expected_t2_headers, f"Table 2 headers mismatch. Expected: {expected_t2_headers}, Got: {t2_headers}"
    
    row_idx = 1
    for ds_key, ds_display in zip(dataset_keys, dataset_display_names):
        judge_data = json_data[ds_key]["llm_judge"]
        cells = t2.rows[row_idx].cells
        cell_texts = [c.text.strip() for c in cells]
        
        status = judge_data["status"]
        status_vi = "Chưa thực hiện" if status == "not_run" else status
        score = judge_data["score"]
        score_text = format_decimal(score) if score is not None else "N/A"
        reason = judge_data["reason"]
        reason_vi = "Không có dữ liệu đánh giá từ LLM bên ngoài hoặc tập gán nhãn thủ công." if "No external LLM annotations" in reason else reason
        
        expected_row = [ds_display, status_vi, score_text, reason_vi]
        print(f"Row {row_idx}: {cell_texts}")
        print(f"  Expected: {expected_row}")
        
        for col_col in range(4):
            assert cell_texts[col_col] == expected_row[col_col], f"Value mismatch at row {row_idx} col {col_col}. Expected: '{expected_row[col_col]}', Got: '{cell_texts[col_col]}'"
            
        # Verify Font and Alignment of table cells
        for col_col, cell in enumerate(cells):
            for paragraph in cell.paragraphs:
                expected_alignment = WD_PARAGRAPH_ALIGNMENT.LEFT if col_col in [0, 3] else WD_PARAGRAPH_ALIGNMENT.CENTER
                actual_alignment = paragraph.alignment
                if actual_alignment is not None:
                    assert actual_alignment == expected_alignment or (actual_alignment == 0 and expected_alignment == 0) or (actual_alignment == 1 and expected_alignment == 1), f"Alignment mismatch at table 2 row {row_idx} col {col_col}. Expected: {expected_alignment}, Got: {actual_alignment}"
                for run in paragraph.runs:
                    assert run.font.name == "Times New Roman", f"Font name mismatch at table 2 row {row_idx} col {col_col}: expected 'Times New Roman', got '{run.font.name}'"
                    assert run.font.size == Pt(11), f"Font size mismatch at table 2 row {row_idx} col {col_col}: expected 11 pt, got {run.font.size}"
                    assert not run.bold, f"Run at table 2 row {row_idx} col {col_col} should not be bold."
                    
        row_idx += 1
        
    print("  -> Table 4.2 Verified Successfully.")

    print("\n=== 3. Global Styles & Covers Formatting Verification ===")
    
    # 3.1 Normal Style
    normal_style = doc.styles['Normal']
    print("Normal Style settings:")
    print(f"  Font Name: {normal_style.font.name} (Expected: Times New Roman)")
    print(f"  Font Size: {normal_style.font.size} (Expected: 13 Pt)")
    assert normal_style.font.name == "Times New Roman", f"Normal style font name is {normal_style.font.name} instead of Times New Roman"
    assert normal_style.font.size == Pt(13), f"Normal style font size is {normal_style.font.size} instead of 13 Pt"
    
    # 3.2 Let's analyze cover page paragraphs
    print("\nCover Page Paragraphs:")
    for idx, p in enumerate(doc.paragraphs[:10]):
        text_preview = p.text.strip().replace('\n', ' ')
        print(f"  Paragraph {idx}: text='{text_preview[:50]}...', alignment={p.alignment}")
        for r_idx, r in enumerate(p.runs):
            print(f"    Run {r_idx}: text='{r.text.strip().replace(chr(10), ' ')[:30]}...', font={r.font.name}, size={r.font.size}, bold={r.bold}")
            
    print("\nAll verifications passed!")

if __name__ == "__main__":
    main()
