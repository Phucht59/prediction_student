import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')

try:
    doc = docx.Document("reports/final/LUAN_VAN_HOAN_CHINH_FINAL.docx")
    print(f"Total tables: {len(doc.tables)}")
    for idx, t in enumerate(doc.tables):
        print(f"\n--- Table {idx} ---")
        for row_idx, row in enumerate(t.rows[:4]):
            try:
                cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
                dedup_cells = []
                for cell in cells:
                    if not dedup_cells or cell != dedup_cells[-1]:
                        dedup_cells.append(cell)
                print(f"Row {row_idx}: {' | '.join(dedup_cells[:6])}")
            except Exception as e:
                print(f"Row {row_idx} error: {e}")
except Exception as e:
    print(f"Error reading file: {e}")
